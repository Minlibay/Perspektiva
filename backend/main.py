"""
FastAPI backend для автоматического заполнения Плана АУДИТА
с использованием GigaChat API для извлечения данных из документов.
"""

import os
import sys
import shutil
import json
import uuid
import re
from datetime import datetime

sys.stdout.reconfigure(encoding='utf-8', errors='replace')
sys.stderr.reconfigure(encoding='utf-8', errors='replace')
from pathlib import Path
from typing import Optional
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Inches
import tempfile
from PyPDF2 import PdfReader
import openpyxl

app = FastAPI(title="Audit Plan Filler", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Папка для загруженных файлов
UPLOAD_DIR = Path(__file__).parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

_SESSION_ID_RE = re.compile(r"^[a-f0-9]{8,64}$")


def _resolve_session_dir(session_id: Optional[str]) -> tuple[str, Path]:
    """Возвращает (session_id, dir). Если id пустой — генерирует новый.
    Защита от path traversal: id должен быть hex.
    """
    if not session_id:
        session_id = uuid.uuid4().hex
    elif not _SESSION_ID_RE.match(session_id):
        raise HTTPException(status_code=400, detail="Некорректный session_id")
    sdir = UPLOAD_DIR / session_id
    sdir.mkdir(parents=True, exist_ok=True)
    return session_id, sdir

# Папка для результатов
OUTPUT_DIR = Path(__file__).parent / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)


def _session_meta_path(session_id: str) -> Path:
    return UPLOAD_DIR / session_id / "_meta.json"


def _save_session_meta(session_id: str, **fields) -> None:
    """Атомарно мержит и сохраняет meta для сессии."""
    if not _SESSION_ID_RE.match(session_id or ""):
        return
    p = _session_meta_path(session_id)
    existing = {}
    if p.exists():
        try:
            existing = json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            existing = {}
    existing.update(fields)
    if "session_id" not in existing:
        existing["session_id"] = session_id
    p.write_text(json.dumps(existing, ensure_ascii=False, indent=2), encoding="utf-8")


def _load_session_meta(session_id: str) -> dict:
    p = _session_meta_path(session_id)
    if not p.exists():
        return {}
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _list_session_files(sdir: Path) -> list[dict]:
    """Перечень файлов сессии (без _meta.json)."""
    files = []
    for f in sorted(sdir.iterdir()):
        if f.name.startswith("_") or not f.is_file():
            continue
        try:
            st = f.stat()
            files.append({"name": f.name, "size": st.st_size, "mtime": st.st_mtime})
        except OSError:
            continue
    return files

# Модель GigaChat по умолчанию. Из списка доступных на ключе:
#   "GigaChat" (Lite), "GigaChat-Pro", "GigaChat-Max",
#   "GigaChat-2", "GigaChat-2-Pro", "GigaChat-2-Max"  — новое поколение
# Реально используемая модель берётся из gigachat_settings.json (поле "model"),
# с фолбэком на это значение, если в настройках поле пустое.
DEFAULT_GIGACHAT_MODEL = "GigaChat-2-Pro"


def _load_settings() -> dict:
    """Прочитать gigachat_settings.json. Возвращает {} если нет или битый."""
    settings_file = Path(__file__).parent / "gigachat_settings.json"
    if not settings_file.exists():
        return {}
    try:
        with open(settings_file, encoding="utf-8") as f:
            return json.load(f) or {}
    except Exception:
        return {}


def _get_active_model() -> str:
    """Активная модель: из настроек, иначе DEFAULT_GIGACHAT_MODEL."""
    return (_load_settings().get("model") or "").strip() or DEFAULT_GIGACHAT_MODEL

# Глобальный статус обработки — для опроса прогресса из UI
processing_status = {
    "stage": "idle",        # idle | header | verify | done | error
    "current": 0,           # текущий пункт чек-листа
    "total": 0,             # всего пунктов
    "message": "",          # текст для отображения
    "detail": "",           # подробность (файлы, этап)
}


class GigaChatSettings(BaseModel):
    """Настройки подключения к GigaChat"""
    api_key: str
    model: Optional[str] = None


class ProcessingResult(BaseModel):
    """Результат обработки"""
    status: str
    message: str
    extracted_data: Optional[dict] = None
    output_file: Optional[str] = None
    validation: Optional[dict] = None
    analyzed_files: Optional[list[str]] = None
    outputs: Optional[list[dict]] = None  # по одному элементу на каждый обработанный чек-лист


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"


def _extract_checkboxes_from_docx_xml(file_path: str) -> str:
    """Достаёт состояние галочек/крестиков из document.xml (SDT-checkbox и Wingdings-символы).
    Подпись каждой отметки — текст параграфа, в котором лежит чекбокс (без самого глифа)."""
    import zipfile
    from xml.etree import ElementTree as ET

    try:
        with zipfile.ZipFile(file_path) as z:
            with z.open('word/document.xml') as f:
                tree = ET.parse(f)
    except Exception:
        return ""

    root = tree.getroot()

    checked_codes = {"f0fe", "f052", "f0fc", "f0d8", "fe", "52", "fc"}
    unchecked_codes = {"f0a8", "a8", "f06f", "6f"}

    findings = []  # list of (is_checked, label)

    # ElementTree не хранит ссылки на родителей — собираем карту.
    parent_map = {child: parent for parent in root.iter() for child in parent}

    def _find_ancestor(elem, tag_suffix: str):
        cur = parent_map.get(elem)
        while cur is not None and not cur.tag.endswith(tag_suffix):
            cur = parent_map.get(cur)
        return cur

    def _label_text(node, exclude_elem=None) -> str:
        """Собирает текст из поддерева node, исключая поддерево exclude_elem."""
        if node is None:
            return ""
        parts = []
        def walk(n):
            if n is exclude_elem:
                return
            if n.text:
                parts.append(n.text)
            for ch in n:
                walk(ch)
                if ch.tail:
                    parts.append(ch.tail)
        walk(node)
        text = re.sub(r"[☐☑☒]", "", "".join(parts)).strip()
        return re.sub(r"\s+", " ", text)

    def _label_near_checkbox(elem) -> str:
        """Подпись чекбокса. Пробуем по нарастающей: текст параграфа, соседняя
        ячейка строки, ближайший непустой текст в той же строке таблицы."""
        para = _find_ancestor(elem, "}p")
        label = _label_text(para, exclude_elem=elem) if para is not None else ""
        if label:
            return label
        # Ближайший контейнер чекбокса — либо ячейка таблицы (}tc), либо
        # сам SDT, который выступает «в позиции ячейки» строки.
        tc = _find_ancestor(elem, "}tc")
        own_slot = tc if tc is not None else _find_ancestor(elem, "}sdt") or elem
        tr = _find_ancestor(own_slot, "}tr")
        if tr is None:
            return ""
        # Соберём всех «соседей по строке»: tc/sdt — то есть все прямые дети tr,
        # которые могут нести текст (исключая trPr).
        siblings = [c for c in tr if c.tag.endswith("}tc") or c.tag.endswith("}sdt")]
        try:
            ix = siblings.index(own_slot)
        except ValueError:
            ix = -1
        # Берём первую непустую соседнюю «ячейку» справа, затем слева.
        order = siblings[ix + 1 :] + list(reversed(siblings[:ix])) if ix >= 0 else siblings
        for cand in order:
            t = _label_text(cand, exclude_elem=elem)
            if t:
                return t
        return ""

    # 1) SDT checkboxes (Word 2010+) — пробегаем по всем SDT, не только внутри параграфов.
    for sdt in root.iter(f"{{{_W_NS}}}sdt"):
        cb = sdt.find(f".//{{{_W14_NS}}}checkbox")
        if cb is None:
            continue
        checked_el = cb.find(f"{{{_W14_NS}}}checked")
        is_checked = False
        if checked_el is not None:
            v = checked_el.get(f"{{{_W14_NS}}}val", "")
            is_checked = v in ("1", "true")
        label = _label_near_checkbox(sdt)
        findings.append((is_checked, (label or "(без подписи)")[:200]))

    # 2) Wingdings/Symbol-символы галочек/крестиков в w:sym.
    for sym in root.iter(f"{{{_W_NS}}}sym"):
        char = sym.get(f"{{{_W_NS}}}char", "").lower()
        font = sym.get(f"{{{_W_NS}}}font", "").lower()
        if "wingdings" not in font and "symbol" not in font:
            continue
        if char in checked_codes or char in unchecked_codes:
            label = _label_near_checkbox(sym)
            is_checked = char in checked_codes
            findings.append((is_checked, (label or "(без подписи)")[:200]))

    if not findings:
        return ""

    lines = ["=== ОТМЕТКИ (галочки/крестики) В ДОКУМЕНТЕ ==="]
    for checked, label in findings[:80]:
        mark = "☒ ОТМЕЧЕНО" if checked else "☐ НЕ ОТМЕЧЕНО"
        lines.append(f"{mark} | {label}")
    return "\n".join(lines)


_CHECKBOX_RE = re.compile(r"([☒☑✓✔✗✘×])\s*([^\n☒☑✓✔✗✘×]{1,150})")


def _scan_text_for_checkbox_symbols(text: str) -> str:
    """Сканирует текст на символы галочек/крестиков (для случаев когда они приходят как Unicode)."""
    if not text:
        return ""
    findings = []
    for m in _CHECKBOX_RE.finditer(text):
        sym = m.group(1)
        label = m.group(2).strip()
        if not label or len(label) < 2:
            continue
        is_checked = sym in "☒☑✓✔"
        mark = "ОТМЕЧЕНО" if is_checked else "НЕ ОТМЕЧЕНО"
        findings.append(f"{sym} {mark} | {label[:120]}")
    if not findings:
        return ""
    return "=== СИМВОЛЫ-ОТМЕТКИ В ТЕКСТЕ ===\n" + "\n".join(findings[:60])


def _collect_docx_part_text(part) -> list:
    """Собирает параграфы и таблицы из произвольной части docx (body/header/footer)."""
    out = []
    for para in getattr(part, "paragraphs", []):
        if para.text.strip():
            out.append(para.text)
    for table in getattr(part, "tables", []):
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            # В docx объединённые ячейки возвращаются по одному инстансу на каждую
            # колонку сетки — итог: '7 Состав | 7 Состав | 7 Состав | ...'.
            # Сжимаем подряд идущие одинаковые непустые значения до одного.
            deduped = []
            for t in row_texts:
                if deduped and t == deduped[-1]:
                    continue
                deduped.append(t)
            if any(deduped):
                out.append(" | ".join(deduped))
    return out


def _open_docx_or_docm(file_path: str):
    """python-docx падает на .docm (другой content-type main part).
    Перепаковываем в памяти: меняем macroEnabled.main → document.main."""
    suffix = Path(file_path).suffix.lower()
    if suffix != ".docm":
        return Document(file_path)
    import zipfile
    from io import BytesIO
    src_macro = b"application/vnd.ms-word.document.macroEnabled.main+xml"
    dst_main = b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
    buf = BytesIO()
    with zipfile.ZipFile(file_path, "r") as zin, zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(src_macro, dst_main)
            zout.writestr(item, data)
    buf.seek(0)
    return Document(buf)


def extract_text_from_docx(file_path: str) -> str:
    """Извлечение текста из .docx/.docm файла + колонтитулы + распознанные галочки/крестики."""
    doc = _open_docx_or_docm(file_path)
    texts = []

    for section in doc.sections:
        for hdr_attr in ("header", "first_page_header", "even_page_header"):
            hdr = getattr(section, hdr_attr, None)
            if hdr is not None:
                hdr_texts = _collect_docx_part_text(hdr)
                if hdr_texts:
                    texts.append(f"=== КОЛОНТИТУЛ ({hdr_attr}) ===")
                    texts.extend(hdr_texts)

    texts.extend(_collect_docx_part_text(doc))

    for section in doc.sections:
        for ftr_attr in ("footer", "first_page_footer", "even_page_footer"):
            ftr = getattr(section, ftr_attr, None)
            if ftr is not None:
                ftr_texts = _collect_docx_part_text(ftr)
                if ftr_texts:
                    texts.append(f"=== НИЖНИЙ КОЛОНТИТУЛ ({ftr_attr}) ===")
                    texts.extend(ftr_texts)

    body = "\n".join(texts)

    extras = []
    cb_xml = _extract_checkboxes_from_docx_xml(file_path)
    if cb_xml:
        extras.append(cb_xml)
    cb_sym = _scan_text_for_checkbox_symbols(body)
    if cb_sym:
        extras.append(cb_sym)
    if extras:
        body = body + "\n\n" + "\n\n".join(extras)
    return body


_OCR_DPI = int(os.environ.get("OCR_DPI", "200"))
_OCR_LANG = os.environ.get("OCR_LANG", "rus+eng")


def _ocr_pdf_pages(file_path: str, dpi: int = None) -> str:
    """OCR всех страниц PDF через PyMuPDF (рендер) + Tesseract (распознавание)."""
    import fitz
    import pytesseract
    from PIL import Image
    import io

    if dpi is None:
        dpi = _OCR_DPI

    pages_text = []
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        print(f"[ocr] Не удалось открыть {file_path}: {e}")
        return ""

    fname = Path(file_path).name
    page_count = doc.page_count
    try:
        for page_idx, page in enumerate(doc):
            try:
                processing_status["detail"] = f"OCR {fname}: стр {page_idx + 1}/{page_count}"
                pix = page.get_pixmap(dpi=dpi, alpha=False)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                page_text = pytesseract.image_to_string(img, lang=_OCR_LANG)
                if page_text.strip():
                    pages_text.append(f"--- Страница {page_idx + 1} (OCR) ---\n{page_text}")
                print(f"[ocr] {fname} стр.{page_idx+1}/{page_count}: {len(page_text)} симв.")
            except Exception as e:
                print(f"[ocr] Ошибка стр.{page_idx+1} {file_path}: {e}")
    finally:
        doc.close()

    return "\n\n".join(pages_text)


def _ocr_pdf_with_cache(file_path: str) -> str:
    """OCR с кэшированием в <file>.ocr.txt — повторно не запускаем."""
    cache = Path(file_path).with_suffix(Path(file_path).suffix + ".ocr.txt")
    if cache.exists():
        try:
            return cache.read_text(encoding="utf-8")
        except Exception:
            pass
    text = _ocr_pdf_pages(file_path)
    if text:
        try:
            cache.write_text(text, encoding="utf-8")
        except Exception as e:
            print(f"[ocr] Не удалось записать кэш {cache}: {e}")
    return text


def _extract_checkboxes_from_pdf_fields(file_path: str) -> str:
    """Достаёт значения чекбоксов из form fields PDF."""
    try:
        reader = PdfReader(file_path)
        fields = reader.get_fields() or {}
    except Exception:
        return ""
    lines = []
    for name, field in fields.items():
        try:
            ft = field.get('/FT')
            if ft != '/Btn':
                continue
            v = field.get('/V')
            v_str = str(v) if v is not None else ""
            is_checked = v_str not in ("", "/Off", "Off", "None")
            mark = "☒ ОТМЕЧЕНО" if is_checked else "☐ НЕ ОТМЕЧЕНО"
            lines.append(f"{mark} | поле '{name}'")
        except Exception:
            continue
    if not lines:
        return ""
    return "=== ЧЕКБОКСЫ PDF (form fields) ===\n" + "\n".join(lines[:80])


def extract_text_from_pdf(file_path: str) -> str:
    """Извлечение текста из .pdf файла + OCR (если текстовый слой пуст) + чекбоксы."""
    reader = PdfReader(file_path)
    page_count = len(reader.pages)
    texts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            texts.append(text)
    body = "\n".join(texts)

    # Триггер OCR: если на страницу приходится меньше 100 символов — считаем что это скан
    avg_per_page = (len(body) / page_count) if page_count else 0
    if page_count > 0 and avg_per_page < 100:
        print(f"[ocr] {Path(file_path).name}: текстовый слой пуст ({avg_per_page:.0f} симв/стр), запускаю OCR...")
        try:
            ocr_text = _ocr_pdf_with_cache(file_path)
            if ocr_text:
                body = (body + "\n\n" if body else "") + "=== OCR-ТЕКСТ ===\n" + ocr_text
        except Exception as e:
            print(f"[ocr] Ошибка OCR {file_path}: {e}")

    extras = []
    cb_fields = _extract_checkboxes_from_pdf_fields(file_path)
    if cb_fields:
        extras.append(cb_fields)
    cb_sym = _scan_text_for_checkbox_symbols(body)
    if cb_sym:
        extras.append(cb_sym)
    if extras:
        body = body + "\n\n" + "\n\n".join(extras)
    return body


def extract_text_from_xlsx(file_path: str) -> str:
    """Извлечение текста из .xlsx файла"""
    wb = openpyxl.load_workbook(file_path, data_only=True)
    texts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        texts.append(f"=== Лист: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            row_vals = [str(cell) if cell is not None else "" for cell in row]
            if any(row_vals):
                texts.append(" | ".join(row_vals))
    return "\n".join(texts)


def fill_plan_template(template_path: str, source_texts: dict, 
                      extracted_data: dict, output_path: str) -> str:
    """Заполнение шаблона Плана АУДИТА извлечёнными данными"""
    shutil.copy2(template_path, output_path)
    
    doc = Document(output_path)
    
    # Проходим по всем таблицам документа
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    text = para.text.strip()
                    
                    # === ШАПКА (первая таблица, 4 строки) ===
                    if row_idx == 0 and col_idx == 0 and "Наименования Заявителя" in text:
                        # Правая ячейка — поле ввода
                        pass
                    elif "Наименования Заявителя" in text:
                        if "Наименование Заявителя" in extracted_data:
                            para.clear()
                            run = para.add_run(f"Наименования Заявителя: {extracted_data['Наименование Заявителя']}")
                            run.font.size = Pt(10)
                    elif "Вид аудита" in text and len(text) < 30:
                        if "Вид аудита" in extracted_data:
                            para.clear()
                            run = para.add_run(f"Вид аудита: {extracted_data['Вид аудита']}")
                            run.font.size = Pt(10)
                    elif "Даты проведения" in text:
                        if "Даты проведения" in extracted_data:
                            para.clear()
                            run = para.add_run(f"Даты проведения: {extracted_data['Даты проведения']}")
                            run.font.size = Pt(10)
                    elif text == "РЭГ" or (len(text) <= 5 and "РЭГ" in text):
                        if "РЭГ" in extracted_data:
                            para.clear()
                            run = para.add_run(f"РЭГ: {extracted_data['РЭГ']}")
                            run.font.size = Pt(10)
    
    doc.save(output_path)
    return output_path


def extract_checklist_from_template(template_path: str) -> list[dict]:
    """Извлечение структуры чек-листа из шаблона"""
    doc = Document(template_path)
    checklist = []
    
    for table in doc.tables:
        first_row = table.rows[0]
        cells_text = [cell.text.strip() for cell in first_row.cells]
        
        # Ищем таблицу чек-листа (>=5 колонок, заголовок "Область проверки").
        # Шаблон бывает на 5 колонок (старый) или 7 (добавились "Доки" и "№").
        if len(first_row.cells) >= 5 and "Область проверки" in cells_text:
            # Пропускаем заголовок, начинаем с row_idx=1
            for row_idx, row in enumerate(table.rows[1:], start=1):
                area = row.cells[0].text.strip()
                comments = row.cells[1].text.strip()
                problems_hint = row.cells[4].text.strip()
                
                # Извлекаем маркеры ИИ из комментариев
                import re
                ii_markers = re.findall(r'ИИ\d+', comments)
                
                if area:  # Пропускаем пустые строки
                    checklist.append({
                        "row_index": row_idx,
                        "area": area,
                        "comments": comments,
                        "problems_hint": problems_hint,
                        "ii_markers": ii_markers  # ['ИИ1', 'ИИ2'] и т.д.
                    })
            break  # Нашли таблицу чек-листа, больше не ищем
    
    return checklist


def extract_ii_references(template_path: str) -> dict:
    """
    Извлечение всех маркеров ИИ из шаблона и их контекста.
    Возвращает словарь: {"ИИ1": "текст рядом", "ИИ2": "текст рядом", ...}
    """
    import re
    doc = Document(template_path)
    ii_refs = {}
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                # Ищем все маркеры ИИ в ячейке
                markers = re.findall(r'ИИ\d+', text)
                for marker in markers:
                    # Запоминаем контекст — что написано рядом с маркером
                    # Берём предложение/фраза вокруг маркера
                    if marker not in ii_refs:
                        # Чистим текст от лишних пробелов и переносов
                        clean = re.sub(r'\s+', ' ', text)
                        ii_refs[marker] = clean

    return ii_refs


def detect_checklist_type(template_path: str) -> str:
    """
    Детерминированно определяет тип чек-листа по структуре таблиц.
      'plan' — «ИИ -ЧК -План АУДИТА»: есть таблица с заголовком «Область проверки».
      'svod' — «ЧК -Сводный акт»: одна таблица 6 колонок с заголовком №/.../ОК/NOK/
               «Проблемные зоны» и строкой-титулом «СВОДНЫЙ АКТ».
      'unknown' — структура не распознана.
    """
    try:
        doc = _open_docx_or_docm(template_path)
    except Exception as e:
        print(f"[detect] не удалось открыть {template_path}: {e}")
        return "unknown"

    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        if any("Область проверки" in c for c in header_cells):
            return "plan"

    # Признаки Сводного акта: 6 колонок, в шапке «Проблемные зоны», и где-то
    # в первых строках титул «СВОДНЫЙ АКТ».
    for table in doc.tables:
        if not table.rows or len(table.rows[0].cells) < 6:
            continue
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        has_problem_col = any("Проблемные зоны" in c for c in header_cells)
        first_rows_text = " ".join(
            c.text.strip().upper()
            for row in table.rows[:3]
            for c in row.cells
        )
        if has_problem_col and "СВОДНЫЙ АКТ" in first_rows_text:
            return "svod"

    return "unknown"


def _looks_like_svod_doc(text: str) -> bool:
    """Эвристика: похож ли текст на «Сводный акт исследования» (главный проверяемый
    документ для чек-листа Сводного акта)."""
    if not text:
        return False
    t = text.lower()
    return (
        "сводный акт" in t
        or "общие сведения об организации" in t
        or ("наименование продукции" in t and "оквэд" in t)
    )


def extract_checklist_svod(template_path: str) -> list[dict]:
    """
    Извлечение пунктов чек-листа «Сводный акт».
    Структура таблицы: [№ | Наименование | Критерий | ОК | NOK | Проблемные зоны].
    Строки-пункты — те, где в колонке № стоит число (1..N).
    Возвращает список в том же формате, что и план-чек-лист, плюс поля для
    заполнения колонок ОК/NOK (table_index / row_index) и подсказку источника.
    """
    doc = _open_docx_or_docm(template_path)
    checklist = []

    for t_idx, table in enumerate(doc.tables):
        if not table.rows or len(table.rows[0].cells) < 6:
            continue
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        if not any("Проблемные зоны" in c for c in header_cells):
            continue

        for row_idx, row in enumerate(table.rows):
            cells = row.cells
            if len(cells) < 6:
                continue
            num = cells[0].text.strip()
            if not re.match(r"^\d+$", num):
                continue  # пропускаем шапку и пустые строки
            area = cells[1].text.strip()
            criterion = cells[2].text.strip()
            source_hint = cells[5].text.strip()
            if not area:
                continue
            checklist.append({
                "item_no": int(num),
                "table_index": t_idx,
                "row_index": row_idx,
                "area": area,
                # criterion кладём в comments — verify_item_strict печатает «комментарии»
                "comments": criterion,
                "problems_hint": source_hint,
                "source_hint": source_hint,
                "ii_markers": [],  # в Сводном акте маркеров ИИ\d+ нет
            })
        if checklist:
            break  # нашли таблицу чек-листа

    return checklist


def fill_svod_with_checklist(template_path: str, extracted_data: dict, output_path: str,
                              checklist_structure: list[dict] = None) -> str:
    """
    Заполнение «ЧК -Сводный акт»: проставляем отметку в колонке ОК (col3) или
    NOK (col4) и пишем текст в «Проблемные зоны» (col5). Шапки Заявителя в этом
    чек-листе нет — заполняем только строки пунктов по их row_index.
    """
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    checklist_data = extracted_data.get("checklist", [])
    if not checklist_structure:
        checklist_structure = extract_checklist_svod(output_path)

    # Таблица чек-листа — та, что хранит структуру (table_index одинаков у всех пунктов)
    if not checklist_structure:
        doc.save(output_path)
        return output_path
    t_index = checklist_structure[0].get("table_index", 0)
    table = doc.tables[t_index]

    def _set_cell(cell, text):
        for para in cell.paragraphs:
            para.clear()
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10)

    for struct, res in zip(checklist_structure, checklist_data):
        r_idx = struct.get("row_index")
        if r_idx is None or r_idx >= len(table.rows):
            continue
        row = table.rows[r_idx]
        if len(row.cells) < 6:
            continue
        is_ok = bool(res.get("ok")) and not bool(res.get("nok"))
        is_nok = bool(res.get("nok")) and not bool(res.get("ok"))
        # Колонки ОК=3, NOK=4. Ручной пункт (ни ok, ни nok) — обе клетки пустые.
        _set_cell(row.cells[3], "V" if is_ok else "")
        _set_cell(row.cells[4], "V" if is_nok else "")
        reason = (res.get("reason") or res.get("problems") or "").strip()
        if reason:
            _set_cell(row.cells[5], reason)

    doc.save(output_path)
    return output_path


def fill_plan_with_checklist(template_path: str, extracted_data: dict, output_path: str,
                              checklist_structure: list[dict] = None) -> str:
    """
    Заполнение шаблона Плана АУДИТА:
    - Шапка (Заявитель, Вид аудита, Даты, РЭГ)
    - Таблица чек-листа: ТОЛЬКО колонки ОК/NOK (текст НЕ трогаем)
    """
    shutil.copy2(template_path, output_path)
    
    doc = Document(output_path)
    
    checklist_data = extracted_data.get("checklist", [])
    header_data = extracted_data.get("header", {})
    
    # Счётчик заполненных строк чек-листа
    checklist_item_idx = 0
    
    for table in doc.tables:
        # Определяем тип таблицы
        first_row = table.rows[0]
        first_cell_text = first_row.cells[0].text.strip()
        cols_count = len(first_row.cells)
        
        # === ШАПКА (2 колонки) — заполняем ===
        if cols_count == 2:
            for row in table.rows:
                cells_text = [cell.text.strip() for cell in row.cells]
                full_row_text = " ".join(cells_text)
                
                if "Наименования Заявителя" in full_row_text:
                    if "Наименование Заявителя" in header_data:
                        for cell in row.cells:
                            if not cell.text.strip() or len(cell.text.strip()) < 5:
                                for para in cell.paragraphs:
                                    para.clear()
                                    run = para.add_run(header_data["Наименование Заявителя"])
                                    run.font.size = Pt(10)
                                break
                
                elif "Вид аудита" in full_row_text and len(full_row_text) < 50:
                    if "Вид аудита" in header_data:
                        for cell in row.cells:
                            if not cell.text.strip() or len(cell.text.strip()) < 5:
                                for para in cell.paragraphs:
                                    para.clear()
                                    run = para.add_run(header_data["Вид аудита"])
                                    run.font.size = Pt(10)
                                break
                
                elif "Даты проведения" in full_row_text:
                    if "Даты проведения" in header_data:
                        for cell in row.cells:
                            if not cell.text.strip() or len(cell.text.strip()) < 5:
                                for para in cell.paragraphs:
                                    para.clear()
                                    run = para.add_run(header_data["Даты проведения"])
                                    run.font.size = Pt(10)
                                break
                
                elif full_row_text.strip() == "РЭГ" or ("РЭГ" in full_row_text and len(full_row_text) < 20):
                    if "РЭГ" in header_data:
                        for cell in row.cells:
                            if not cell.text.strip() or len(cell.text.strip()) < 3:
                                for para in cell.paragraphs:
                                    para.clear()
                                    run = para.add_run(header_data["РЭГ"])
                                    run.font.size = Pt(10)
                                break
        
        # === ЧЕК-ЛИСТ (>=5 колонок, заголовок "Область проверки") ===
        elif cols_count >= 5 and "Область проверки" in first_cell_text:
            # Пропускаем заголовок (row_idx=0), начинаем с row_idx=1
            for row in table.rows[1:]:
                if checklist_item_idx >= len(checklist_data):
                    break
                
                item = checklist_data[checklist_item_idx]
                
                # Колонки: 0=Область, 1=Комментарии, 2=ОК, 3=NOK, 4=Проблемные зоны
                # ВАЖНО: Трогаем ТОЛЬКО колонки 2 (ОК) и 3 (NOK)
                if len(row.cells) >= 5:
                    # === Колонка ОК — ставим ОДИН маркер ===
                    ok_cell = row.cells[2]
                    # Полная очистка ячейки — удаляем всё содержимое
                    for para in ok_cell.paragraphs:
                        for run in para.runs:
                            run.text = ""
                        para.text = ""
                    
                    # Если OK — записываем маркер
                    if item.get("ok"):
                        # Убедимся что есть хотя бы один параграф
                        if not ok_cell.paragraphs:
                            ok_cell.add_paragraph()
                        para = ok_cell.paragraphs[0]
                        run = para.add_run("☑")
                        run.font.size = Pt(11)
                    
                    # === Колонка NOK — ставим ОДИН маркер ===
                    nok_cell = row.cells[3]
                    # Полная очистка
                    for para in nok_cell.paragraphs:
                        for run in para.runs:
                            run.text = ""
                        para.text = ""
                    
                    # Если NOK — записываем маркер
                    if item.get("nok"):
                        if not nok_cell.paragraphs:
                            nok_cell.add_paragraph()
                        para = nok_cell.paragraphs[0]
                        run = para.add_run("☒")
                        run.font.size = Pt(11)
                    
                    # Колонки 0, 1, 4 НЕ ТРОГАЕМ — текст остаётся оригинальный из шаблона
                
                checklist_item_idx += 1
    
    doc.save(output_path)
    return output_path


def gigachat_preflight(api_key: str, timeout: float = 8.0, model: Optional[str] = None) -> dict:
    """
    Быстрая диагностика доступности GigaChat.
    Шаги:
      1. OAuth-запрос на порт 9443 (ngw.devices.sberbank.ru)
      2. Запрос списка моделей на порт 443 (gigachat.devices.sberbank.ru)
    Возвращает: {ok, stage, detail, models, current_model}.
    """
    import uuid
    active_model = (model or "").strip() or DEFAULT_GIGACHAT_MODEL
    try:
        import httpx
    except ImportError:
        return {"ok": False, "stage": "no_httpx", "detail": "httpx не установлен",
                "models": [], "current_model": active_model}

    result = {"ok": False, "stage": "", "detail": "", "models": [], "current_model": active_model}

    # Шаг 1: OAuth
    try:
        resp = httpx.post(
            "https://ngw.devices.sberbank.ru:9443/api/v2/oauth",
            headers={
                "Authorization": f"Basic {api_key}",
                "RqUID": str(uuid.uuid4()),
                "Content-Type": "application/x-www-form-urlencoded",
            },
            data={"scope": "GIGACHAT_API_PERS"},
            timeout=timeout,
            verify=False,
        )
    except httpx.ConnectTimeout:
        result["stage"] = "oauth_timeout"
        result["detail"] = ("Таймаут подключения к ngw.devices.sberbank.ru:9443. "
                            "Порт OAuth закрыт: firewall/антивирус/VPN/ISP. "
                            "Откройте исходящий TCP на 185.157.96.243:9443.")
        return result
    except httpx.ConnectError as e:
        result["stage"] = "oauth_connect"
        result["detail"] = f"Не удалось установить соединение с 9443: {str(e)[:150]}"
        return result
    except Exception as e:
        result["stage"] = "oauth_error"
        result["detail"] = f"{type(e).__name__}: {str(e)[:200]}"
        return result

    if resp.status_code == 401:
        result["stage"] = "oauth_unauthorized"
        result["detail"] = "Ключ отклонён (HTTP 401). Проверьте корректность авторизационных данных."
        return result
    if resp.status_code >= 400:
        result["stage"] = "oauth_http_error"
        result["detail"] = f"OAuth вернул HTTP {resp.status_code}: {resp.text[:200]}"
        return result

    try:
        access_token = resp.json().get("access_token")
    except Exception as e:
        result["stage"] = "oauth_parse"
        result["detail"] = f"Не удалось распарсить ответ OAuth: {str(e)[:120]}"
        return result

    if not access_token:
        result["stage"] = "oauth_no_token"
        result["detail"] = "OAuth ответил 200, но в ответе нет access_token"
        return result

    # Шаг 2: Список моделей
    try:
        models_resp = httpx.get(
            "https://gigachat.devices.sberbank.ru/api/v1/models",
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=timeout,
            verify=False,
        )
    except Exception as e:
        result["stage"] = "models_request_error"
        result["detail"] = f"Ошибка запроса /models: {type(e).__name__}: {str(e)[:150]}"
        return result

    if models_resp.status_code != 200:
        result["stage"] = "models_http_error"
        result["detail"] = f"/models вернул HTTP {models_resp.status_code}: {models_resp.text[:200]}"
        return result

    try:
        result["models"] = [m.get("id") for m in models_resp.json().get("data", []) if m.get("id")]
    except Exception as e:
        result["stage"] = "models_parse"
        result["detail"] = f"Не удалось распарсить список моделей: {str(e)[:120]}"
        return result

    result["ok"] = True
    result["stage"] = "ready"
    if active_model in result["models"]:
        result["detail"] = f"OAuth и список моделей получены. Текущая модель '{active_model}' доступна."
    else:
        result["detail"] = (f"OAuth работает, но модель '{active_model}' отсутствует в списке доступных. "
                            f"Выберите одну из: {', '.join(result['models'])}.")
        result["ok"] = False
        result["stage"] = "model_not_available"
    return result


def _gigachat_call(api_key: str, system_prompt: str, user_prompt: str,
                   model: str = "GigaChat", temperature: float = 0.0,
                   max_tokens: int = 2000) -> str:
    """Один вызов GigaChat с retry до 3 попыток."""
    import time
    from gigachat import GigaChat

    last_err = None
    for attempt in range(3):
        try:
            gc = GigaChat(credentials=api_key, verify_ssl_certs=False)
            response = gc.chat({
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                "temperature": temperature,
                "max_tokens": max_tokens,
            })
            return response.choices[0].message.content
        except Exception as e:
            last_err = e
            print(f"[gigachat] Попытка {attempt+1} провалилась: {e}")
            time.sleep(1.5 * (attempt + 1))
    raise RuntimeError(f"GigaChat недоступен после 3 попыток: {last_err}")


def _parse_json_response(text: str) -> dict:
    """Извлечь и распарсить JSON из ответа GigaChat.

    Толерантно к частым отклонениям моделей:
    - markdown-обёртка ```json ... ```
    - Python-словарь с одинарными кавычками
    - True/False/None вместо true/false/null
    - висячие запятые перед закрывающей скобкой
    """
    import re, ast
    cleaned = text.strip()
    cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
    cleaned = re.sub(r'\s*```$', '', cleaned)

    start = cleaned.find('{')
    if start < 0:
        raise ValueError(f"JSON не найден: {text[:200]}")

    depth = 0
    in_string = False
    escape = False
    end = -1
    for i in range(start, len(cleaned)):
        ch = cleaned[i]
        if escape:
            escape = False
            continue
        if ch == '\\':
            escape = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                end = i
                break
    if end < 0:
        # Ответ обрезан max_tokens — попробуем восстановить.
        repaired = cleaned[start:]
        if in_string:
            repaired += '"'
        repaired = re.sub(r',\s*$', '', repaired.rstrip())
        repaired += '}' * max(depth, 1)
        try:
            return json.loads(repaired)
        except Exception:
            try:
                fixed = re.sub(r'\bTrue\b', 'true', repaired)
                fixed = re.sub(r'\bFalse\b', 'false', fixed)
                fixed = re.sub(r'\bNone\b', 'null', fixed)
                return json.loads(fixed)
            except Exception:
                pass
        raise ValueError(f"Ответ модели обрезан, восстановить не удалось: {text[:300]}")

    blob = cleaned[start:end + 1]

    try:
        return json.loads(blob)
    except Exception:
        pass

    try:
        result = ast.literal_eval(blob)
        if isinstance(result, dict):
            return result
    except Exception:
        pass

    fixed = re.sub(r'\bTrue\b', 'true', blob)
    fixed = re.sub(r'\bFalse\b', 'false', fixed)
    fixed = re.sub(r'\bNone\b', 'null', fixed)
    fixed = re.sub(r',(\s*[}\]])', r'\1', fixed)
    try:
        return json.loads(fixed)
    except Exception as e:
        raise ValueError(f"JSON не парсится ({e}). Ответ модели: {blob[:400]}") from e


def _find_plan_applicant(plan_text: str) -> str:
    """Достать первое появление 'ОПФ «Название»' в начале Плана, исключив орган по сертификации.

    Возвращает строку как есть, включая ОПФ-опечатки (ПААО, ОООО, АОО) — НЕ нормализуем.
    Орган по сертификации в шаблонах Газпрома — это «ОС СМК ООО «КЦ «Перспектива»»,
    он обычно идёт первым в колонтитуле. Заявитель — следующий.
    """
    if not plan_text:
        return ""
    head = plan_text[:15000]
    # Названия, которые НЕ являются заявителем (это орган по сертификации или служебные).
    blacklist_substrings = ("кц", "перспектива", "ос смк", "орган по сертифик", "интергазсерт")
    # ОПФ: 2-6 заглавных русских букв подряд + пробел + название в кавычках.
    for m in re.finditer(r'\b([А-ЯЁ]{2,6})\s+[«"]([^»"\n]{2,80})[»"]', head):
        candidate = f'{m.group(1)} «{m.group(2)}»'
        low = candidate.lower()
        if any(b in low for b in blacklist_substrings):
            continue
        return candidate
    return ""


_STANDARD_OPFS = {"ооо", "оао", "зао", "пао", "нао", "ао", "ип"}


def _detect_audit_type(all_texts: dict) -> str:
    """Детерминированно определить «Вид аудита» из Приказа ЭГ (приоритет) или Расчёта трудоёмкости.

    Модель часто хватает первое попавшееся слово вида аудита из перечня в Плане,
    игнорируя что галочка не стоит. В Приказе ЭГ формулировка стандартная и однозначная.
    """
    prikaz_text = ""
    raschet_text = ""
    for fname, text in all_texts.items():
        lf = fname.lower()
        if not text:
            continue
        if not prikaz_text and ("приказ" in lf or " эг" in lf or "_эг" in lf or "назнач" in lf):
            prikaz_text = text
        if not raschet_text and ("трудоемкост" in lf or "трудоёмкост" in lf):
            raschet_text = text

    # Паттерны в порядке приоритета (специфичные → общие).
    patterns = [
        (r"внеплановы\w*\s+инспекционн\w*\s+контрол", "Внеплановый инспекционный контроль"),
        (r"(?<!\d)2\s*ик\b|втор\w+\s+инспекционн\w*\s+контрол", "Второй инспекционный контроль"),
        (r"(?<!\d)1\s*ик\b|перв\w+\s+инспекционн\w*\s+контрол", "Первый инспекционный контроль"),
        (r"ресертификационн\w*\s+аудит", "Ресертификационный аудит"),
        (r"расширени\w*\s+области\s+сертификаци", "Расширение области сертификации"),
        (r"дополнительн\w*\s+аудит", "Дополнительный аудит"),
        (r"втор\w+\s+этап\s+первичного", "Второй этап первичного сертификационного аудита"),
        (r"перв\w+\s+этап\s+первичного", "Первый этап первичного сертификационного аудита"),
        (r"сертификационн\w*\s+аудит", "Сертификационный аудит"),
    ]

    for source_text in (prikaz_text, raschet_text):
        if not source_text:
            continue
        low = source_text.lower()
        for rx, label in patterns:
            if re.search(rx, low):
                return label
    return ""


def _find_plan_opf_typos(plan_text: str) -> list[str]:
    """Найти в Плане упоминания ОПФ, которые не входят в стандартный список (ПААО, ОООО, АОО...).

    Возвращает список уникальных 'ОПФ «Название»', отсортирован по первому появлению.
    Использует тот же чёрный список (КЦ/Перспектива/ОС СМК), чтобы не ловить орган по сертификации.
    """
    if not plan_text:
        return []
    blacklist = ("кц", "перспектива", "ос смк", "орган по сертифик", "интергазсерт")
    seen: list[str] = []
    seen_set: set[str] = set()
    for m in re.finditer(r'\b([А-ЯЁ]{2,6})\s+[«"]([^»"\n]{2,80})[»"]', plan_text[:30000]):
        opf_raw = m.group(1)
        name = m.group(2)
        full = f'{opf_raw} «{name}»'
        full_lc = full.lower()
        if any(b in full_lc for b in blacklist):
            continue
        if opf_raw.lower() in _STANDARD_OPFS:
            continue
        # Это нестандартная ОПФ — кандидат на опечатку.
        if full not in seen_set:
            seen_set.add(full)
            seen.append(full)
    return seen


def _find_plan_audit_dates(plan_text: str) -> str:
    """Достать строку дат после 'Сроки проведения аудита' из Плана."""
    if not plan_text:
        return ""
    # Ищем в первых 30k — обычно шапка/таблица идут в начале.
    head = plan_text[:30000]
    idx = head.lower().find("сроки проведения")
    if idx < 0:
        return ""
    window = head[idx:idx + 600]
    # Диапазон вида "17-20.02.2026г." опционально с "г." и несколькими через запятую.
    pattern = r'(\d{1,2}[-–]\d{1,2}\.\d{1,2}\.\d{4}\s*г?\.?(?:\s*,\s*\d{1,2}[-–]\d{1,2}\.\d{1,2}\.\d{4}\s*г?\.?)*)'
    m = re.search(pattern, window)
    if m:
        return m.group(1).strip()
    # Запасной вариант: одиночная дата ДД.ММ.ГГГГ-ДД.ММ.ГГГГ
    m2 = re.search(r'(\d{1,2}\.\d{1,2}\.\d{4}\s*[-–]\s*\d{1,2}\.\d{1,2}\.\d{4})', window)
    if m2:
        return m2.group(1).strip()
    return ""


def extract_header_info(api_key: str, all_texts: dict, model: str = "GigaChat") -> dict:
    """Извлечь данные шапки (Заявитель, Вид аудита, Даты, РЭГ) из пакета документов."""
    # Бюджет на файл: Плану даём больше, потому что в начале повторяется
    # перечень продукции (структура таблицы) и нужные поля «Сроки проведения»,
    # «Состав ЭГ» оказываются глубоко в тексте.
    summaries = []
    for fname, text in all_texts.items():
        is_plan = "план" in fname.lower()
        budget = 20000 if is_plan else 3000
        summaries.append(f"=== {fname} ===\n{text[:budget]}")
    combined = "\n\n".join(summaries)[:120000]

    system_prompt = """Ты извлекаешь 4 поля для ШАПКИ Плана АУДИТА из документов.

Верни ТОЛЬКО JSON:
{
  "Наименование Заявителя": "полное название организации-заявителя на сертификацию СМК (с ОПФ как в тексте — НЕ нормализуй опечатки типа ПААО→ПАО)",
  "Вид аудита": "один из: 'Сертификационный аудит', 'Аудит', 'Расширение', 'Дополнительный аудит', 'Инспекционный контроль'",
  "Даты проведения": "даты проведения аудита из поля 'Сроки проведения аудита' Плана — ВЫПИШИ КАК ЕСТЬ в документе, не нормализуй формат",
  "РЭГ": "регистрационный номер дела (обычно начинается с номера типа 01-01-2025 или подобного)"
}

Правила:
- Если конкретное поле не найдено в документах — напиши "не найдено". НЕ выдумывай данные.
- Для «Даты проведения» принимаются ЛЮБЫЕ форматы: «17-20.02.2026г.», «17.03.2026-27.03.2026», «17 марта — 27 марта 2026», «17.02.2026-20.02.2026, 21.01.2025-24.01.2025» и т.п. Просто скопируй строку из поля «Сроки проведения» как она написана. Если там диапазон или несколько диапазонов — копируй полностью.
- Если в Плане в шапке даты одни, а в таблице день-за-днём — другие, бери из поля «Сроки проведения аудита» (это шапка); противоречие — забота другой проверки."""

    user_prompt = f"Документы:\n\n{combined}"
    response = _gigachat_call(api_key, system_prompt, user_prompt, model=model, max_tokens=800)
    try:
        result = _parse_json_response(response)
    except Exception as e:
        print(f"[header] LLM JSON-парс не удался: {e}")
        result = {}

    # Детерминированные оверрайды из Плана: модель часто нормализует ОПФ
    # (ПААО→ПАО) и пропускает поле дат — оба заменяем сырыми значениями из текста Плана.
    plan_text = ""
    for fname, text in all_texts.items():
        if "план" in fname.lower():
            plan_text = text or ""
            break
    if plan_text:
        det_applicant = _find_plan_applicant(plan_text)
        if det_applicant:
            result["Наименование Заявителя"] = det_applicant
        det_dates = _find_plan_audit_dates(plan_text)
        cur_dates = (result.get("Даты проведения") or "").strip().lower()
        if det_dates and (not cur_dates or cur_dates in ("не найдено", "—", "-")):
            result["Даты проведения"] = det_dates
        print(f"[header] детерминированно: applicant='{det_applicant}', dates='{det_dates}'")

    # Вид аудита — из Приказа ЭГ / Расчёта трудоёмкости (надёжнее, чем модель по Плану).
    det_audit_type = _detect_audit_type(all_texts)
    if det_audit_type:
        result["Вид аудита"] = det_audit_type
        print(f"[header] детерминированно: audit_type='{det_audit_type}'")

    # РЭГ — формат «ОГН<цифра>.RU.<номер> (<суффикс>)». Берём из Приказа ЭГ,
    # фолбэк — любой другой файл (кроме Расчёта трудоёмкости, где встречается
    # «ОГН0.RU.0125 Порядок оплат» — это код регламента, не РЭГ).
    cur_reg = (result.get("РЭГ") or "").strip().lower()
    if not cur_reg or cur_reg in ("не найдено", "—", "-"):
        reg_re = re.compile(
            r"(ОГН\d*\s*\.\s*RU\s*\.\s*\d{2,6})(\s*\([^)]{1,40}\))?",
            re.IGNORECASE,
        )
        det_reg = None
        # 1) Приоритет — файлы с «приказ» или «эг» в имени.
        for fname, text in all_texts.items():
            fn_lc = fname.lower()
            if "трудоемкост" in fn_lc or "трудоёмкост" in fn_lc:
                continue
            if "приказ" in fn_lc or " эг" in f" {fn_lc} " or "_эг" in fn_lc:
                m = reg_re.search(text or "")
                if m:
                    det_reg = (m.group(1) + (m.group(2) or "")).strip()
                    break
        # 2) Фолбэк — остальные файлы (всё ещё пропуская Расчёт).
        if not det_reg:
            for fname, text in all_texts.items():
                fn_lc = fname.lower()
                if "трудоемкост" in fn_lc or "трудоёмкост" in fn_lc:
                    continue
                m = reg_re.search(text or "")
                if m:
                    det_reg = (m.group(1) + (m.group(2) or "")).strip()
                    break
        if det_reg:
            # Нормализуем пробелы внутри «ОГН0 . RU . 0136» → «ОГН0.RU.0136».
            det_reg = re.sub(r"\s*\.\s*", ".", det_reg)
            det_reg = re.sub(r"\s+", " ", det_reg).strip()
            result["РЭГ"] = det_reg
            print(f"[header] детерминированно: РЭГ='{det_reg}'")

    # Дефолты если ничего не нашли
    for key in ("Наименование Заявителя", "Вид аудита", "Даты проведения", "РЭГ"):
        if not (result.get(key) or "").strip():
            result[key] = "не найдено"

    return result


_OPF_EXPANSIONS = [
    ("публичное акционерное общество", "пао"),
    ("непубличное акционерное общество", "нао"),
    ("акционерное общество", "ао"),
    ("общество с ограниченной ответственностью", "ооо"),
    ("закрытое акционерное общество", "зао"),
    ("открытое акционерное общество", "оао"),
    ("индивидуальный предприниматель", "ип"),
]


def _normalize_applicant_name(name: str) -> tuple[str, str]:
    """Нормализует наименование юр.лица.

    Возвращает (opf, core) — токен ОПФ (ооо/ао/пао/...) и «ядро» (название без ОПФ, без кавычек/пробелов).
    Любое отклонение в ОПФ (ПААО vs ПАО) или в ядре (Ромашкаа vs Ромашка) даст разные значения.
    """
    if not name:
        return ("", "")
    s = name.lower().strip()
    # убираем все виды кавычек и лишние знаки
    for ch in ['«', '»', '"', '"', '"', "'", '`', '(', ')', '\n', '\t']:
        s = s.replace(ch, ' ')
    # схлопываем пробелы
    s = re.sub(r'\s+', ' ', s).strip()

    # развёрнутая ОПФ → аббревиатура (сравниваем полные формы до коротких, чтобы «акционерное общество» не съело «публичное акционерное общество»)
    opf = ""
    for full, short in _OPF_EXPANSIONS:
        if s.startswith(full + ' ') or s == full:
            opf = short
            s = s[len(full):].strip()
            break
    if not opf:
        # Если ОПФ-аббревиатура — первое короткое «слово» (2–6 букв) перед названием.
        # Сюда попадает и валидное «ПАО», и опечатки вида «ПААО» — то, что и нужно ловить.
        parts = s.split(' ', 1)
        head = parts[0]
        if 2 <= len(head) <= 6 and head.isalpha():
            opf = head
            s = parts[1] if len(parts) > 1 else ''

    # ядро — буквы и цифры (выкидываем пунктуацию и пробелы для устойчивого сравнения)
    core = re.sub(r'[^0-9a-zа-яё]+', '', s)
    return (opf, core)


def cross_check_applicant_name(api_key: str, plan_text: str, sources_texts: dict,
                                model: str = "GigaChat") -> Optional[dict]:
    """Сравнить наименование юр.лица заявителя между блоком 2 (План аудита) и блоком 3 (источники).

    Возвращает dict {plan_name, sources_name, match, note, plan_norm, sources_norm} либо None.
    Решение о match принимается ДЕТЕРМИНИРОВАННО локально (нормализация ОПФ + ядро),
    модели доверяем только извлечение текста наименования.
    """
    if not plan_text or not sources_texts:
        return None

    plan_excerpt = plan_text[:6000]
    src_summaries = []
    for fname, text in sources_texts.items():
        if not text:
            continue
        src_summaries.append(f"=== {fname} ===\n{text[:2500]}")
    sources_excerpt = "\n\n".join(src_summaries)[:60000]
    if not sources_excerpt:
        return None

    system_prompt = """Ты извлекаешь наименование юридического лица (заявителя СМК) из двух источников.

КРИТИЧНО (главное правило):
Ты НЕ СРАВНИВАЕШЬ и НЕ НОРМАЛИЗУЕШЬ. ЗАПРЕЩЕНО исправлять «ПААО» на «ПАО», «ОООО» на «ООО», «АОО» на «АО» и т.п., даже если ты думаешь что это опечатка. Копируй ОПФ ровно теми буквами, что стоят в тексте. Любая «помощь» в виде нормализации = провал задачи.

Примеры (корректное поведение):
- В тексте «ПААО „Газпром автоматизация“» → plan_name = «ПААО „Газпром автоматизация“» (НЕ «ПАО»).
- В тексте «Общество с ограниченной ответственностью „Ромашка“» → пиши именно так (раскрытие можно оставить — это не нормализация, а копия).
- В тексте «ОООО „Север“» → пиши «ОООО „Север“» (4 буквы О, как в оригинале).

ЗАДАЧА:
1) Из БЛОКА 2 (текст «Плана аудита») выпиши наименование организации-заявителя ДОСЛОВНО (буква в букву, с учётом всех ОПФ-аббревиатур как они есть).
2) Из БЛОКА 3 (Заявка, Договор, Приказ и т.п.) выпиши наименование заявителя ДОСЛОВНО. Если в разных файлах разные написания — выбери то, что ближе к нормальной грамотной форме (без явных опечаток в ОПФ).

Верни ТОЛЬКО JSON:
{
  "plan_name": "точная строка из Плана (или 'не найдено')",
  "sources_name": "точная строка из источников (или 'не найдено')"
}"""

    user_prompt = f"""БЛОК 2 — ПЛАН АУДИТА:
{plan_excerpt}

=====

БЛОК 3 — ИСТОЧНИКИ:
{sources_excerpt}"""

    try:
        response = _gigachat_call(api_key, system_prompt, user_prompt, model=model, max_tokens=300)
        result = _parse_json_response(response)
        if not isinstance(result, dict):
            return None
        plan_name = str(result.get("plan_name") or "").strip()
        sources_name = str(result.get("sources_name") or "").strip()

        # Plan name берём детерминированно из самого текста Плана — модель часто нормализует ОПФ.
        det_plan = _find_plan_applicant(plan_text)
        if det_plan:
            plan_name = det_plan

        # Отдельно ищем в Плане НЕСТАНДАРТНЫЕ ОПФ (ПААО, ОООО, АОО и т.п.) — это явные
        # опечатки, которые надо вынести в примечание независимо от итога сравнения.
        opf_typos = _find_plan_opf_typos(plan_text)
        if opf_typos:
            # Подставляем первую опечатку как plan_name, чтобы сравнение точно дало mismatch.
            plan_name = opf_typos[0]

        not_found = {"не найдено", "не указано", "отсутствует", ""}
        plan_missing = plan_name.lower() in not_found
        src_missing = sources_name.lower() in not_found

        if plan_missing or src_missing:
            match = None
            note = ""
            if plan_missing and src_missing:
                note = "наименование не найдено ни в Плане, ни в источниках"
            elif plan_missing:
                note = "наименование не найдено в Плане"
            else:
                note = "наименование не найдено в источниках"
        else:
            plan_opf, plan_core = _normalize_applicant_name(plan_name)
            src_opf, src_core = _normalize_applicant_name(sources_name)
            opf_match = (plan_opf == src_opf)
            core_match = (plan_core == src_core)
            match = opf_match and core_match
            note = ""
            if not match:
                parts = []
                if not opf_match:
                    parts.append(f"ОПФ отличается: «{plan_opf or '?'}» vs «{src_opf or '?'}»")
                if not core_match:
                    parts.append(f"название отличается: «{plan_core}» vs «{src_core}»")
                note = "; ".join(parts)

        print(f"[applicant-crosscheck] plan='{plan_name}' src='{sources_name}' match={match}")
        return {
            "plan_name": plan_name,
            "sources_name": sources_name,
            "match": match,
            "note": note,
        }
    except Exception as e:
        print(f"[applicant-crosscheck] Ошибка: {e}")
        return None


def classify_files_for_item(api_key: str, item: dict, ii_references: dict,
                             file_names: list, model: str = "GigaChat") -> list:
    """
    Определить, в каких файлах искать данные для конкретного пункта чек-листа.
    Возвращает список имён файлов (подмножество file_names).
    """
    ii_markers = item.get("ii_markers", [])
    ii_context = "\n".join([f"- {m}: {ii_references.get(m, '')[:150]}" for m in ii_markers])

    system_prompt = """Ты определяешь, в каких файлах искать данные для одного пункта чек-листа аудита.

Имена файлов содержат описание содержимого (например "4.1. Приказ об ЭГ..."). Используй это.

Верни ТОЛЬКО JSON:
{
  "relevant_files": ["имя_файла_1", "имя_файла_2", ...]
}

ПРАВИЛА:
- Включи ВСЕ потенциально релевантные файлы (лучше лишний, чем пропущенный).
- Имена файлов должны быть точно такими же, как в списке ниже (копируй буква в букву).
- Максимум 20 файлов.
- Если не уверен какие — включи все файлы с совпадающим номером раздела (например, все файлы "4.x" для пункта раздела 4)."""

    files_list = "\n".join([f"- {f}" for f in file_names])

    user_prompt = f"""ПУНКТ ЧЕК-ЛИСТА:
ОБЛАСТЬ: {item['area']}
КОММЕНТАРИИ: {item['comments']}
ПОДСКАЗКА: {item['problems_hint']}
МАРКЕРЫ ИИ: {', '.join(ii_markers) if ii_markers else 'нет'}

СПРАВОЧНИК МАРКЕРОВ (что они означают):
{ii_context}

ДОСТУПНЫЕ ФАЙЛЫ:
{files_list}"""

    try:
        response = _gigachat_call(api_key, system_prompt, user_prompt, model=model, max_tokens=800)
        result = _parse_json_response(response)
        relevant = result.get("relevant_files", [])
        # Валидация имён (точное совпадение) + case-insensitive фолбэк
        valid = [f for f in relevant if f in file_names]
        if not valid:
            lower_map = {f.lower(): f for f in file_names}
            valid = [lower_map[r.lower()] for r in relevant if r.lower() in lower_map]
        return valid if valid else _fallback_files_by_section(item, file_names)
    except Exception as e:
        print(f"[classify] Ошибка ({item.get('area', '')[:30]}): {e}. Fallback по номеру раздела.")
        return _fallback_files_by_section(item, file_names)


def _section_prefix(s: str) -> str:
    """Достаёт ведущий номер раздела вида '4', '4.1', '4.1.2' из строки.
    Возвращает '' если не найден.
    """
    m = re.match(r'\s*(\d+(?:\.\d+)*)', s or '')
    return m.group(1) if m else ''


# Стоп-слова: служебные/общие, не несут различающего смысла для поиска по аудиту.
_RU_STOPWORDS = {
    "который", "которая", "которое", "которые", "если", "может", "должен", "должна",
    "должны", "должно", "также", "более", "менее", "очень", "когда", "пока",
    "соответствие", "соответствует", "соответствующий", "наличие", "иметь",
    "иметься", "проверка", "проверяется", "пункт", "пункта", "пункту", "область",
    "областью", "комментарий", "комментарии", "проблемы", "проблема", "зона",
    "зоны", "данные", "данных", "документ", "документа", "документы", "документов",
    "файл", "файла", "файлы", "файлов", "требования", "требование", "требуется",
    "только", "нужно", "необходимо", "следует", "часть", "части", "случае",
    "процесс", "процесса", "вопрос", "вопросы", "наличие", "необходимость",
    "осуществляется", "проведение", "проведения", "проводится", "выполнение",
    "выполнения", "выполняется", "включая", "относится", "связан", "связана",
    "перечень", "перечня", "форма", "формы", "образец", "образца",
}


def _extract_keywords(*texts: str) -> set:
    """Достаёт значимые слова (≥4 символов) из переданных строк, без стоп-слов."""
    out = set()
    for t in texts:
        for w in re.findall(r'[А-Яа-яЁёA-Za-z][А-Яа-яЁёA-Za-z\-]{3,}', t or ''):
            wl = w.lower()
            if wl in _RU_STOPWORDS:
                continue
            out.add(wl)
    return out


def find_relevant_files_for_item(item: dict, all_texts: dict) -> list:
    """
    Детерминированный подбор файлов под пункт чек-листа.
    Сигналы:
      A) Номер раздела пункта совпадает с префиксом имени файла,
         либо встречается отдельной "лексемой" в тексте файла.
      B) ≥2 ключевых слов из area/comments/problems_hint встречаются в тексте файла.
    Файл попадает в результат, если сработал хотя бы один сигнал.
    Если ничего не найдено — возвращает все файлы (лучше избыточно, чем пропустить).
    Возвращает список, отсортированный по убыванию релевантности.
    """
    item_section = _section_prefix(item.get('area', ''))
    item_top = item_section.split('.')[0] if item_section else ''
    section_re = re.compile(rf'(?<!\d){re.escape(item_section)}(?!\d)') if item_section else None

    keywords = _extract_keywords(
        item.get('area', ''), item.get('comments', ''), item.get('problems_hint', '')
    )

    scored = []
    for fname, text in all_texts.items():
        score = 0
        # Сигнал A: совпадение по номеру раздела
        fname_section = _section_prefix(fname)
        if item_section and fname_section:
            if fname_section == item_section:
                score += 3
            elif fname_section.startswith(item_section + '.') or item_section.startswith(fname_section + '.'):
                score += 2
            elif item_top and fname_section.split('.')[0] == item_top:
                score += 1
        # Номер раздела упомянут в самом тексте файла
        if section_re and text and section_re.search(text):
            score += 1

        # Сигнал B: совпадение ключевых слов в тексте
        if keywords and text:
            text_lower = text.lower()
            hits = sum(1 for kw in keywords if kw in text_lower)
            if hits >= 2:
                score += min(hits // 2, 3)  # ограничиваем вклад чтобы не перебивать раздел

        if score > 0:
            scored.append((fname, score))

    scored.sort(key=lambda x: -x[1])
    if scored:
        return [f for f, _ in scored]
    return list(all_texts.keys())


def _fallback_files_by_section(item: dict, file_names: list) -> list:
    """Фолбэк-выбор файлов: те, чей префикс номера совпадает с разделом пункта.
    Если совпадений нет — возвращает все файлы (лучше избыточно, чем пропустить).
    """
    item_section = _section_prefix(item.get('area', ''))
    if not item_section:
        return list(file_names)
    # Берём верхний уровень раздела ('4.1' -> '4', '4' -> '4')
    top = item_section.split('.')[0]
    matched = [f for f in file_names if _section_prefix(f).split('.')[0] == top]
    return matched if matched else list(file_names)


def build_evidence_pack(relevant_files: list, all_texts: dict, max_chars: int = 180000) -> str:
    """
    Собрать пак доказательств из релевантных файлов.
    Бюджет делится поровну между файлами, минимальный потолок снят —
    при большом числе файлов лучше короткий фрагмент из каждого, чем пропуск.
    """
    parts = []
    if not relevant_files:
        return ""
    per_file_budget = max(max_chars // len(relevant_files), 2000)
    for fname in relevant_files:
        if fname not in all_texts:
            continue
        text = all_texts[fname]
        chunk = text[:per_file_budget]
        parts.append(f"=== ФАЙЛ: {fname} ===\n{chunk}")
    return "\n\n".join(parts)


def verify_item_strict(api_key: str, item: dict, ii_references: dict,
                        evidence: str, model: str = "GigaChat",
                        extra_instructions: str = "",
                        comparison: bool = False,
                        compare_fields: str = "",
                        expected_file_keywords: Optional[list] = None) -> dict:
    """
    Строгая проверка пункта чек-листа с NOK-first логикой.
    Возвращает: {ok, nok, reason, ii_data_found, evidence_quote, source_file}.
    """
    ii_markers = item.get("ii_markers", [])
    ii_context = "\n".join([f"- {m}: {ii_references.get(m, '')[:150]}" for m in ii_markers])

    system_prompt = """Ты — аудитор СМК, который САМ выполняет проверку по одному пункту чек-листа. Не ищи "доказательства, что кто-то уже проверил". Твоя задача — выполнить сверку между документами и вынести вердикт.

КАК РАБОТАТЬ:
1. Прочитай требование пункта.
2. Найди в загруженных документах фактические данные, относящиеся к этому требованию (даты, номера, ФИО, пункты СТО, состав ЭГ, перечень площадок, расчёты и т.п.).
3. Если требование подразумевает СРАВНЕНИЕ (например, "сравнить состав ЭГ в Плане аудита и Приказе") — сам сравни данные между файлами.
4. Понимай суть, а не буквальные совпадения: "Объект заявителя" может встречаться как "объекты сертификации", "контролируемые объекты", в табличной шапке и т.п. Маркеры ИИ1, ИИ2 — указатели где искать, отсутствие самого слова "ИИ1" не критично.

ЧИСЛОВЫЕ И ВРЕМЕННЫЕ УСЛОВИЯ:
Если требование содержит конкретный порог ("не позднее чем за 5 рабочих дней", "не менее 8 часов", "за 14 дней до...") — ВЫПОЛНИ РАСЧЁТ:
1. Найди в документах обе даты/числа.
2. Рассчитай фактическую разницу (для дней — учитывай только Пн–Пт, исключай Сб и Вс).
3. Сравни с порогом. Если факт нарушает порог — это NOK, даже если данные присутствуют.
В reason обязательно покажи расчёт: "Дата утверждения 03.12.2025, начало аудита 12.12.2025, между ними 7 рабочих дней — соответствует требованию ≥5".

КОГДА СТАВИТЬ OK:
- Нужные данные присутствуют в документах И (если требуется сверка) совпадают между собой.
- Если есть числовой порог — фактическое значение его удовлетворяет.
- Семантически близкие формулировки засчитываются как совпадение, если смысл тот же.
- Допустимы небольшие расхождения форматирования (даты "12.12.2025" vs "12 декабря 2025", регистр, пробелы).

КОГДА СТАВИТЬ NOK:
- Нужных данных НЕТ ни в одном из загруженных файлов.
- Данные есть, но числовой порог нарушен (например, дата утверждения за 3 рабочих дня вместо требуемых ≥5).
- Между файлами явное и существенное противоречие (другая дата, другие ФИО, другой номер, разный состав).
- Документ, на который ссылается требование (например, "Приказ об ЭГ"), не загружен.

ВАЖНО: если данные есть, числовые пороги выполнены и нет противоречий — это OK. Не требуй дополнительных "доказательств соответствия".

КРИТИЧНО (часто нарушаемое правило):
Все файлы, нужные для проверки, УЖЕ переданы тебе ниже в секции «ДОКУМЕНТЫ-ИСТОЧНИКИ ДЛЯ ПРОВЕРКИ». Каждый файл начинается с маркера `=== ФАЙЛ: <имя> ===`. Любой файл, имя которого ты видишь в этих маркерах, — присутствует и его текст следует за маркером.
ЗАПРЕЩЕНО писать «файл не передан», «не удалось найти файл», «файл отсутствует», «среди переданных документов нет файла X», если этот файл присутствует в списке маркеров `=== ФАЙЛ: ... ===`.
Если данных конкретно ВНУТРИ текста файла нет — пиши «в файле X нет данных о ...» (фокус на содержимом), а не «файла нет».

Верни ТОЛЬКО JSON (без markdown-обёртки, без комментариев):
{
  "ok": true/false,
  "nok": true/false,
  "evidence_quote": "конкретные данные/цитаты из файлов (до 350 символов), на основе которых вынесен вердикт",
  "source_file": "имя файла (или несколько через запятую) откуда взяты данные, или 'не найдено'",
  "reason": "обоснование (1-3 предложения): какие данные нашёл, как сверил, почему такой вердикт",
  "ii_data_found": "список найденных маркеров через запятую (ИИ1, ИИ2...), или 'не найдено'",
  "extracted_values": []
}

Поле "extracted_values" заполняй ТОЛЬКО если в дополнительных правилах указан "РЕЖИМ СВЕРКИ". Иначе оставляй пустой массив [].

ok и nok — взаимоисключающие: ровно один true, другой false."""

    extra_block = ("\n\nДОПОЛНИТЕЛЬНЫЕ ПРАВИЛА ДЛЯ ЭТОГО ПУНКТА:\n" + extra_instructions) if extra_instructions else ""

    comparison_block = ""
    if comparison:
        fields_hint = compare_fields or "ключевые реквизиты пункта"
        comparison_block = f"""

РЕЖИМ СВЕРКИ (ОБЯЗАТЕЛЬНО):
Этот пункт требует ПРЯМОЙ СВЕРКИ значений между двумя или более файлами ({fields_hint}).
Прежде чем вынести вердикт — ВЫПИШИ дословные цитаты из КАЖДОГО файла-источника. Не перефразируй, не сокращай, копируй как есть. Это твоё доказательство, что ты реально посмотрел в документы.

ФОРМАТ В ПОЛЕ extracted_values (массив объектов):
[
  {{ "field": "название поля (напр. '№ договора', 'дата договора', 'ФИО эксперта 1')",
     "value_in_plan": "дословная цитата из Плана аудита",
     "value_in_other": "дословная цитата из второго файла",
     "other_file": "имя второго файла",
     "match": true/false }},
  ...
]

ПРАВИЛА СОВПАДЕНИЯ:
- Номера документов и даты — сравнивай ПОСИМВОЛЬНО (после удаления пробелов и приведения к одному формату даты). Любое расхождение хотя бы в одном символе номера или одной цифре даты → match: false.
- ФИО — нормализуй к виду "Фамилия И.О." и сравнивай. Разный человек, опечатка в фамилии, другой инициал → match: false. Допустимо только различие "Иванов И.И." vs "Иванов Иван Иванович".
- Если значение НЕ НАЙДЕНО в одном из файлов — match: false, в недостающее поле напиши "не найдено".

АНТИ-ГАЛЛЮЦИНАЦИЯ (КРИТИЧНО):
Цитаты `value_in_plan` и `value_in_other` должны быть СКОПИРОВАНЫ БУКВА-В-БУКВУ из текста соответствующего файла, который тебе передан в evidence ниже. НЕ выдумывай номера, даты, ФИО. НЕ догадывайся по контексту или из общих знаний. НЕ "примерно", не "вероятно", не "по аналогии с другим документом".
Если в тексте файла данного значения НЕТ — пиши ровно "не найдено" (без подмены похожим из памяти или из соседнего файла).
Любая выдуманная цитата = провал проверки. Система будет постфактум искать каждую твою цитату в тексте файла. Если не найдёт — пункт будет помечен как фальсификация.

ВЕРДИКТ:
- OK — ТОЛЬКО если ВСЕ строки extracted_values имеют match: true.
- NOK — если хотя бы одна строка имеет match: false, ИЛИ если массив extracted_values пуст (не смог извлечь — значит проверку не провёл).

В reason кратко перечисли поля и результат сверки: "№ договора: 2604 (План) vs 2605 (Договор) — НЕ совпадает; дата: 26.04.2024 vs 26.04.2024 — совпадает".
"""

    markers_str = ", ".join(ii_markers) if ii_markers else "нет"
    user_prompt = f"""ПУНКТ ДЛЯ ПРОВЕРКИ:
ОБЛАСТЬ: {item['area']}
ТРЕБОВАНИЯ (КОММЕНТАРИИ В ШАБЛОНЕ): {item['comments']}
ПОДСКАЗКА (ПРОБЛЕМНЫЕ ЗОНЫ): {item['problems_hint']}
МАРКЕРЫ ИИ В ПУНКТЕ: {markers_str}

СПРАВОЧНИК МАРКЕРОВ ИИ:
{ii_context if ii_context else '(нет маркеров)'}

ДОКУМЕНТЫ-ИСТОЧНИКИ ДЛЯ ПРОВЕРКИ:
{evidence if evidence else '(файлы не переданы)'}
{extra_block}{comparison_block}

Проведи проверку: сам сравни данные между файлами. OK — если данные есть и согласуются. NOK — только если данных нет или есть явное противоречие."""

    try:
        response = _gigachat_call(api_key, system_prompt, user_prompt, model=model, max_tokens=3500)
        result = _parse_json_response(response)
    except ValueError as e:
        # Чаще всего — ответ обрезан на полпути или модель не закрыла JSON.
        # Один ретрай с повышенным лимитом и явным напоминанием.
        print(f"[verify_item] JSON парсинг упал ({str(e)[:100]}), ретрай с max_tokens=6000")
        retry_user = (
            user_prompt
            + "\n\nВНИМАНИЕ: предыдущий ответ был невалидным JSON или оборвался. "
              "Верни КОРОТКИЙ и ПОЛНОСТЬЮ валидный JSON. Закрой все скобки. "
              "Не приводи лишних цитат — только обязательные поля."
        )
        try:
            response = _gigachat_call(api_key, system_prompt, retry_user, model=model, max_tokens=6000)
            result = _parse_json_response(response)
        except (ValueError, Exception) as e2:
            print(f"[verify_item] Ретрай тоже упал ({str(e2)[:100]}). Graceful NOK.")
            return {
                "ok": False,
                "nok": True,
                "reason": (
                    "[авто-NOK: модель дважды вернула невалидный JSON. "
                    f"Последняя ошибка: {str(e2)[:120]}. "
                    "Содержательную проверку провести не удалось — рекомендуется ручная сверка пункта.]"
                ),
                "evidence_quote": "",
                "source_file": "—",
                "ii_data_found": ", ".join(ii_markers) if ii_markers else "нет",
                "extracted_values": [],
            }

    # Нормализация: ровно один из ok/nok должен быть true.
    # При неоднозначности — NOK (нет уверенного подтверждения).
    ok = bool(result.get("ok", False))
    nok = bool(result.get("nok", False))
    if ok == nok:
        reason_text = (result.get("reason", "") + " " + result.get("evidence_quote", "")).lower()
        nok_signals = ["не найден", "отсутств", "не соответ", "противоречи", "не совпад", "нет данных",
                       "не обнаруж", "не указан", "не подтвержд", "не загружен", "не предостав",
                       "нарушен", "меньше", "позже", "позднее"]
        has_nok_signal = any(sig in reason_text for sig in nok_signals)
        ok_signals = ["соответствует", "удовлетворяет", "совпада", "выполнен", "подтвержд"]
        has_ok_signal = any(sig in reason_text for sig in ok_signals)
        if has_ok_signal and not has_nok_signal:
            ok, nok = True, False
        else:
            ok, nok = False, True
    # Жёсткая пост-валидация для пунктов в режиме сверки.
    # Цель — поймать случаи, когда модель ставит OK без реального извлечения значений
    # или выдумывает (галлюцинирует) цитаты.
    if comparison:
        extracted = result.get("extracted_values") or []

        def _norm(s: str) -> str:
            # Нормализация для поиска подстроки: убираем пробелы, ё→е, lower.
            return re.sub(r"\s+", "", (s or "").lower().replace("ё", "е"))

        evidence_norm = _norm(evidence or "")

        if ok and (not isinstance(extracted, list) or len(extracted) == 0):
            ok, nok = False, True
            result["reason"] = ("[авто-NOK: режим сверки требует извлечь и сравнить значения, "
                                "но модель не вернула extracted_values] "
                                + (result.get("reason") or ""))
        elif isinstance(extracted, list) and extracted:
            mismatches = []
            hallucinations = []
            wrong_files = []
            kw_norm = [k.lower() for k in (expected_file_keywords or []) if k]
            # Слова, по которым понимаем что other_file — это «План аудита» (его не считаем нужным источником)
            plan_markers = ("план", "plan")

            # Детерминированная нормализация для пункта по договору:
            # модель ставит match=false из-за форматных различий ('17 60-25' vs '17-60-25',
            # '03.09.25' vs '03.09.2025'), хотя правило явно велит нормализовать.
            is_contract_item = any("договор" in k or "соглашен" in k or "контракт" in k for k in kw_norm)
            if is_contract_item:
                _MONTHS_RU = {
                    "январ": 1, "феврал": 2, "март": 3, "апрел": 4, "ма": 5, "июн": 6,
                    "июл": 7, "август": 8, "сентябр": 9, "октябр": 10, "ноябр": 11, "декабр": 12,
                }

                def _norm_contract_num(s: str) -> str:
                    return re.sub(r"[^0-9a-zа-я]", "", (s or "").lower().replace("ё", "е"))

                def _parse_date(s: str) -> Optional[tuple]:
                    s = (s or "").strip().lower().replace("ё", "е")
                    if not s:
                        return None
                    m = re.search(r"(\d{1,2})[.\-/_ ](\d{1,2})[.\-/_ ](\d{2,4})", s)
                    if m:
                        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
                        if y < 100:
                            y += 2000
                        return (d, mo, y)
                    m = re.search(r"(\d{1,2})\s+([а-я]+)\s+(\d{2,4})", s)
                    if m:
                        d, mon_word, y = int(m.group(1)), m.group(2), int(m.group(3))
                        for stem, mo in _MONTHS_RU.items():
                            if mon_word.startswith(stem):
                                if y < 100:
                                    y += 2000
                                return (d, mo, y)
                    return None

                def _contract_row_equivalent(row: dict) -> bool:
                    a = (row.get("value_in_plan") or "").strip()
                    b = (row.get("value_in_other") or "").strip()
                    if not a or not b:
                        return False
                    field_lc = (row.get("field") or "").lower()
                    if "дата" in field_lc:
                        da, db = _parse_date(a), _parse_date(b)
                        return da is not None and da == db
                    # Номера: сравниваем по «значащим» символам.
                    na, nb = _norm_contract_num(a), _norm_contract_num(b)
                    return bool(na) and na == nb

                normalized_any = False
                for row in extracted:
                    if isinstance(row, dict) and row.get("match") is False and _contract_row_equivalent(row):
                        row["match"] = True
                        row["_normalized_override"] = True
                        normalized_any = True

                # Если после нормализации все строки match=true, а модель поставила NOK
                # из-за форматного расхождения — переворачиваем вердикт.
                if normalized_any and nok and all(
                    (r.get("match") is True) for r in extracted if isinstance(r, dict)
                ):
                    ok, nok = True, False
                    result["reason"] = (
                        "Номер/дата договора совпадают (различия только в форматировании — "
                        "пробелы/разделители/двузначный vs четырёхзначный год)."
                    )

            for row in extracted:
                if not isinstance(row, dict):
                    continue
                # 1) Проверка на галлюцинации: каждая непустая цитата должна быть в evidence.
                # Кроме «технических» меток модели о том, что данные в файле отсутствуют —
                # это служебные фразы, а не цитаты, и их в evidence искать бессмысленно.
                _ABSENCE_MARKERS = (
                    "не найдено", "не указано", "не указан", "не указаны",
                    "отсутств", "не содержит", "нет данных", "нет в файле",
                    "не приведен", "не приведено", "не приведены",
                )
                for key in ("value_in_plan", "value_in_other"):
                    raw = (row.get(key) or "").strip()
                    raw_lc = raw.lower()
                    if not raw or raw_lc in ("не найдено", "—", "-", "n/a", "нет"):
                        continue
                    if any(m in raw_lc for m in _ABSENCE_MARKERS):
                        continue
                    if _norm(raw) and _norm(raw) not in evidence_norm:
                        hallucinations.append(
                            f"{row.get('field', 'поле')} [{key}]: '{raw[:80]}' нет в тексте файлов"
                        )
                # 2) Сбор расхождений.
                if row.get("match") is False:
                    field = row.get("field", "поле")
                    a = (row.get("value_in_plan") or "").strip()
                    b = (row.get("value_in_other") or "").strip()
                    other = row.get("other_file", "")
                    mismatches.append(f"{field}: '{a}' (План) vs '{b}' ({other})")
                # 3) Проверка, что other_file — именно ожидаемый источник (а не подменённый
                #    моделью «что есть на руках»).
                if kw_norm:
                    other_name = (row.get("other_file") or "").strip().lower()
                    if not other_name or other_name in ("не найдено", "—", "-", "n/a", "нет"):
                        # Источник не назван — это проблема только если value_in_other непустое
                        if (row.get("value_in_other") or "").strip().lower() not in ("", "не найдено"):
                            wrong_files.append(f"{row.get('field', 'поле')}: other_file пуст")
                    else:
                        # other_file должен содержать хотя бы одно ключевое слово ИЛИ быть Планом
                        # (если поле сравнивается «План vs План» — норм). Но обычно нужен внешний.
                        is_expected = any(k in other_name for k in kw_norm)
                        is_plan = any(p in other_name for p in plan_markers) and "плана" not in other_name
                        if not is_expected and not is_plan:
                            wrong_files.append(
                                f"{row.get('field', 'поле')}: сверка с '{other_name[:60]}' "
                                f"вместо ожидаемого источника по ключам {kw_norm[:3]}"
                            )

            if hallucinations:
                ok, nok = False, True
                result["reason"] = (
                    "[авто-NOK: галлюцинация цитаты — значения не найдены в тексте файлов] "
                    + "; ".join(hallucinations[:5])
                    + (" | " + result.get("reason", "") if result.get("reason") else "")
                )
            elif wrong_files:
                ok, nok = False, True
                result["reason"] = (
                    "[авто-NOK: нужный файл-источник не загружен — модель сверила с тем что нашла] "
                    + "; ".join(wrong_files[:5])
                    + (" | " + result.get("reason", "") if result.get("reason") else "")
                )
            elif mismatches and ok:
                ok, nok = False, True
                result["reason"] = ("[авто-NOK: расхождения по сверке] "
                                    + "; ".join(mismatches[:5])
                                    + (" | " + result.get("reason", "") if result.get("reason") else ""))

    # Пост-проверка для пункта 1 (даты утверждения/начала).
    # Распознаём по специфичной фразе из требования.
    area_text = (item.get("area") or "") + " " + (item.get("comments") or "")
    if ("утверждения" in area_text.lower() or "согласования" in area_text.lower()) and ok:
        reason_text = result.get("reason") or ""
        reason_lower = reason_text.lower()

        # 1) Жёсткая проверка по самопризнанному числу: если модель сама написала
        #    "N рабочих дней" с N > 5 — форс-NOK независимо от того, как она это
        #    интерпретирует ("но требование выполнено" — нет, требование нарушено).
        num_match = re.search(
            r"(\d{1,3})\s*(?:рабочих)?\s*раб?(?:очих)?\.?\s*д(?:н(?:[еия][йвм]?)?|\.)",
            reason_lower
        )
        if not num_match:
            # Запасной паттерн: "разница ... N ..."
            num_match = re.search(r"разниц\w*[^0-9]{1,40}(\d{1,3})", reason_lower)
        if num_match:
            try:
                n_days = int(num_match.group(1))
                if n_days > 5:
                    ok, nok = False, True
                    # Достаём даты из reason для чистого вывода, выкидываем противоречивые
                    # рассуждения модели («7 рабочих дней соответствует требованию ≤5» и т.п.).
                    found_dates = re.findall(r"\b\d{1,2}[.\-]\d{1,2}[.\-]\d{2,4}\b", reason_text)
                    dates_str = ", ".join(found_dates[:2]) if len(found_dates) >= 2 else ""
                    clean_reason = (
                        f"[авто-NOK] Разница между датой утверждения Плана и началом аудита — "
                        f"{n_days} рабочих дней, что превышает допустимое окно (≤5). "
                        f"Требование: утвердить не более чем за 5 рабочих дней до начала."
                    )
                    if dates_str:
                        clean_reason += f" Найденные даты: {dates_str}."
                    result["reason"] = clean_reason
            except ValueError:
                pass

        # 2) Если расчёт вообще не приведён (нет дат и/или слова "разница") — также NOK.
        if ok:
            date_re = re.compile(r"\b\d{1,2}[.\-/]\d{1,2}[.\-/]\d{2,4}\b")
            dates_found = date_re.findall(reason_text)
            has_diff_word = any(w in reason_lower for w in ("разниц", "рабочих дн", "рабочих д.", "раб.дн"))
            if len(dates_found) < 2 or not has_diff_word:
                ok, nok = False, True
                result["reason"] = (
                    "[авто-NOK: для пункта 1 в reason обязаны быть обе даты (D_утв, D_начало) "
                    "и явная разница в рабочих днях. Модель не привела расчёт.] "
                    + reason_text
                )

    # Пост-проверка: модель иногда заявляет «файл не передан / не удалось найти файл X»,
    # хотя файл реально был в evidence. Это галлюцинация — подменяем формулировку,
    # чтобы пользователь видел истинную причину (модель отказалась читать содержимое).
    reason_check = (result.get("reason") or "").lower()
    file_missing_phrases = [
        "не удалось найти файл",
        "файл не передан",
        "не передан файл",
        "файл отсутствует среди",
        "файла нет среди",
        "файл не загружен",
        "среди переданных документов нет файла",
        "среди переданных документов не найден",
    ]
    if any(p in reason_check for p in file_missing_phrases):
        evidence_files = re.findall(r"===\s*ФАЙЛ:\s*([^=\n]+?)\s*===", evidence or "")
        if evidence_files:
            files_list = "; ".join(evidence_files[:6])
            ok, nok = False, True
            result["reason"] = (
                f"[авто-NOK: модель заявила, что файл не передан, но он реально был в evidence "
                f"({files_list}). Это означает, что модель не прочитала содержимое — данные могут быть, "
                f"но проверку она не провела.] "
                + (result.get("reason") or "")
            )

    result["ok"] = ok
    result["nok"] = nok
    result.setdefault("reason", result.get("evidence_quote", ""))
    result.setdefault("ii_data_found", ", ".join(ii_markers) if ii_markers else "нет")
    result.setdefault("evidence_quote", "")
    result.setdefault("source_file", "не найдено")
    return result


def adversarial_recheck(api_key: str, item: dict, ii_references: dict,
                         evidence: str, prior_verdict: dict,
                         model: str = "GigaChat",
                         comparison: bool = False,
                         compare_fields: str = "") -> dict:
    """
    Второй проход для OK-вердиктов: попытаться найти причины перевести в NOK.
    Это основная защита от ложных OK.
    """
    ii_markers = item.get("ii_markers", [])

    if comparison:
        fields_hint = compare_fields or "ключевые реквизиты"
        system_prompt = f"""Ты — второй аудитор, который перепроверяет вердикт OK первого аудитора по пункту, требующему ПРЯМОЙ СВЕРКИ значений между файлами ({fields_hint}).

Твоя задача в этом режиме — НЕ подтверждать на доверии, а ИСКАТЬ расхождения, которые первый мог пропустить или замаскировать.

ОБЯЗАТЕЛЬНО ВЫПОЛНИ:
1. Найди в документах-источниках КАЖДОЕ значение, которое заявил первый аудитор в своём массиве extracted_values (если массив был).
2. Сравни их сам, посимвольно для номеров/дат и в нормализованной форме "Фамилия И.О." для ФИО.
3. Если первый аудитор НЕ предоставил extracted_values — это уже основание для NOK (он не доказал, что реально сверил).
4. Если хотя бы одно значение не совпадает (другой номер, другая цифра в дате, другая фамилия, другой инициал, другой человек) — NOK.
5. Если значение есть только в одном из файлов, а во втором отсутствует — NOK.

ДЕФОЛТ В РЕЖИМЕ СВЕРКИ: NOK. Подтверждай OK ТОЛЬКО если ты сам перепроверил все значения и они идентичны.

Верни ТОЛЬКО JSON:
{{
  "ok": true/false,
  "nok": true/false,
  "evidence_quote": "конкретные значения, которые ты сравнил, через '|' (напр. 'План: №2604 от 26.04.2024 | Договор: №2605 от 26.04.2024')",
  "source_file": "имена файлов через запятую",
  "reason": "если NOK — что именно не совпало. Если OK — какие значения ты сверил и подтвердил.",
  "ii_data_found": "..."
}}

ВАЖНО ПРО REASON: пиши для клиента, не для аудитора-коллеги. ЗАПРЕЩЕНО упоминать «первый аудитор», «второй аудитор», «перепроверка», «адверсариальный», «модель сказала», «коллега», ссылаться на процесс проверки или на предыдущий вердикт. Только итог: что именно совпало или не совпало, со ссылкой на значения из файлов."""
    else:
        system_prompt = """Ты — второй аудитор, который делает критический ревью вердикта OK от первого аудитора. Твоя задача — объективно подтвердить OK или перевернуть в NOK, если есть СУЩЕСТВЕННЫЕ основания.

СУЩЕСТВЕННЫЕ основания для NOK (только такие!):
1. Первый аудитор сослался на файл/цитату, которой на самом деле нет в документах-источниках ниже.
2. Явное противоречие между разными документами, которое первый пропустил.
3. Отсутствие критичного реквизита, без которого пункт не может считаться выполненным (например, нет подписи на приказе, нет даты на документе, не тот номер).
4. Данные относятся к ДРУГОМУ аудиту/заявителю/периоду.
5. Фальсификация цитаты первым аудитором.

Несущественные причины (НЕ переворачивать OK в NOK):
- "Мог бы быть более подробный" документ — не причина.
- "Первый аудитор сформулировал обоснование коротко" — не причина.
- "Не все маркеры ИИ подтверждены" — если найдены основные, этого достаточно.
- "Теоретически где-то мог бы быть нюанс" — спекуляция, не переворачивай.

ДЕФОЛТ: подтверждать OK. NOK — только при чётких существенных основаниях, которые ты можешь указать цитатой из файла.

Верни ТОЛЬКО JSON (та же схема):
{
  "ok": true/false,
  "nok": true/false,
  "evidence_quote": "...",
  "source_file": "...",
  "reason": "если OK — 'Подтверждено: ...'. Если NOK — укажи КОНКРЕТНОЕ существенное основание.",
  "ii_data_found": "..."
}

ВАЖНО ПРО REASON: пиши для клиента-заявителя, не для аудитора-коллеги. ЗАПРЕЩЕНО упоминать «первый аудитор», «второй аудитор», «перепроверка», «адверсариальный», «модель сказала», «коллега», ссылаться на процесс проверки или на предыдущий вердикт. Только итог: что именно подтверждено или что именно не выполнено, со ссылкой на содержимое файлов."""

    prior_extracted = prior_verdict.get("extracted_values") or []
    try:
        import json as _json
        prior_extracted_str = _json.dumps(prior_extracted, ensure_ascii=False, indent=2) if prior_extracted else "(не предоставлены)"
    except Exception:
        prior_extracted_str = "(не предоставлены)"

    extracted_block = f"\nИЗВЛЕЧЁННЫЕ ИМ ЗНАЧЕНИЯ (extracted_values):\n{prior_extracted_str}\n" if comparison else ""

    user_prompt = f"""ПУНКТ:
ОБЛАСТЬ: {item['area']}
ТРЕБОВАНИЯ: {item['comments']}
ПОДСКАЗКА: {item['problems_hint']}
МАРКЕРЫ ИИ: {', '.join(ii_markers) if ii_markers else 'нет'}

ВЕРДИКТ ПЕРВОГО АУДИТОРА: OK
ЕГО ОБОСНОВАНИЕ: {prior_verdict.get('reason', '')}
ЦИТАТА, КОТОРУЮ ОН ПРИВЁЛ: {prior_verdict.get('evidence_quote', '')}
ФАЙЛ-ИСТОЧНИК: {prior_verdict.get('source_file', '')}{extracted_block}

ДОКУМЕНТЫ-ИСТОЧНИКИ (проверь их ещё раз):
{evidence if evidence else '(нет)'}

Найди причины поменять на NOK. Если не нашёл — подтверди OK."""

    try:
        response = _gigachat_call(api_key, system_prompt, user_prompt, model=model, max_tokens=3500)
        result = _parse_json_response(response)

        ok = bool(result.get("ok", False))
        nok = bool(result.get("nok", False))
        if ok == nok:
            # Неоднозначный ответ ревью.
            # В режиме сверки — переворачиваем в NOK (дефолт строгий).
            # В обычном режиме — оставляем исходный OK.
            if comparison:
                ok, nok = False, True
                result["reason"] = ("[авто-NOK: ревью в режиме сверки не дало однозначного подтверждения] "
                                    + (result.get("reason") or ""))
            else:
                ok, nok = True, False
        result["ok"] = ok
        result["nok"] = nok
        result.setdefault("reason", prior_verdict.get("reason", ""))
        result.setdefault("ii_data_found", prior_verdict.get("ii_data_found", ""))
        result.setdefault("evidence_quote", prior_verdict.get("evidence_quote", ""))
        result.setdefault("source_file", prior_verdict.get("source_file", "не найдено"))
        return result
    except Exception as e:
        print(f"[adversarial] Ошибка: {e}. Оставляю исходный вердикт OK.")
        # Если ревью не прошло (сетевая ошибка, парсинг) — доверяем первому аудитору
        return prior_verdict


# Правила обработки пунктов чек-листа.
# Ключ — индекс пункта (0 = пункт 1 в плане), значение — словарь:
#   file_keywords: список подстрок (lowercase) для поиска в имени файла; файл попадает в evidence если совпала любая.
#   include_template: добавлять ли текст самого "ПЛАН аудита" в evidence (по умолчанию True для всех пунктов в правилах).
#   extra_instructions: блок специфичных правил, дописывается в промпт verify_item_strict.
ITEM_RULES = {
    0: {
        "file_keywords": ["план"],
        "extra_instructions": (
            "Этот пункт проверяет САМ документ \"ПЛАН аудита\" (включён в evidence как файл \"ПЛАН АУДИТА (проверяемый документ)\").\n\n"
            "ТРЕБОВАНИЕ (точная формулировка из шаблона): дата утверждения/согласования в шапке Плана должна быть НЕ ПОЗДНЕЕ 5 рабочих дней до начала аудита. "
            "Это понимается как окно: разница между датой утверждения и датой начала аудита НЕ ДОЛЖНА ПРЕВЫШАТЬ 5 рабочих дней. "
            "Слишком ранее утверждение (>5 рабочих дней до начала) — тоже нарушение.\n"
            "Пример из шаблона: если начало аудита 10.02.2026, дата утверждения должна быть не раньше 03.02.2026 и не позже 10.02.2026.\n\n"
            "АЛГОРИТМ (выполни ВСЕ шаги по порядку):\n"
            "ШАГ 1. Найди в шапке/колонтитуле Плана дату утверждения/согласования (рядом со словами \"утверждаю\", \"утверждено\", \"согласовано\", \"УТВ\", \"СОГЛ\"). Запиши её как D_утв.\n"
            "ШАГ 2. Найди дату НАЧАЛА аудита (поле \"Сроки проведения аудита\"). Если указан диапазон дат — бери самую раннюю. Запиши как D_начало.\n"
            "ШАГ 3. Если D_утв или D_начало отсутствуют → NOK.\n"
            "ШАГ 4. Если D_утв >= D_начало (утверждение в день начала или ПОСЛЕ) → NOK. Утверждать после начала аудита недопустимо.\n"
            "ШАГ 5. Посчитай число РАБОЧИХ дней (Пн–Пт, без Сб/Вс) от D_утв до D_начало включительно ПО ОДНУ из границ "
            "(т.е. число рабочих дней между ними + 1, считая саму D_начало).\n"
            "ШАГ 6. Если рабочих дней > 5 → NOK (утверждение слишком рано, выходит за окно). "
            "Если рабочих дней <= 5 (и при этом D_утв < D_начало) → OK.\n\n"
            "В reason ОБЯЗАТЕЛЬНО приведи обе даты и явный расчёт в формате: "
            "\"D_утв = ДД.ММ.ГГГГ, D_начало = ДД.ММ.ГГГГ, разница N рабочих дней. Требование: не позднее 5 рабочих дней до начала аудита (окно ≤5). Вердикт: OK/NOK.\"\n"
            "Если расчёт не приведён — это автоматически считается ошибкой проверки."
        ),
    },
    1: {
        # Имена файлов в этой области часто сокращают: «ОС ПАО ...» = Орг.Структура ПАО.
        # Включаем варианты: 'ос ', 'ос_', 'ос-' (как начало имени), плюс полные формы.
        "file_keywords": [
            "оргструктур", "орг.структур", "орг структур", "организационн",
            "ос ", "ос_", "ос-", "ос.",
        ],
        "extra_instructions": (
            "Орг.структура — это графическая схема (квадраты с названиями подразделений). "
            "После OCR получается плоский поток названий в произвольном порядке, часто с искажениями. "
            "Это НОРМАЛЬНО.\n\n"
            "ВЫПОЛНИ РОВНО ОДНУ ПРОВЕРКУ — больше ничего:\n"
            "  Совпадает ли название заявителя из Плана с названием организации, упомянутым "
            "  в файле орг.структуры? Допускаются сокращения, кавычки, регистр, разные варианты "
            "  кавычек («» vs \"\"), ОПФ перед именем или после (ПАО / ООО / АО).\n\n"
            "СТРОГО ЗАПРЕЩЕНО:\n"
            "  – Сравнивать конкретные названия подразделений между Планом и Орг.структурой. "
            "    Подразделения в Плане упоминаются в контексте проверяемых отделов; их состав НЕ обязан "
            "    совпадать с орг.структурой, потому что аудит может покрывать только часть.\n"
            "  – Требовать наличия в орг.структуре подразделений, упомянутых в Плане.\n"
            "  – Делать вывод о расхождении, если орг.структура «выглядит иначе» — она и должна.\n\n"
            "OK — если название заявителя из Плана найдено в тексте файла орг.структуры (по сути, "
            "не обязательно дословно).\n"
            "NOK — ТОЛЬКО если: файл орг.структуры не загружен, ИЛИ название заявителя в файле "
            "совершенно другое (другая организация).\n\n"
            "Reason: одна короткая фраза вида «Заявитель «ПАО Газпром автоматизация» подтверждён "
            "в файле орг.структуры — OK». Никаких рассуждений про подразделения. extracted_values НЕ заполняй."
        ),
    },
    2: {
        "file_keywords": ["трудоемкост", "трудоёмкост", "расчет труд", "расчёт труд"],
        "extra_instructions": (
            "Сравниваем РАСЧЁТ ТРУДОЁМКОСТИ с фактическим Планом аудита.\n"
            "Вердикт ставится детерминированно в коде после твоего ответа — твоя задача "
            "просто извлечь и привести числа в reason, чтобы оператор мог их проверить.\n\n"
            "ИЗ ФАЙЛА «Расчёт трудоёмкости» возьми строку, соответствующую виду аудита из Плана "
            "(Первый/Второй инспекционный контроль, Ресертификационный, Первый/Второй этап "
            "первичной сертификации и т.п.). В этой строке обычно два значения чел.-дней: "
            "«Вне территории Заказчика» и «На территории Заказчика». R_total = их сумма.\n\n"
            "ИЗ ПЛАНА:\n"
            "  P_days = число дней в графике (раздел «Сроки проведения», диапазон).\n"
            "  P_experts = число членов экспертной группы (раздел «Состав ЭГ»: руководитель + "
            "  эксперты + технические эксперты + кандидаты + наблюдатели — все непустые).\n"
            "  P_total = P_days × P_experts.\n\n"
            "В reason приведи все числа: R_total (с разбивкой на вне/на территории), P_days, "
            "P_experts, P_total. Сравни R_total и P_total. Если расходятся — NOK с конкретными "
            "числами; если совпадают — OK. extracted_values НЕ заполняй (вердикт всё равно "
            "детерминированный)."
        ),
    },
    3: {
        "file_keywords": ["план", "приказ", "эг", "назначен"],
        "extra_instructions": (
            "ЭТОТ ПУНКТ — про ВИД АУДИТА (галочки в блоке '1 Цель и область аудита' Плана), "
            "НЕ про сроки и НЕ про состав ЭГ. Не отвечай про другие пункты.\n\n"
            "В evidence Плана есть секция «ОТМЕТКИ (галочки/крестики)» или «СИМВОЛЫ-ОТМЕТКИ» / «ЧЕКБОКСЫ PDF» — "
            "в ней список символов: ☐ (пусто, НЕ отмечено), ☑ или ☒ (отмечено).\n"
            "АЛГОРИТМ:\n"
            "1. Посчитай количество отмеченных (☑/☒) и пустых (☐) чекбоксов в этой секции.\n"
            "2. Если отмеченных = 0 (все ☐) — это ЗАВЕДОМО NOK: вид аудита не выбран.\n"
            "3. Если хотя бы одна отметка стоит — определи, какая (например, «Второй инспекционный контроль»), "
            "и сверь с видом аудита из файла «Приказ о назначении ЭГ».\n"
            "4. NOK если отметка не соответствует тому, что в Приказе.\n"
            "В reason укажи: количество ☐, количество ☑/☒, какой вид отмечен в Плане, какой в Приказе."
        ),
    },
    4: {
        "file_keywords": [
            "оквэд", "сертификат", "игс", "разбивк", "макет",
            "область сертификаци", "область применени",
            "scope", "смк ", "смк_", "смк-",
        ],
        "comparison": True,
        "compare_fields": "формулировка области применения (сертификации) СМК — дословно",
        "extra_instructions": (
            "В шаблоне Плана найди строку НЕПОСРЕДСТВЕННО НАД фразой \"область применения СМК (область сертификации)\". "
            "Эта строка должна ДОСЛОВНО (буква в букву) совпадать с формулировкой из файлов «Разбивка кодов ОКВЭД» и «Сертификат ИГС».\n\n"
            "ОБЯЗАТЕЛЬНО заполни extracted_values (иначе пункт автоматически NOK). Для каждого источника отдельная строка:\n"
            "1) field='Формулировка области (План vs Разбивка ОКВЭД)', value_in_plan = ДОСЛОВНАЯ цитата из Плана, "
            "value_in_other = ДОСЛОВНАЯ цитата из «Разбивка кодов ОКВЭД» (или 'не найдено'), "
            "other_file = имя файла Разбивки, match = true только при ПОЛНОМ буквальном совпадении строк.\n"
            "2) field='Формулировка области (План vs Сертификат ИГС)', value_in_plan = та же дословная цитата из Плана, "
            "value_in_other = ДОСЛОВНАЯ цитата из «Сертификат ИГС» (или 'не найдено'), other_file = имя файла Сертификата, "
            "match = true только при ПОЛНОМ буквальном совпадении.\n\n"
            "ПРАВИЛА СОВПАДЕНИЯ:\n"
            "- Любое отличие в словах (есть «конструкторской» в Плане, нет в Разбивке; есть «шефмонтажу/монтажу/пусконаладке» в Плане, нет в Разбивке) → match=false.\n"
            "- Отличие в порядке слов или знаках препинания внутри значимой части → match=false.\n"
            "- Регистр и кавычки можно игнорировать.\n\n"
            "Если в источниках нет ни «Разбивки ОКВЭД», ни «Сертификата ИГС» — NOK с reason: "
            "\"Не загружены файлы для сверки: «Разбивка кодов ОКВЭД» и/или «Сертификат ИГС»\". "
            "«Макет сертификата» НЕ является заменой Сертификата ИГС — если есть только Макет, NOK с явным указанием."
        ),
    },
    5: {
        "file_keywords": [
            "договор", "дог.", "дог №", "дог. №", "дог_", "контракт",
            "доп.соглашен", "доп соглашен", "доп. соглашен", "соглашен",
        ],
        "comparison": True,
        "compare_fields": "№ договора, дата договора, № и дата доп.соглашения (если есть)",
        "extra_instructions": (
            "В шаблоне Плана пункт 2 — \"Основание (номер и дата заявки и/или договора)\". "
            "В этой строке найди номер и дату договора (обычно в форме 'Договор №<номер> от <дата>'). "
            "Они должны совпадать с реквизитами из файла \"Договор\" (а если есть доп.соглашение — то и с ним). "
            "Если в источниках несколько договоров — для сверки выбирай тот, чей номер/дата совпадают с указанными в пункте 2 Плана; "
            "остальные договоры в реквизиты этого пункта не вмешивай. "
            "NOK если номер/дата договора не совпали или не упомянуто доп.соглашение, когда оно реально присутствует в источниках.\n\n"
            "НОРМАЛИЗАЦИЯ ПРИ СРАВНЕНИИ НОМЕРОВ ДОГОВОРА И ДАТ (обязательно):\n"
            "- Игнорируй пробелы внутри номера: '17 60-25' ≡ '17-60-25' ≡ '17/60/25'.\n"
            "- Считай эквивалентными разделители: пробел, '-', '_', '/', '.'. То есть '05_09-2023' ≡ '05/09/2023' ≡ '05.09.2023' ≡ '05 09 2023'.\n"
            "- Игнорируй регистр и наличие/отсутствие символа '№' перед номером.\n"
            "- Скобки вокруг внутреннего шифра '(17-60-25)' можно игнорировать — это тот же номер.\n"
            "- Даты сравнивай по значению ДД.ММ.ГГГГ независимо от формата записи ('03.09.2025' ≡ '03.09.25' ≡ '03 сентября 2025 г.').\n"
            "Если различие ТОЛЬКО в форматировании (пробелы/разделители/скобки/регистр) — это НЕ расхождение, ставь match=true. "
            "NOK только когда отличаются сами цифры/значащие части номера или сама дата.\n\n"
            "В extracted_values заведи ОТДЕЛЬНЫЕ строки для '№ договора', 'дата договора', "
            "и для каждого доп.соглашения — '№ доп.соглашения N', 'дата доп.соглашения N'. "
            "В value_in_plan / value_in_other пиши значения как они выглядят в каждом источнике (без нормализации, "
            "чтобы было видно фактическое написание), а match рассчитывай по нормализованному сравнению."
        ),
    },
    6: {
        # Пункт 7 чек-листа: «п.4 Площадки». Источник по чек-листу — «Информация
        # для подготовки к аудиту». Если такого файла нет — это НЕ автоматический NOK,
        # т.к. сам Плана содержит перечень площадок, и базовую проверку можно сделать по нему.
        "file_keywords": ["план", "информация для подготовки", "инфо для подгот", "подготовк"],
        "extra_instructions": (
            "Проверяем раздел «4 Наименование и адрес аудитируемых производственных площадок» Плана.\n\n"
            "Проверь по Плану:\n"
            "1. Перечислены ли площадки с адресами (хотя бы одна).\n"
            "2. Если упомянут строительный объект — рядом должна быть пометка «Объект заявителя».\n"
            "3. Если есть файл «Информация для подготовки к аудиту» — площадки и адреса не должны расходиться.\n\n"
            "OK — площадки перечислены, противоречий с «Информацией для подготовки» нет, "
            "строительные объекты (если есть) помечены корректно.\n"
            "NOK — нет ни одной площадки, либо есть стройка без пометки, либо адреса расходятся с источником.\n\n"
            "REASON — короткая фраза для клиента БЕЗ внутренних терминов правила. "
            "Пиши по-человечески: «В Плане 2 площадки с адресами, расхождений нет», "
            "а не «следовательно указание Объект заявителя не требуется». "
            "Технические маркеры («ИИ\\d+», названия шагов алгоритма) в reason НЕ упоминать."
        ),
    },
    7: {
        "file_keywords": ["план", "приказ", "эг", "назначен", "трудоемк", "трудоёмк", "заявк"],
        "comparison": True,
        "compare_fields": "даты начала и окончания аудита, перечень дней",
        "extra_instructions": (
            "В шаблоне Плана найди пункт 5 \"Сроки проведения аудита\" — там диапазон дат "
            "(например, «09-11.12.2025г.»). Это эталон.\n\n"
            "Проверь, что те же даты упомянуты ХОТЯ БЫ В ОДНОМ из следующих файлов:\n"
            "  – «Приказ о назначении ЭГ» (если в нём указаны даты аудита);\n"
            "  – «Расчёт трудоёмкости»;\n"
            "  – «Заявка» / договор.\n\n"
            "ВАЖНО: Приказ о назначении ЭГ часто содержит ТОЛЬКО дату самого приказа "
            "(когда он подписан), а не даты проведения аудита. Если в Приказе дат аудита нет — "
            "это НЕ повод для NOK. Просто сверь с другим источником (Расчёт/Заявка) или, "
            "если ни один сторонний источник не содержит дат аудита, пиши reason: "
            "«Даты аудита указаны только в Плане; стороннего источника для перекрёстной сверки нет — "
            "отмечено как OK по самому Плану».\n\n"
            "OK — даты в Плане согласуются хотя бы с одним сторонним источником, ИЛИ ни один сторонний "
            "источник дат не содержит (нечего противоречить).\n"
            "NOK — даты в Плане ПРЯМО ПРОТИВОРЕЧАТ датам, указанным в стороннем источнике.\n\n"
            "В extracted_values заведи строки 'дата начала аудита' и 'дата окончания аудита'. "
            "Если в файле даты отсутствуют — value_in_other='не найдено в этом файле', match=true "
            "(отсутствие данных ≠ расхождение)."
        ),
    },
    8: {
        "file_keywords": ["план", "приказ", "эг", "назначен"],
        "comparison": True,
        "compare_fields": "ФИО каждого члена ЭГ, роль (руководитель / эксперт / стажёр / наблюдатель)",
        "extra_instructions": (
            "В шаблоне Плана найди раздел \"Состав экспертной группы\" (ФИО, роли). "
            "Сравни с составом из файла \"Приказ о назначении ЭГ\". NOK при расхождении состава или ролей. "
            "В extracted_values заведи отдельную строку для КАЖДОГО человека: field='ФИО — <роль>', "
            "value_in_plan = ФИО как написано в Плане, value_in_other = ФИО как написано в Приказе. "
            "Сравнивай нормализованно: 'Иванов И.И.' ≡ 'Иванов Иван Иванович', но 'Иванов И.И.' ≠ 'Иванов И.А.' и ≠ 'Иванова И.И.'. "
            "Если человек есть в одном файле и нет в другом — match: false."
        ),
    },
    9: {
        "file_keywords": ["акт", "предыдущ", "результат", "сводный акт", "1 этап", "1-го этапа", "2 этап", "2-го этапа", "отчет"],
        "extra_instructions": (
            "Сверяем раздел «8 Замечания/несоответствия по результатам предыдущего аудита» Плана "
            "с файлом «Акт предыдущего аудита» (или «Сводный акт», «Отчёт по 1/2 этапу»).\n\n"
            "Проверь:\n"
            "1. В Плане в п.8 перечислены замечания/несоответствия из предыдущего аудита (если их не было — должна быть пометка «отсутствуют»/«не выявлены»).\n"
            "2. Этот перечень совпадает по существу с данными из Акта предыдущего аудита.\n\n"
            "ВАЖНО: НЕ сравнивай состав экспертной группы — это другой пункт. Здесь только замечания/несоответствия.\n\n"
            "OK — данные в п.8 Плана соответствуют Акту (либо обе стороны подтверждают «замечаний не было»).\n"
            "NOK — данные расходятся, либо в Акте нет нужной информации, либо файл Акта не загружен.\n\n"
            "REASON — прямой вердикт одной фразой для клиента. "
            "НЕ пиши «требуется сравнить...» — ты УЖЕ выполняешь сравнение, в reason должен быть ИТОГ. "
            "Пример OK: «Замечаний в Акте не выявлено, в п.8 Плана соответствует». "
            "Пример NOK: «В Акте указано 3 замечания, в п.8 Плана упомянуто только 1»."
        ),
    },
    10: {
        # Проверка по содержимому Плана — отдельный файл с самим СТО загружать НЕ требуется
        # (это стандарт, не дозагружаемый документ). Подключаем План и при наличии — Сертификат.
        "file_keywords": ["план", "сертификат", "игс"],
        "extra_instructions": (
            "Этот пункт — про список ИСКЛЮЧЕНИЙ из СТО, указанный в Плане в разделе "
            "«8 Исключения из требований стандартов».\n\n"
            "Допустимые исключения по бизнес-правилу:\n"
            "- п. 7.1.3.5 — только если исключён 1 абзац (не весь пункт);\n"
            "- п. 7.1.4.3 — только если исключён 1 абзац (не весь пункт);\n"
            "- никаких других исключений быть не должно.\n\n"
            "АЛГОРИТМ:\n"
            "1. Найди в Плане раздел «8 Исключения из требований стандартов».\n"
            "2. Выпиши все перечисленные там пункты СТО (например, '7.1.3.5 (абзац 1)', '7.1.4.3 (абзац 1)', '8.3.5.4').\n"
            "3. Сравни со списком допустимых: только 7.1.3.5 (1 абзац) и 7.1.4.3 (1 абзац).\n"
            "4. OK — если в исключениях только 7.1.3.5 и/или 7.1.4.3 с пометкой про 1 абзац.\n"
            "5. NOK — если есть любой другой пункт (8.3.5.4 и т.п.), или 7.1.3.5/7.1.4.3 исключены целиком (без пометки про 1 абзац).\n"
            "В reason приведи полный список найденных исключений и какие из них нелегитимны.\n\n"
            "Отдельный файл со стандартом «СТО Газпром 9001-2018» загружать НЕ нужно — это сам стандарт, "
            "проверка идёт по содержимому Плана. Не пиши «не загружен файл СТО»."
        ),
    },
    11: {
        "file_keywords": ["план"],
        "extra_instructions": (
            "В шаблоне Плана проверь наличие совещаний. Минимальные требования:\n"
            "- «Предварительное совещание» — в НАЧАЛЕ первого дня на КАЖДОЙ уникальной площадке (новом адресе).\n"
            "- «Заключительное совещание» — в КОНЦЕ последнего дня на КАЖДОЙ уникальной площадке (для последней площадки это конец всего аудита).\n"
            "- В конце прочих дней (не последних) — должно быть «Рабочее совещание» И/ИЛИ «Промежуточное совещание». "
            "Достаточно хотя бы одного из этих двух; одновременно их наличие НЕ обязательно.\n"
            "- В дни с «Заключительным совещанием» отдельные «Промежуточное»/«Рабочее» НЕ требуются — Заключительное их заменяет.\n\n"
            "OK если для каждой площадки есть Предварительное и Заключительное, и для остальных дней — хотя бы одно из Рабочее/Промежуточное. "
            "NOK только если отсутствует Предварительное на площадке, или Заключительное на последнем дне площадки, или если в обычный день не запланировано НИ одного совещания.\n"
            "В reason укажи: перечень площадок и какие совещания на каких днях найдены."
        ),
    },
    12: {
        "file_keywords": ["план"],
        "extra_instructions": (
            "Инструктаж — это ПРОСТО СТРОКА в таблице Плана в столбце «Вид деятельности, процесс» "
            "со словом «Инструктаж» (обычно «Инструктаж по технике безопасности и охране труда»). "
            "ОТДЕЛЬНЫХ документов/записей об инструктаже искать НЕ НУЖНО — проверка только по тексту Плана.\n\n"
            "АЛГОРИТМ:\n"
            "1. Найди в тексте Плана ВСЕ строки со словом «Инструктаж».\n"
            "2. Каждая такая строка — это запланированный инструктаж на одной площадке для всех экспертов сразу.\n"
            "3. Площадки в Плане различаются по адресам (заголовки дней содержат адрес).\n"
            "4. OK если строка «Инструктаж» есть как минимум по одной для КАЖДОГО уникального адреса площадки в Плане.\n"
            "5. NOK ТОЛЬКО если на каком-то адресе площадки строки «Инструктаж» нет вообще, либо если в Плане нет ни одного инструктажа.\n"
            "В reason укажи: число найденных строк «Инструктаж» и список адресов, на которых они стоят."
        ),
    },
    13: {
        "file_keywords": ["акт", "предыдущ", "сводный акт", "1 этап", "1-го этапа", "2 этап", "2-го этапа", "отчет"],
        "extra_instructions": (
            "Сравни процессы из файла «акт предыдущего аудита» с процессами в шаблоне Плана.\n"
            "ВАЖНО: процессы в Плане ищи в колонке «Вид деятельности, процесс» (это столбец 2 таблицы Плана) — "
            "строки начинаются со слов «Процесс П1», «Процесс П2» и т.п. В колонке «Пункт стандарта» (столбец 3) — "
            "номера пунктов СТО (4.4.1, 4.4.2 и т.п.), процессов там НЕТ.\n"
            "Алгоритм:\n"
            "1. Выпиши все процессы из акта предыдущего аудита (по их названиям/номерам).\n"
            "2. Найди в Плане в столбце «Вид деятельности, процесс» все строки со словом «Процесс».\n"
            "3. Сверь: каждый процесс из акта должен встречаться в Плане (с тем же названием или явным эквивалентом).\n"
            "OK — если все процессы из акта найдены в Плане. NOK — если хотя бы один процесс из акта отсутствует в Плане.\n\n"
            "REASON — прямой вердикт для клиента. ЗАПРЕЩЕНО упоминать «первый аудитор», «второй аудитор», "
            "«перепроверка», «адверсариальный», или ссылаться на сам процесс проверки. "
            "Пиши только результат: «OK: в Плане найдены процессы П1, П2, П3 — совпадают с актом» "
            "или «NOK: в акте есть процесс П4, в Плане отсутствует»."
        ),
    },
    14: {
        "file_keywords": ["план"],
        "extra_instructions": (
            "В шаблоне Плана в колонке \"Пункт стандарта\": в каждой строке, где встречается слово \"Процесс\", "
            "должны быть указаны пункты 4.4.1, 4.4.2 и 4.4.3 (все три). "
            "NOK если в строке с \"Процесс\" отсутствует хотя бы один из 4.4.1 / 4.4.2 / 4.4.3."
        ),
    },
    15: {
        "file_keywords": ["план"],
        "extra_instructions": (
            "В шаблоне Плана найди временные интервалы перерыва (формат вида \"12:00 - 12:30\"). "
            "Каждый перерыв должен длиться 30 минут или больше (то есть >= 30 минут). "
            "Перерыв ровно 30 минут — это OK, требование выполнено. "
            "NOK ставь ТОЛЬКО если есть хотя бы один перерыв СТРОГО короче 30 минут (29 минут и меньше). "
            "В reason укажи найденные перерывы и их длительность в минутах."
        ),
    },
}


# ===== Fuzzy-сверка подразделений (для п.2 «Орг.структура») =====
# OCR орг.структуры часто содержит шум (склейки слов, замены букв), поэтому
# буквальное сравнение названий не работает. Используем сверку по корням слов
# с допуском в 1 редакторскую ошибку на каждые 4 символа.

_ORG_STOP_WORDS = {
    "и", "по", "в", "для", "с", "из", "на", "от", "к", "до", "за", "над",
    "под", "при", "о", "об", "у", "не", "также", "или", "но", "а", "их",
    "то", "что", "как",
}
_ORG_GENERIC_PREFIXES = (
    "отдел", "отделом", "управлен", "департамент", "служб",
    "групп", "сектор", "центр", "бюро",
)


def _org_normalize(text: str) -> str:
    s = (text or "").lower().replace("ё", "е")
    s = re.sub(r"[^а-я0-9]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _org_significant_stems(division_name: str) -> list:
    """Возвращает «значащие» корни (по 6 символов) из названия подразделения,
    отбрасывая стоп-слова и родовые («отдел», «управление» и т.п.)."""
    norm = _org_normalize(division_name)
    stems = []
    seen = set()
    for tok in norm.split():
        if len(tok) < 5:
            continue
        if tok in _ORG_STOP_WORDS:
            continue
        if any(tok.startswith(g) for g in _ORG_GENERIC_PREFIXES):
            continue
        stem = tok[:6]
        if stem not in seen:
            seen.add(stem)
            stems.append(stem)
    return stems


def _edit_distance(a: str, b: str, cutoff: int = 3) -> int:
    """Усечённое расстояние Левенштейна. Возвращает cutoff+1, если превысили."""
    if a == b:
        return 0
    if abs(len(a) - len(b)) > cutoff:
        return cutoff + 1
    la, lb = len(a), len(b)
    if la > lb:
        a, b = b, a
        la, lb = lb, la
    prev = list(range(la + 1))
    for i, cb in enumerate(b, 1):
        curr = [i] + [0] * la
        best = curr[0]
        for j, ca in enumerate(a, 1):
            curr[j] = min(
                curr[j - 1] + 1,
                prev[j] + 1,
                prev[j - 1] + (0 if ca == cb else 1),
            )
            if curr[j] < best:
                best = curr[j]
        if best > cutoff:
            return cutoff + 1
        prev = curr
    return prev[-1]


def _stem_in_text(stem: str, ocr_text_norm: str, ocr_tokens: list) -> bool:
    """Проверяет, встречается ли корень в OCR-тексте — подстрока или fuzzy."""
    if stem in ocr_text_norm:
        return True
    max_dist = max(1, len(stem) // 4)
    for tok in ocr_tokens:
        if len(tok) < len(stem):
            continue
        for i in range(len(tok) - len(stem) + 1):
            if _edit_distance(stem, tok[i : i + len(stem)], cutoff=max_dist) <= max_dist:
                return True
    return False


_PLAN_DIVISION_RE = re.compile(
    r"([А-ЯЁ][А-Яа-яЁё0-9 ,\-«»\"'()/]{8,200}?)"
    r"\s+[–—-]\s+"
    r"[А-ЯЁ][а-яё]+(?:ов|ев|ин|кий|ова|ева|ина|ская|ко|ян|ук|юк|ий)\s*[А-ЯЁ]\.\s*[А-ЯЁ]\."
)


def _extract_plan_divisions(plan_text: str) -> list:
    """Вытаскивает подразделения из раздела 10 «Объекты аудита»:
    строки вида '<Подразделение> – <Фамилия И.О.>'."""
    if not plan_text:
        return []
    divisions = []
    seen = set()
    for m in _PLAN_DIVISION_RE.finditer(plan_text):
        d = m.group(1).strip(" ,.|-")
        if not re.search(r"(отдел|управлен|департамент|служб|групп|сектор|центр|бюро|цех)", d, re.IGNORECASE):
            continue
        key = _org_normalize(d)
        if key in seen or len(key) < 10:
            continue
        seen.add(key)
        divisions.append(d)
    return divisions


_APPLICANT_OPF_RE = re.compile(r"\b(пао|ао|ооо|зао|оао|нко|фгуп|пкф|ип)\b", re.IGNORECASE)


def _applicant_core_name(name: str) -> str:
    """Из «ПАО \"Газпром автоматизация\"» оставляет «газпром автоматизация» для нечёткой сверки."""
    if not name:
        return ""
    s = name.lower().replace("ё", "е")
    s = _APPLICANT_OPF_RE.sub(" ", s)
    s = re.sub(r"[«»\"'()]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def fuzzy_orgstructure_check(plan_text: str, org_text: str, applicant_name: str = "") -> dict:
    """Детерминированная сверка орг.структуры с Планом.
    Возвращает {ok, reason, matched, total, missing}."""
    if not org_text:
        return {
            "ok": False,
            "reason": "Файл организационной структуры не загружен или его текст не извлёкся.",
            "matched": 0, "total": 0, "missing": [],
        }

    org_norm = _org_normalize(org_text)
    org_tokens = [t for t in org_norm.split() if len(t) >= 4]

    # 1) Проверка названия заявителя (если оно известно).
    if applicant_name:
        core = _applicant_core_name(applicant_name)
        core_tokens = [t for t in core.split() if len(t) >= 4]
        applicant_hits = sum(1 for t in core_tokens if _stem_in_text(t[:6], org_norm, org_tokens))
        if core_tokens and applicant_hits < max(1, len(core_tokens) // 2 + (len(core_tokens) % 2)):
            return {
                "ok": False,
                "reason": (
                    f"Название заявителя «{applicant_name}» не найдено в файле орг.структуры — "
                    f"возможно, загружена орг.структура другой организации."
                ),
                "matched": 0, "total": 0, "missing": [],
            }

    # 2) Сверка подразделений.
    divisions = _extract_plan_divisions(plan_text)
    if not divisions:
        return {
            "ok": True,
            "reason": (
                "В Плане не удалось выделить перечень подразделений (раздел 10 «Объекты аудита» "
                "пуст или имеет нестандартный формат). Сверка с орг.структурой пропущена; "
                "название заявителя в орг.структуре подтверждено."
            ),
            "matched": 0, "total": 0, "missing": [],
        }

    matched, missing = [], []
    for div in divisions:
        stems = _org_significant_stems(div)
        if not stems:
            continue
        need = max(1, (len(stems) + 1) // 2)
        hits = sum(1 for s in stems if _stem_in_text(s, org_norm, org_tokens))
        if hits >= need:
            matched.append(div)
        else:
            missing.append(div)

    total = len(matched) + len(missing)
    if total == 0:
        return {
            "ok": True,
            "reason": "Подразделения из Плана содержат только родовые слова — сверка пропущена.",
            "matched": 0, "total": 0, "missing": [],
        }

    ratio = len(matched) / total
    ok = ratio >= 0.6
    if ok:
        reason = (
            f"В орг.структуре найдено {len(matched)} из {total} подразделений из Плана "
            f"({int(ratio * 100)}%)."
        )
        if missing:
            shown = missing[:3]
            reason += f" Не нашлись: {'; '.join(shown)}"
            if len(missing) > 3:
                reason += f" (и ещё {len(missing) - 3})"
            reason += " — возможно, незначимые отклонения OCR."
    else:
        shown = missing[:5]
        reason = (
            f"В орг.структуре найдено только {len(matched)} из {total} подразделений из Плана "
            f"({int(ratio * 100)}%, порог 60%). Не нашлись: {'; '.join(shown)}"
        )
        if len(missing) > 5:
            reason += f" (и ещё {len(missing) - 5})"
        reason += "."
    return {"ok": ok, "reason": reason, "matched": len(matched), "total": total, "missing": missing}


# ============================================================================
# Правила для чек-листа «Сводный акт» (14 пунктов). Индексы 0..13 = пункты №1..14.
# Главный проверяемый документ — «Сводный акт исследования (итог)» — подставляется
# в evidence каждого пункта движком process_checklist_svod (как SVOD_DOC).
# file_keywords — по каким подстрокам имени искать файл-источник для сверки.
# manual=True — пункт по методике проверяется вручную: ставим «на ручную проверку».
# ============================================================================
SVOD_ITEM_RULES = {
    0: {  # №1 Соответствие формы утверждённому шаблону
        "extra_instructions": (
            "Проверь нижний колонтитул (footer) Сводного акта: в левом углу должна быть "
            "надпись о версии формы, например «вер. 2 от 14.11.25г.». "
            "OK — если такая отметка о версии присутствует. NOK — если её нет."
        ),
    },
    1: {  # №2 Заполнение всех разделов согласно шаблону
        "file_keywords": ["шаблон"],
        "extra_instructions": (
            "Сверь Сводный акт с файлом-ШАБЛОНОМ «Шаблон -Сводный акт». В шаблоне разделы "
            "содержат курсивные инструкции «Указать ...». Проверь, что в Сводном акте все "
            "разделы фактически заполнены данными (а не остались инструкции/пустые ячейки). "
            "NOK — если есть незаполненные обязательные разделы; перечисли их."
        ),
    },
    2: {  # №3 Область распространения СМК
        "file_keywords": ["акт р", "акт 2", "акт по результ", "результатам аудита"],
        "comparison": True,
        "compare_fields": "формулировка области распространения (применения) СМК",
        "extra_instructions": (
            "Сверь формулировку области распространения СМК в Сводном акте с пунктом 1 "
            "файла «3.1 Акт Р» (раздел «1 Цель и область аудита» → «область применения СМК "
            "(область сертификации)»). Должны совпадать по существу. NOK — при смысловом "
            "расхождении формулировок; приведи обе цитаты."
        ),
    },
    3: {  # №4 Коды ОКВЭД
        "file_keywords": ["разбивк", "оквэд", "акт р"],
        "comparison": True,
        "compare_fields": "перечень кодов ОКВЭД",
        "extra_instructions": (
            "Сравни перечень кодов ОКВЭД в Сводном акте (колонка ОКВЭД в таблицах продукции "
            "и услуг) с файлом «4.1 Разбивка области по кодам ОКВЭД» и п.5 файла «3.1 Акт Р» "
            "(«Подтвержденные коды видов экономической деятельности»). Перечисли коды из "
            "каждого источника. NOK — если множества кодов различаются; укажи какие коды "
            "лишние/отсутствуют."
        ),
    },
    4: {  # №5 Коды ОКВЭД — наличие услуг
        "extra_instructions": (
            "Внутренняя сверка Сводного акта: сравни раздел продукции (п.1) и раздел услуг/"
            "работ (п.2). Проверь, что коды ОКВЭД услуг (п.2) соответствуют заявленной "
            "области и не противоречат продукции (п.1). NOK — при явном расхождении."
        ),
    },
    5: {  # №6 Количество площадок (GigaChat — по решению пользователя, Заявка из OCR)
        "file_keywords": ["заявк"],
        "extra_instructions": (
            "Сверь количество производственных площадок: в Сводном акте раздел «Общие "
            "сведения об организации» (адреса/площадки) и файл «6.1 Заявка». ВНИМАНИЕ: текст "
            "Заявки получен через OCR и может содержать искажения — ориентируйся на смысл "
            "(адрес: Московская область, Подольск, Слащево). OK — если число площадок и "
            "основной адрес совпадают; NOK — при явном расхождении количества площадок."
        ),
    },
    6: {  # №7 Адреса мест осуществления деятельности
        "file_keywords": ["акт р", "акт 2", "результатам аудита", "заявк"],
        "extra_instructions": (
            "Проверь адреса мест осуществления деятельности в Сводном акте (раздел «Общие "
            "сведения об организации»): они должны быть заполнены и согласовываться с адресом "
            "производственных площадок из «3.1 Акт Р» (п.4 «Наименование и адрес ... площадок») "
            "и/или Заявки. Приведи адрес из Сводного акта и из источника. OK — если адреса "
            "заполнены и совпадают; NOK — если пусто или адреса расходятся."
        ),
    },
    7: {  # №8 Численность сотрудников
        "file_keywords": ["отчет", "отчёт", "1 этап", "первого этап", "ра "],
        "comparison": True,
        "compare_fields": "численность персонала",
        "extra_instructions": (
            "Сверь численность персонала: в Сводном акте раздел «Общие сведения об "
            "организации» (строки про численность рабочих/специалистов/персонала в СМК) и в "
            "файле «4.2 Отчет 1 этап» (приложение, пункт 4). Приведи числа из обоих "
            "источников. NOK — если значения расходятся."
        ),
    },
    8: {  # №9 Инфраструктура и оборудование
        "manual": True,
        "extra_instructions": (
            "Этот пункт по методике проверяется ВРУЧНУЮ. Кратко выведи содержание раздела "
            "«Инфраструктура»/«Оборудование» Сводного акта для ручной сверки. Вердикт "
            "оставь на оператора."
        ),
    },
    9: {  # №10 Режим работы (количество смен / часы работы)
        "file_keywords": ["трудоемкост", "трудоёмкост", "расчет труд", "расчёт труд"],
        "comparison": True,
        "compare_fields": "режим работы (количество смен)",
        "extra_instructions": (
            "Сверь режим работы: в Сводном акте раздел «Общие сведения об организации» строка "
            "«Режим работы (количество смен / часы работы)» и в файле «10.1 Трудоёмкость» "
            "строка «Количество смен». Приведи число смен из обоих источников. NOK — если "
            "количество смен различается."
        ),
    },
    10: {  # №11 Количество рекламаций
        "file_keywords": ["акт р", "акт 2", "акт по результ", "результатам аудита"],
        "comparison": True,
        "compare_fields": "количество рекламаций",
        "extra_instructions": (
            "Сверь количество рекламаций: в Сводном акте раздел про рекламации (П.9) и п.15 "
            "файла «3.1 Акт Р» (Акт 2 этапа). Приведи числа из обоих источников. NOK — если "
            "значения расходятся."
        ),
    },
    11: {  # №12 Несоответствия
        "manual": True,
        "extra_instructions": (
            "Этот пункт по методике проверяется ВРУЧНУЮ (несоответствия отображаются курсивом "
            "в форме). Выведи найденные пометки о несоответствиях для ручной сверки. Вердикт "
            "оставь на оператора."
        ),
    },
    12: {  # №13 Область особого внимания
        "extra_instructions": (
            "Проверь, что в Сводном акте раздел «Область особого внимания» ЗАПОЛНЕН (не пуст, "
            "не остался инструкцией шаблона). OK — если заполнен содержательно; NOK — если "
            "пуст или содержит только шаблонную инструкцию."
        ),
    },
    13: {  # №14 Дата подписания, наличие подписей РЭГ и руководителя ОС
        "file_keywords": ["акт р", "акт 2", "акт по результ", "результатам аудита"],
        "extra_instructions": (
            "Проверь две вещи в Сводном акте:\n"
            "1) Наличие подписей и дат: Руководитель экспертной группы (РЭГ) и Руководитель "
            "органа по сертификации — ФИО и дата подписания должны быть проставлены.\n"
            "2) Дата подписания Сводного акта должна быть НЕ ПОЗДНЕЕ 20 РАБОЧИХ дней от даты "
            "Акта 2 этапа (дата «утверждаю» в шапке файла «3.1 Акт Р»). Выполни расчёт в "
            "рабочих днях (Пн–Пт) и приведи обе даты и число рабочих дней между ними.\n"
            "NOK — если подписи/даты отсутствуют ИЛИ просрочка более 20 рабочих дней."
        ),
    },
}


def _files_by_keyword(file_names: list, keywords: list) -> list:
    """Возвращает файлы, в имени которых встречается ЛЮБАЯ из подстрок (без учёта регистра)."""
    if not keywords:
        return []
    kws = [k.lower() for k in keywords]
    matched = []
    for f in file_names:
        lf = f.lower()
        if any(kw in lf for kw in kws):
            matched.append(f)
    return matched


def process_checklist_advanced(api_key: str, all_texts: dict,
                                checklist_structure: list,
                                ii_references: dict,
                                model: str = "GigaChat",
                                template_text: str = "") -> dict:
    """
    Per-item обработка чек-листа:
      1. Извлечение шапки (1 вызов)
      2. Для каждого пункта:
         a. Классификация релевантных файлов (1 вызов)
         b. Строгая проверка NOK-first (1 вызов)
         c. Adversarial-перепроверка если OK (1 вызов)
    Всего: 1 + N*(2..3) вызовов GigaChat.
    """
    global processing_status

    total = len(checklist_structure)
    file_names = list(all_texts.keys())

    processing_status.update({
        "stage": "header",
        "current": 0,
        "total": total,
        "message": "Извлечение шапки документа...",
        "detail": "",
    })

    try:
        header_data = extract_header_info(api_key, all_texts, model=model)
    except Exception as e:
        print(f"[header] Ошибка: {e}")
        header_data = {
            "Наименование Заявителя": "не найдено",
            "Вид аудита": "не найдено",
            "Даты проведения": "не найдено",
            "РЭГ": "не найдено",
        }

    checklist_results = []

    for idx, item in enumerate(checklist_structure):
        processing_status.update({
            "stage": "verify",
            "current": idx + 1,
            "message": f"Пункт {idx+1}/{total}: {item['area'][:60]}",
            "detail": "классификация файлов...",
        })

        try:
            rule = ITEM_RULES.get(idx)
            file_names = list(all_texts.keys())

            keyword_hit = False
            if rule and rule.get("file_keywords"):
                # Детерминированный отбор по ключевым подстрокам в имени файла
                relevant = _files_by_keyword(file_names, rule["file_keywords"])
                if relevant:
                    keyword_hit = True
                else:
                    # Фоллбэк: общий поиск по разделу/ключевым словам
                    relevant = find_relevant_files_for_item(item, all_texts)
            else:
                relevant = find_relevant_files_for_item(item, all_texts)

            print(f"[item {idx+1}/{total}] релевантных файлов: {len(relevant)} (keyword_hit={keyword_hit})")

            # Быстрый NOK без вызова модели: пункт в режиме сверки, а целевого файла-источника нет.
            # Раньше в таких случаях модель уходила в режим comparison и часто возвращала битый JSON.
            if rule and rule.get("comparison") and rule.get("file_keywords") and not keyword_hit:
                kw_display = ", ".join(rule["file_keywords"][:5])
                checklist_results.append({
                    "ok": False,
                    "nok": True,
                    "reason": (
                        f"[авто-NOK: не загружен файл-источник для сверки. "
                        f"Ожидалось имя с подстрокой из: {kw_display}. "
                        f"Проверка содержательно невозможна.]"
                    ),
                    "ii_data_found": ", ".join(item.get("ii_markers", [])) or "не найдено",
                    "evidence_quote": "",
                    "source_file": "не найдено",
                    "_files_checked": [],
                })
                print(f"[item {idx+1}/{total}] NOK (нет файла-источника по ключевым словам)")
                continue

            evidence = build_evidence_pack(relevant, all_texts, max_chars=180000)

            extra_rules = ""
            comparison_flag = False
            compare_fields = ""
            expected_file_keywords = None
            if rule:
                extra_rules = rule.get("extra_instructions", "")
                comparison_flag = bool(rule.get("comparison", False))
                compare_fields = rule.get("compare_fields", "")
                expected_file_keywords = rule.get("file_keywords") or None
                # Для пунктов с правилами всегда добавляем текст самого Плана
                if template_text:
                    template_block = f"=== ФАЙЛ: ПЛАН АУДИТА (проверяемый документ) ===\n{template_text[:30000]}"
                    evidence = template_block + ("\n\n" + evidence if evidence else "")

            processing_status["detail"] = f"проверка по {len(relevant)} файлу(ам)..."
            verdict = verify_item_strict(api_key, item, ii_references, evidence, model=model,
                                          extra_instructions=extra_rules,
                                          comparison=comparison_flag,
                                          compare_fields=compare_fields,
                                          expected_file_keywords=expected_file_keywords)

            if verdict.get("ok"):
                processing_status["detail"] = "adversarial-перепроверка OK-вердикта..."
                verdict = adversarial_recheck(api_key, item, ii_references, evidence, verdict, model=model,
                                               comparison=comparison_flag,
                                               compare_fields=compare_fields)

            # Детерминированные пост-проверки идут по чистому тексту Плана
            # (а не по evidence — там template_block обрезан до 30k и ОТМЕТКИ могут
            # выпасть).
            plan_raw = ""
            for fn, txt in all_texts.items():
                if "план" in fn.lower() and txt:
                    plan_raw = txt
                    break

            # Пункт 1 (idx=0): дата утверждения Плана vs дата начала аудита.
            # Считаем сами — модель регулярно либо галлюцинирует даты, либо
            # пишет «не позднее пяти рабочих дней» словами, и старый
            # пост-чек на цифры в reason такие случаи пропускает.
            if idx == 0 and plan_raw:
                _MONTHS_RU = {
                    "январ": 1, "феврал": 2, "март": 3, "апрел": 4, "ма": 5, "июн": 6,
                    "июл": 7, "август": 8, "сентябр": 9, "октябр": 10, "ноябр": 11, "декабр": 12,
                }

                def _month_num(word: str):
                    w = (word or "").lower().replace("ё", "е")
                    for stem, n in _MONTHS_RU.items():
                        if w.startswith(stem):
                            return n
                    return None

                d_utv = None
                # «01» декабря 2025 г.  /  « 01 » декабря 20 25 г.
                m = re.search(
                    r"«\s*(\d{1,2})\s*»\s*([А-Яа-яё]+)\s*20\s*(\d{2})\s*г",
                    plan_raw,
                )
                if m:
                    mon = _month_num(m.group(2))
                    if mon:
                        d_utv = (int(m.group(1)), mon, 2000 + int(m.group(3)))

                d_start = None
                # «5 Сроки проведения аудита 09-11.12.2025г.» — диапазон.
                # Сначала пытаемся распознать диапазон DD-DD.MM.YYYY (берём первую дату),
                # затем — одиночную дату DD.MM.YYYY.
                i_srok = plan_raw.find("Сроки проведения")
                if i_srok >= 0:
                    seg = plan_raw[i_srok : i_srok + 400]
                    md = re.search(r"(\d{1,2})\s*[-–]\s*\d{1,2}\.(\d{1,2})\.(\d{2,4})", seg)
                    if not md:
                        md = re.search(r"(\d{1,2})\.(\d{1,2})\.(\d{2,4})", seg)
                    if md:
                        d, mo, y = int(md.group(1)), int(md.group(2)), int(md.group(3))
                        if y < 100:
                            y += 2000
                        d_start = (d, mo, y)

                if d_utv and d_start:
                    from datetime import date as _date, timedelta
                    try:
                        du = _date(d_utv[2], d_utv[1], d_utv[0])
                        ds = _date(d_start[2], d_start[1], d_start[0])
                    except ValueError:
                        du = ds = None
                    if du and ds:
                        # Считаем рабочие дни от du до ds включительно (Пн–Пт).
                        if du > ds:
                            verdict["ok"] = False
                            verdict["nok"] = True
                            verdict["reason"] = (
                                f"План утверждён ПОСЛЕ начала аудита: "
                                f"D_утв={du.strftime('%d.%m.%Y')}, "
                                f"D_начало={ds.strftime('%d.%m.%Y')}. Недопустимо."
                            )
                        else:
                            cnt = 0
                            cur = du
                            while cur <= ds:
                                if cur.weekday() < 5:
                                    cnt += 1
                                cur += timedelta(days=1)
                            if cnt > 5:
                                verdict["ok"] = False
                                verdict["nok"] = True
                                verdict["reason"] = (
                                    f"План утверждён слишком рано: "
                                    f"D_утв={du.strftime('%d.%m.%Y')}, "
                                    f"D_начало={ds.strftime('%d.%m.%Y')}, "
                                    f"между ними {cnt} рабочих дней (требование ≤5)."
                                )
                            else:
                                verdict["ok"] = True
                                verdict["nok"] = False
                                verdict["reason"] = (
                                    f"Дата утверждения Плана соответствует требованию: "
                                    f"D_утв={du.strftime('%d.%m.%Y')}, "
                                    f"D_начало={ds.strftime('%d.%m.%Y')}, "
                                    f"между ними {cnt} рабочих дней (≤5)."
                                )
                            print(
                                f"[item 1 deterministic] D_утв={du.isoformat()} "
                                f"D_начало={ds.isoformat()} рабочих_дней={cnt} -> "
                                f"{'OK' if verdict['ok'] else 'NOK'}"
                            )

            # Пункт 3 (idx=2, трудоёмкость): детерминированная сверка чисел.
            #   R_total = сумма «Вне территории» + «На территории» из строки Расчёта,
            #            соответствующей виду аудита из Плана.
            #   P_total = (число дней в графике Плана) × (число членов ЭГ в Плане).
            #   Расходятся → NOK с конкретными числами; совпадают → OK.
            if idx == 2 and plan_raw:
                # P_days: парсим диапазон в «Сроки проведения».
                p_days = None
                i_sr = plan_raw.find("Сроки проведения")
                if i_sr >= 0:
                    seg = plan_raw[i_sr : i_sr + 300]
                    mr = re.search(r"(\d{1,2})\s*[-–]\s*(\d{1,2})\.(\d{1,2})\.(\d{2,4})", seg)
                    if mr:
                        p_days = int(mr.group(2)) - int(mr.group(1)) + 1
                    else:
                        if re.search(r"\d{1,2}\.\d{1,2}\.\d{2,4}", seg):
                            p_days = 1

                # P_experts: считаем непустые роли 7.1..7.5 (Фамилия И.О.).
                p_experts = None
                i_eg = plan_raw.find("Состав экспертной группы")
                if i_eg >= 0:
                    seg = plan_raw[i_eg : i_eg + 2000]
                    fios = re.findall(
                        r"7\.\d+\s+[^|\n]+?\|\s*([А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.)",
                        seg,
                    )
                    p_experts = len(fios)

                # Вид аудита: используем общую функцию (есть в шапке).
                audit_type = (header_data or {}).get("Вид аудита", "") or ""
                audit_type_lc = audit_type.lower().replace("ё", "е")

                # Сопоставление видов с подстроками в Расчёте.
                TYPE_TO_KEYS = {
                    "первый инспекционный": "первый инспекционный",
                    "второй инспекционный": "второй инспекционный",
                    "внеплановый инспекционный": "внеплановый инспекционный",
                    "ресертификационный": "ресертифика",
                    "ресертификация": "ресертифика",
                    "первый этап": "первый этап первичного",
                    "второй этап": "второй этап первичного",
                    "дополнительный": "дополнительный",
                }
                rasch_key = None
                for typ, key in TYPE_TO_KEYS.items():
                    if typ in audit_type_lc:
                        rasch_key = key
                        break

                # Если по «Виду аудита» из шапки не получилось — попробуем из меток в Плане.
                if rasch_key is None:
                    for m in re.finditer(r"[☑☒]\s*([^☐☑☒\n]{4,120})", plan_raw):
                        label = m.group(1).lower().replace("ё", "е")
                        for typ, key in TYPE_TO_KEYS.items():
                            if typ in label and "расширение" not in label:
                                rasch_key = key
                                break
                        if rasch_key:
                            break

                rasch_files = [
                    f for f in relevant
                    if any(k in f.lower() for k in ("трудоемкост", "трудоёмкост", "расчет труд", "расчёт труд"))
                ]
                rasch_text = "\n".join(all_texts.get(f, "") for f in rasch_files)
                rasch_lc = rasch_text.lower().replace("ё", "е")

                r_total = None
                r_parts = []
                if rasch_key and rasch_text:
                    for m in re.finditer(re.escape(rasch_key), rasch_lc):
                        seg = rasch_text[m.start() : m.start() + 2000]
                        found_any = False
                        local_total = 0.0
                        local_parts = []
                        for territ in ("Вне территории", "На территории"):
                            mt = re.search(
                                territ + r"\s*Заказчика\s*:?\s*\|\s*([\d,\.]+|[-–—])",
                                seg,
                            )
                            if mt:
                                val = mt.group(1).replace(",", ".")
                                if val in ("-", "–", "—"):
                                    continue
                                try:
                                    n = float(val)
                                    local_total += n
                                    local_parts.append(f"{territ}: {mt.group(1)}")
                                    found_any = True
                                except ValueError:
                                    pass
                        if found_any:
                            r_total = local_total
                            r_parts = local_parts
                            break

                print(
                    f"[item 3 numbers] audit_type='{audit_type}' rasch_key='{rasch_key}' "
                    f"P_days={p_days} P_experts={p_experts} R_total={r_total}"
                )

                if p_days and p_experts and r_total is not None:
                    p_total = p_days * p_experts
                    # Сравнение с допуском 0.01 (на случай 9.5 vs 9.50).
                    if abs(p_total - r_total) < 0.01:
                        verdict["ok"] = True
                        verdict["nok"] = False
                        verdict["reason"] = (
                            f"Трудоёмкость согласована: Расчёт R_total={r_total:g} чел.-дн. "
                            f"({'; '.join(r_parts)}) = План P_days × P_experts = "
                            f"{p_days} × {p_experts} = {p_total:g}."
                        )
                    else:
                        verdict["ok"] = False
                        verdict["nok"] = True
                        verdict["reason"] = (
                            f"Трудоёмкость НЕ совпадает: в Расчёте R_total={r_total:g} чел.-дн. "
                            f"({'; '.join(r_parts)}); по Плану P_days × P_experts = "
                            f"{p_days} × {p_experts} = {p_total:g}. Разница {abs(p_total - r_total):g}."
                        )
                elif not rasch_files:
                    verdict["ok"] = False
                    verdict["nok"] = True
                    verdict["reason"] = (
                        "Файл «Расчёт трудоёмкости» не загружен — сверить трудоёмкость не с чем."
                    )
                # Если что-то из чисел не извлеклось — оставляем вердикт модели как есть.

            # Пункт 2 (idx=1, орг.структура): детерминированная fuzzy-сверка.
            # Перебивает вердикт модели, потому что OCR орг.структуры всегда
            # шумный и LLM регулярно ошибается в обе стороны.
            if idx == 1:
                org_files = [
                    f for f in relevant
                    if any(k in f.lower() for k in (
                        "оргструктур", "орг.структур", "орг структур",
                        "организацион", "ос ", "ос_", "ос-", "ос.",
                    ))
                ]
                if not org_files:
                    verdict["ok"] = False
                    verdict["nok"] = True
                    verdict["reason"] = (
                        "Файл организационной структуры не загружен. "
                        "Загрузите файл с подстрокой в имени: «Орг.структура», «Организационная», «ОС…»."
                    )
                else:
                    org_text = all_texts.get(org_files[0], "")
                    applicant = (header_data or {}).get("Наименование Заявителя", "") or ""
                    fz = fuzzy_orgstructure_check(plan_raw, org_text, applicant_name=applicant)
                    verdict["ok"] = bool(fz["ok"])
                    verdict["nok"] = not verdict["ok"]
                    verdict["reason"] = fz["reason"]
                    verdict["source_file"] = org_files[0]
                    print(
                        f"[item 2 fuzzy] matched={fz['matched']}/{fz['total']} "
                        f"missing={len(fz['missing'])} -> {'OK' if fz['ok'] else 'NOK'}"
                    )

            # Пункт 4 (idx=3): вид аудита — в Плане обязана стоять галочка (☑/☒).
            # Считаем напрямую в plan_raw. Если есть секция «ОТМЕТКИ» — берём её,
            # иначе сканируем весь текст Плана (на случай если docx-парсер не выделил секцию).
            if idx == 3 and plan_raw:
                otmetki_idx = plan_raw.find("ОТМЕТКИ")
                if otmetki_idx >= 0:
                    block = plan_raw[otmetki_idx:otmetki_idx + 10000]
                    src = "секция ОТМЕТКИ"
                else:
                    block = plan_raw
                    src = "весь Плана"
                checked = block.count('☑') + block.count('☒')
                unchecked = block.count('☐')
                print(f"[item 4 checkbox] source={src}, checked={checked}, unchecked={unchecked}")
                if unchecked > 0 and checked == 0:
                    verdict["ok"] = False
                    verdict["nok"] = True
                    verdict["reason"] = (
                        f"[авто-NOK: в Плане {src}: ни одного отмеченного чекбокса "
                        f"(☐×{unchecked}, ☑/☒×0). Вид аудита обязан быть выбран галочкой — "
                        f"если модель сказала иначе, это ошибка распознавания.]"
                    )
                else:
                    # Достаём отмеченные виды из секции ОТМЕТКИ (если есть подписи)
                    # и из текста Плана. Считаем OK, если вид из Приказа есть среди
                    # отмеченных в Плане. «Расширение области сертификации» —
                    # допустимое дополнение, не повод для NOK.
                    AUDIT_TYPE_KEYWORDS = {
                        "первый инспекционный": ["первый инспекцион", "1 ик", "1ик", "первого инспекцион"],
                        "второй инспекционный": ["второй инспекцион", "2 ик", "2ик", "второго инспекцион"],
                        "внеплановый инспекционный": ["внеплановый инспекцион", "внеплановый ик"],
                        "первый этап первичного": ["первый этап первичного", "1 этап первичной", "1-й этап перв"],
                        "второй этап первичного": ["второй этап первичного", "2 этап первичной", "2-й этап перв"],
                        "ресертификационный": ["ресертификационн"],
                        "дополнительный": ["дополнительный аудит"],
                        "расширение области": ["расширение области"],
                    }
                    # Виды, отмеченные в Плане: ищем по соседству с ☑/☒.
                    marked_in_plan = set()
                    for m in re.finditer(r"[☑☒]\s*([^☐☑☒\n]{4,120})", plan_raw):
                        label = m.group(1).lower().replace("ё", "е")
                        for canon, kws in AUDIT_TYPE_KEYWORDS.items():
                            if any(k in label for k in kws):
                                marked_in_plan.add(canon)

                    # Вид аудита из Приказа: ищем в файлах с «приказ» в имени.
                    prikaz_text = ""
                    for fn, txt in all_texts.items():
                        if ("приказ" in fn.lower() or "эг" in fn.lower()) and txt:
                            prikaz_text += " " + txt
                    prikaz_lc = prikaz_text.lower().replace("ё", "е")
                    type_in_prikaz = None
                    # Приоритет «инспекционный контроль» (формальный вид аудита),
                    # «расширение области» в Приказе — расширение, не сам тип.
                    priority = [
                        "первый инспекционный", "второй инспекционный", "внеплановый инспекционный",
                        "ресертификационный", "первый этап первичного", "второй этап первичного",
                        "дополнительный",
                    ]
                    for canon in priority:
                        kws = AUDIT_TYPE_KEYWORDS[canon]
                        if any(k in prikaz_lc for k in kws):
                            type_in_prikaz = canon
                            break

                    print(
                        f"[item 4 type-check] marked_in_plan={sorted(marked_in_plan)} "
                        f"type_in_prikaz={type_in_prikaz}"
                    )

                    if type_in_prikaz and type_in_prikaz in marked_in_plan:
                        extras = sorted(marked_in_plan - {type_in_prikaz})
                        verdict["ok"] = True
                        verdict["nok"] = False
                        reason = (
                            f"Вид аудита из Приказа «{type_in_prikaz}» отмечен в Плане."
                        )
                        if extras:
                            reason += (
                                f" Дополнительно в Плане отмечено: {', '.join(extras)} "
                                f"(допустимое сопровождение)."
                            )
                        verdict["reason"] = reason
                    elif type_in_prikaz and marked_in_plan and type_in_prikaz not in marked_in_plan:
                        verdict["ok"] = False
                        verdict["nok"] = True
                        verdict["reason"] = (
                            f"Вид аудита из Приказа «{type_in_prikaz}» НЕ отмечен в Плане. "
                            f"В Плане отмечено: {', '.join(sorted(marked_in_plan)) or 'ничего'}."
                        )

            # Пункт 6 (idx=5, договор): детерминированная сверка реквизитов.
            # Из Плана извлекаем номер и дату договора (пункт 2 «Основание»),
            # затем ищем эти же реквизиты в любом из загруженных контрактных файлов.
            # Если совпадает — форс-OK; модель часто выдаёт самопротиворечивый NOK
            # с reason «...совпадает...совпадает».
            if idx == 5 and plan_raw:
                m = re.search(
                    r"Договор[^№\n]*(?:№\s*)?([A-Za-zА-Яа-я0-9\-/_\.\s\(\)]{4,40}?)\s*от\s*(\d{1,2}[.\-/_ ]\d{1,2}[.\-/_ ]\d{2,4})",
                    plan_raw,
                    flags=re.IGNORECASE,
                )
                if m:
                    plan_num_raw = m.group(1).strip()
                    plan_date_raw = m.group(2).strip()
                    _digits = lambda s: re.sub(r"\D+", "", s or "")
                    plan_num_norm = _digits(plan_num_raw)
                    pd = re.match(r"(\d{1,2})[.\-/_ ](\d{1,2})[.\-/_ ](\d{2,4})", plan_date_raw)
                    plan_date_norm = None
                    if pd:
                        d, mo, y = int(pd.group(1)), int(pd.group(2)), int(pd.group(3))
                        if y < 100:
                            y += 2000
                        plan_date_norm = (d, mo, y)

                    contract_files = [
                        f for f in relevant
                        if any(k in f.lower() for k in ("договор", "дог.", "дог_", "контракт"))
                    ]
                    matched_file = None
                    for cf in contract_files:
                        txt = all_texts.get(cf, "")
                        if not txt:
                            continue
                        # Хватает совпадения нормализованного номера ИЛИ даты в тексте файла.
                        num_hit = plan_num_norm and plan_num_norm in _digits(txt)
                        date_hit = False
                        if plan_date_norm:
                            for dm in re.finditer(r"(\d{1,2})[.\-/_ ](\d{1,2})[.\-/_ ](\d{2,4})", txt):
                                d, mo, y = int(dm.group(1)), int(dm.group(2)), int(dm.group(3))
                                if y < 100:
                                    y += 2000
                                if (d, mo, y) == plan_date_norm:
                                    date_hit = True
                                    break
                        if num_hit and date_hit:
                            matched_file = cf
                            break

                    if matched_file:
                        verdict["ok"] = True
                        verdict["nok"] = False
                        verdict["reason"] = (
                            f"Реквизиты договора из Плана (№{plan_num_raw} от {plan_date_raw}) "
                            f"подтверждены в файле «{matched_file}»."
                        )

            # Пункт 5 (idx=4, ОКВЭД/область сертификации):
            # Если из двух нужных файлов (Разбивка ОКВЭД и Сертификат ИГС) загружен
            # только один — переписать reason на понятный для аудитора, без
            # технического префикса «авто-NOK: режим сверки требует ...».
            if idx == 4:
                names_lc = [f.lower() for f in relevant]
                has_razbivka = any(
                    ("разбивк" in n) or ("оквэд" in n) for n in names_lc
                )
                has_cert = any(
                    ("сертификат" in n and "макет" not in n)
                    or (" игс" in f" {n} ") or ("(игс)" in n)
                    for n in names_lc
                )
                if (has_razbivka and not has_cert) or (has_cert and not has_razbivka):
                    missing = "Сертификат ИГС" if has_razbivka else "Разбивка кодов ОКВЭД"
                    verdict["ok"] = False
                    verdict["nok"] = True
                    verdict["reason"] = (
                        f"Не загружен файл «{missing}» — для сверки формулировки "
                        f"области применения СМК нужны оба источника (Разбивка ОКВЭД и Сертификат ИГС). "
                        f"Загрузите недостающий файл и повторите проверку."
                    )
                elif not has_razbivka and not has_cert:
                    verdict["ok"] = False
                    verdict["nok"] = True
                    verdict["reason"] = (
                        "Не загружены файлы «Разбивка кодов ОКВЭД» и «Сертификат ИГС» — "
                        "сверка формулировки области применения СМК невозможна."
                    )

            # Пункт 11 (idx=10): исключения из СТО. Только п. 7.1.3.5 (1 абзац) и
            # п. 7.1.4.3 (1 абзац) легитимны. Любой другой → NOK.
            if idx == 10 and plan_raw:
                # Находим раздел «8 Исключения из требований стандартов».
                # Раньше брали 2000 символов — это захватывало соседний раздел графика
                # с пунктами процессов (4.4.1, 4.3.2 и т.п.) и давало ложно-NOK.
                # Сужаем: берём до следующего нумерованного раздела ("9 ", "10 ", ...) или 700 символов — что раньше.
                excl_idx = plan_raw.find("Исключения из требований")
                if excl_idx >= 0:
                    raw_block = plan_raw[excl_idx:excl_idx + 2000]
                    # Обрезаем на ближайшем следующем заголовке раздела вида "9 Заголовок" / "10 Заголовок"
                    next_section = re.search(r'\n\s*(?:9|10|11|12)\s+[А-ЯA-Z]', raw_block)
                    if next_section:
                        block = raw_block[: next_section.start()]
                    else:
                        block = raw_block[:700]
                    # Учитываем только пункты, рядом с которыми есть контекст исключения
                    # ('п.', 'пункт', 'абзац', 'не применяется', 'исключ').
                    raw_pts_all = re.findall(r'\b(\d{1,2}(?:\.\d{1,2}){2,4})\b', block)
                    raw_pts = []
                    for p in raw_pts_all:
                        pos = block.find(p)
                        ctx = block[max(0, pos - 40): pos + 80].lower()
                        if any(k in ctx for k in ("п.", "пункт", "абзац", "не применя", "исключ")):
                            raw_pts.append(p)
                    # Дедуп с сохранением порядка
                    found_pts = list(dict.fromkeys(raw_pts))
                    # Допустимы только 7.1.3.5 и 7.1.4.3, и только если рядом есть "абзац"
                    illegitimate: list[str] = []
                    for p in found_pts:
                        is_allowed = p in ("7.1.3.5", "7.1.4.3")
                        if not is_allowed:
                            illegitimate.append(p)
                            continue
                        # Проверим, есть ли «абзац» в окрестности этого пункта
                        p_pos = block.find(p)
                        ctx = block[max(0, p_pos - 30):p_pos + 80].lower()
                        if "абзац" not in ctx:
                            illegitimate.append(f"{p} (исключён целиком, без пометки про 1 абзац)")
                    print(f"[item 11 exclusions] all={found_pts}, illegitimate={illegitimate}")
                    if illegitimate:
                        verdict["ok"] = False
                        verdict["nok"] = True
                        verdict["reason"] = (
                            f"[авто-NOK: нелегитимные исключения из СТО: {', '.join(illegitimate)}. "
                            f"Допустимы только п. 7.1.3.5 и п. 7.1.4.3, и только с пометкой «1 абзац».]"
                        )
                    elif found_pts and not verdict.get("ok"):
                        # Только допустимые исключения, но модель сказала NOK — переворачиваем
                        verdict["ok"] = True
                        verdict["nok"] = False
                        verdict["reason"] = (
                            f"[авто-OK: исключения в Плане — {', '.join(found_pts)}, все легитимны "
                            f"(7.1.3.5/7.1.4.3 с пометкой «1 абзац»).]"
                        )

            # Пункт 15 (idx=14): для каждой строки с «Процесс» в Плане должны быть
            # пункты 4.4.1, 4.4.2, 4.4.3.
            # Пункт 10 (idx=9, замечания/несоответствия предыдущего аудита):
            # детерминированная сверка АКТа и Плана.
            #   В Акте парсим п.24 «Количество значительных/малозначительных
            #   несоответствий» и общие фразы об отсутствии замечаний.
            #   В Плане ищем раздел «Замечания/несоответствия по результатам
            #   предыдущего аудита» (по тексту, не по номеру — в разных редакциях
            #   шаблона номер раздела может отличаться).
            if idx == 9:
                akt_files = [f for f in relevant if "акт" in f.lower()]
                akt_text = "\n".join(all_texts.get(f, "") for f in akt_files)

                def _count_nc(text: str, kind: str) -> int:
                    pat = re.compile(
                        r"Количество\s+" + kind + r"\s+несоответствий"
                        r"(?!\s*,\s*устран)\s*([0-9]+|\s*\|)",
                        re.IGNORECASE,
                    )
                    total = 0
                    for m in pat.finditer(text):
                        v = m.group(1).strip()
                        if v and v != "|":
                            try:
                                total += int(v)
                            except ValueError:
                                pass
                    return total

                if not akt_text:
                    verdict["ok"] = False
                    verdict["nok"] = True
                    verdict["reason"] = (
                        "Файл «Акт предыдущего аудита» не загружен — сверить замечания не с чем."
                    )
                else:
                    sig = _count_nc(akt_text, "значительных")
                    minor = _count_nc(akt_text, "малозначительных")
                    akt_lc = akt_text.lower()
                    explicit_clean = any(
                        ph in akt_lc
                        for ph in (
                            "несоответствий не выявлен",
                            "замечаний не выявлен",
                            "замечания не выявлен",
                            "без замечаний",
                            "несоответствия не выявлен",
                        )
                    )
                    plan_lc = (plan_raw or "").lower()
                    has_plan_section = bool(
                        re.search(
                            r"замечани[яе]\s*[\\/]?\s*несоответстви",
                            plan_lc,
                        )
                        or re.search(
                            r"замечани[яе]\s+.{0,60}предыдущ",
                            plan_lc,
                        )
                        or re.search(
                            r"несоответстви[яе]\s+.{0,60}предыдущ",
                            plan_lc,
                        )
                    )

                    total_nc = sig + minor
                    akt_file_label = akt_files[0] if akt_files else "Акт"

                    if total_nc == 0 and explicit_clean:
                        verdict["ok"] = True
                        verdict["nok"] = False
                        verdict["reason"] = (
                            f"В Акте предыдущего аудита («{akt_file_label}») замечаний/"
                            f"несоответствий не выявлено — раздел в Плане не требуется."
                        )
                    elif total_nc == 0 and not explicit_clean:
                        # Чисел нет, но и явной формулировки «не выявлено» тоже.
                        # Считаем OK, но в reason указываем, что проверка по факту прошла мягко.
                        verdict["ok"] = True
                        verdict["nok"] = False
                        verdict["reason"] = (
                            f"В Акте предыдущего аудита («{akt_file_label}») числовых "
                            f"показателей несоответствий не обнаружено и явных упоминаний "
                            f"замечаний нет — отмечено как OK."
                        )
                    else:
                        # В Акте есть несоответствия — проверяем, отражены ли они в Плане.
                        parts = []
                        if sig:
                            parts.append(f"значительных: {sig}")
                        if minor:
                            parts.append(f"малозначительных: {minor}")
                        counts_str = ", ".join(parts)
                        if has_plan_section:
                            verdict["ok"] = True
                            verdict["nok"] = False
                            verdict["reason"] = (
                                f"В Акте предыдущего аудита («{akt_file_label}») выявлено "
                                f"{counts_str} несоответствий; в Плане раздел про замечания/"
                                f"несоответствия предыдущего аудита присутствует."
                            )
                        else:
                            verdict["ok"] = False
                            verdict["nok"] = True
                            verdict["reason"] = (
                                f"В Акте предыдущего аудита («{akt_file_label}») выявлено "
                                f"{counts_str} несоответствий, но в Плане раздел про "
                                f"замечания/несоответствия предыдущего аудита отсутствует — "
                                f"нечем подтвердить, что они учтены/закрыты."
                            )
                print(
                    f"[item 10 nc] sig={ _count_nc(akt_text, 'значительных') if akt_text else '-'} "
                    f"minor={ _count_nc(akt_text, 'малозначительных') if akt_text else '-'} "
                    f"-> {'OK' if verdict.get('ok') else 'NOK'}"
                )

            if idx == 14 and plan_raw:
                # Окно после «Процесс П<цифра>» — допускаем переносы строк,
                # т.к. в извлечённом docx ячейки таблицы могут разделяться \n.
                process_rows = re.findall(
                    r'(Процесс\s+П\d[\s\S]{0,500})', plan_raw
                )
                if process_rows:
                    missing = []
                    for row in process_rows:
                        # Достаточно проверить наличие всех трёх номеров в окне 500 символов
                        if not all(p in row for p in ("4.4.1", "4.4.2", "4.4.3")):
                            missing.append(row[:60])
                    print(f"[item 15 processes] rows={len(process_rows)}, missing_clauses={len(missing)}")
                    if not missing:
                        verdict["ok"] = True
                        verdict["nok"] = False
                        if "[авто-OK" not in (verdict.get("reason") or ""):
                            verdict["reason"] = (
                                f"[авто-OK: найдено {len(process_rows)} строк с «Процесс», "
                                f"в каждой присутствуют пункты 4.4.1, 4.4.2, 4.4.3.] "
                                + (verdict.get("reason") or "")
                            )
                    elif missing and verdict.get("ok"):
                        verdict["ok"] = False
                        verdict["nok"] = True
                        verdict["reason"] = (
                            f"[авто-NOK: в {len(missing)} из {len(process_rows)} строк с «Процесс» "
                            f"отсутствует хотя бы один из 4.4.1/4.4.2/4.4.3.]"
                        )

            # Пункт 12 (idx=11): совещания. Модель часто цепляется за маркеры ИИ13/ИИ9
            # из problems_hint и заявляет «отсутствует», даже когда в Плане совещания есть.
            # Считаем явные вхождения и переопределяем при достаточном покрытии.
            if idx == 11 and not verdict.get("ok") and evidence:
                ev_lc = evidence.lower()
                preliminary = ev_lc.count("предварительное совещан")
                final_meet = ev_lc.count("заключительное совещан")
                intermediate = ev_lc.count("промежуточное совещан") + ev_lc.count("рабочее совещан")
                reason_lc = (verdict.get("reason") or "").lower()
                deny_signals = ("отсутств", "не подтвержд", "не найден", "не запланирован", "не обнаруж")
                model_denied = any(s in reason_lc for s in deny_signals)
                # Базовая проверка: есть хотя бы одно Предварительное, одно Заключительное
                # и какие-то Рабочие/Промежуточные. Этого минимума достаточно.
                if model_denied and preliminary >= 1 and final_meet >= 1 and intermediate >= 1:
                    verdict["ok"] = True
                    verdict["nok"] = False
                    verdict["reason"] = (
                        f"[авто-OK: в Плане найдены Предварительные совещания ({preliminary}), "
                        f"Заключительные ({final_meet}), Рабочие/Промежуточные ({intermediate}) — "
                        f"минимальные требования выполнены] "
                        + (verdict.get("reason") or "")
                    )

            # Пункт 13 (idx=12): инструктаж. Часто модель пишет «инструктаж не подтверждён»,
            # хотя строка «Инструктаж по технике безопасности» физически присутствует в Плане.
            # Если в evidence слово «инструктаж» встречается, а модель сказала NOK по причине
            # «не подтверждено / не найдено / отсутствует» — переопределяем на OK.
            if idx == 12 and not verdict.get("ok") and evidence:
                inst_count = evidence.lower().count("инструктаж")
                reason_lc = (verdict.get("reason") or "").lower()
                deny_signals = ("не подтвержд", "не найден", "отсутств", "пропущен", "не обнаруж", "маркер")
                if inst_count >= 1 and any(s in reason_lc for s in deny_signals):
                    verdict["ok"] = True
                    verdict["nok"] = False
                    verdict["reason"] = (
                        f"[авто-OK: модель не засчитала инструктаж, но в evidence слово 'Инструктаж' "
                        f"встречается {inst_count} раз — требование выполнено в самом Плане] "
                        + (verdict.get("reason") or "")
                    )

            verdict["_files_checked"] = relevant
            # Дописываем в обоснование список файлов, в которых искалась информация
            if relevant:
                shown = relevant[:10]
                files_note = "; ".join(shown)
                if len(relevant) > len(shown):
                    files_note += f" (и ещё {len(relevant) - len(shown)})"
                suffix = f" Искалось в файлах ({len(relevant)}): {files_note}."
            else:
                suffix = " Искалось во всех загруженных файлах (релевантных по разделу/ключевым словам не найдено)."
            verdict["reason"] = (verdict.get("reason") or "").rstrip() + suffix
            checklist_results.append(verdict)

            print(f"[item {idx+1}/{total}] {'OK' if verdict['ok'] else 'NOK'}: {item['area'][:50]}")
        except Exception as e:
            print(f"[item {idx+1}] Ошибка анализа: {e}")
            checklist_results.append({
                "ok": False,
                "nok": True,
                "reason": f"Ошибка анализа: {str(e)[:120]}",
                "ii_data_found": ", ".join(item.get("ii_markers", [])) or "нет",
                "evidence_quote": "ошибка обработки",
                "source_file": "—",
                "_files_checked": [],
            })

    processing_status.update({
        "stage": "done",
        "current": total,
        "total": total,
        "message": "Обработка завершена",
        "detail": "",
    })

    return {"header": header_data, "checklist": checklist_results}


def _svod_extract_sign_date(text: str):
    """Дата подписи из блока «И.О. Фамилия … ДД.ММ.ГГГГ». Возвращает max date или None.
    Сроки/дедлайны вида «до 09.04.2026» не матчатся (нет ФИО перед датой)."""
    from datetime import date
    pat = re.compile(r'[А-ЯЁ]\.\s*[А-ЯЁ]\.\s*[А-ЯЁ][а-яё]+[^\d]{0,12}(\d{2}\.\d{2}\.\d{4})')
    dates = []
    for m in pat.finditer(text or ""):
        try:
            d, mo, y = m.group(1).split('.')
            dates.append(date(int(y), int(mo), int(d)))
        except ValueError:
            pass
    return max(dates) if dates else None


def _svod_working_days(d1, d2) -> int:
    """Число рабочих дней (Пн–Пт) между датами, не считая первую, включая последнюю."""
    from datetime import timedelta
    if d2 < d1:
        d1, d2 = d2, d1
    n, cur = 0, d1 + timedelta(days=1)
    while cur <= d2:
        if cur.weekday() < 5:
            n += 1
        cur += timedelta(days=1)
    return n


def _svod_shift_count(text: str):
    """Максимальное число перед словом «смен(а)» в тексте, либо None."""
    nums = [int(m.group(1)) for m in re.finditer(r'(\d+)\s*смен', (text or "").lower())]
    return max(nums) if nums else None


def _svod_deterministic_postcheck(idx: int, verdict: dict, svod_doc_text: str, all_texts: dict) -> dict:
    """Детерминированные пост-проверки поверх вердикта модели. Меняют вердикт ТОЛЬКО
    когда удаётся уверенно извлечь нужные значения; иначе оставляют ответ GigaChat."""
    # №10 (idx 9) — режим работы: число смен в Сводном акте vs Расчёт трудоёмкости
    if idx == 9:
        i = (svod_doc_text or "").lower().find("режим работы")
        svod_shifts = _svod_shift_count(svod_doc_text[i:i + 250]) if i >= 0 else None
        trud_shifts = None
        for fn, txt in all_texts.items():
            if any(k in fn.lower() for k in ("трудоемкост", "трудоёмкост")):
                m = re.search(r"количество смен\s*\|\s*(\d+)", (txt or "").lower())
                if m:
                    trud_shifts = int(m.group(1))
                break
        if svod_shifts is not None and trud_shifts is not None:
            ok = svod_shifts == trud_shifts
            verdict["ok"], verdict["nok"] = ok, not ok
            verdict["reason"] = (
                f"[детерм.] Смен в Сводном акте (режим работы): {svod_shifts}; "
                f"в Расчёте трудоёмкости: {trud_shifts}. "
                + ("Совпадает — OK." if ok else "Расхождение — NOK.")
            )
        return verdict

    # №14 (idx 13) — дата подписания ≤ 20 рабочих дней от даты Акта 2 этапа
    if idx == 13:
        d_svod = _svod_extract_sign_date(svod_doc_text)
        d_akt = None
        for fn, txt in all_texts.items():
            if any(k in fn.lower() for k in ("акт р", "акт 2", "результатам аудита", "акт по результ")):
                d_akt = _svod_extract_sign_date(txt)
                if d_akt:
                    break
        if d_svod and d_akt:
            late = d_svod < d_akt
            wd = _svod_working_days(d_akt, d_svod)
            ok = (not late) and wd <= 20
            verdict["ok"], verdict["nok"] = ok, not ok
            verdict["reason"] = (
                f"[детерм.] Дата подписания Сводного акта: {d_svod.strftime('%d.%m.%Y')}; "
                f"дата Акта 2 этапа: {d_akt.strftime('%d.%m.%Y')}; между ними {wd} рабочих дней. "
                "Требование ≤20 рабочих дней. "
                + ("OK." if ok else ("NOK (Сводный акт подписан раньше Акта 2 этапа)." if late
                                     else "NOK (просрочка более 20 рабочих дней)."))
            )
        return verdict

    return verdict


def process_checklist_svod(api_key: str, all_texts: dict,
                            checklist_structure: list,
                            model: str = "GigaChat",
                            svod_doc_text: str = "") -> dict:
    """
    Per-item обработка чек-листа «Сводный акт» (14 пунктов).
    Без шапки Заявителя. Главный проверяемый документ — Сводный акт (svod_doc_text) —
    подставляется в evidence каждого пункта. Источники для сверки берутся из all_texts
    по SVOD_ITEM_RULES[idx]['file_keywords']. Пункты с manual=True не гоняются через
    модель — помечаются «на ручную проверку».
    """
    global processing_status

    total = len(checklist_structure)
    checklist_results = []

    for idx, item in enumerate(checklist_structure):
        item_no = item.get("item_no", idx + 1)
        processing_status.update({
            "stage": "verify",
            "current": idx + 1,
            "total": total,
            "message": f"[Сводный акт] Пункт {item_no}/{total}: {item['area'][:50]}",
            "detail": "",
        })

        rule = SVOD_ITEM_RULES.get(idx, {})
        is_manual = bool(rule.get("manual"))

        try:
            file_names = list(all_texts.keys())
            keyword_hit = False
            if rule.get("file_keywords"):
                relevant = _files_by_keyword(file_names, rule["file_keywords"])
                if relevant:
                    keyword_hit = True
                else:
                    relevant = find_relevant_files_for_item(item, all_texts)
            else:
                relevant = find_relevant_files_for_item(item, all_texts)

            # Авто-NOK: пункт требует сверки, но файл-источник не загружен
            if rule.get("comparison") and rule.get("file_keywords") and not keyword_hit:
                kw_display = ", ".join(rule["file_keywords"][:5])
                checklist_results.append({
                    "ok": False,
                    "nok": True,
                    "reason": (
                        f"[авто-NOK: не загружен файл-источник для сверки. "
                        f"Ожидалось имя с подстрокой из: {kw_display}.]"
                    ),
                    "ii_data_found": "",
                    "evidence_quote": "",
                    "source_file": "не найдено",
                })
                print(f"[svod {item_no}/{total}] NOK (нет файла-источника)")
                continue

            evidence = build_evidence_pack(relevant, all_texts, max_chars=150000)
            if svod_doc_text:
                svod_block = (
                    "=== ФАЙЛ: СВОДНЫЙ АКТ (проверяемый документ) ===\n"
                    + svod_doc_text[:40000]
                )
                evidence = svod_block + ("\n\n" + evidence if evidence else "")

            processing_status["detail"] = f"проверка по {len(relevant)} файлу(ам)..."
            verdict = verify_item_strict(
                api_key, item, {}, evidence, model=model,
                extra_instructions=rule.get("extra_instructions", ""),
                comparison=bool(rule.get("comparison", False)),
                compare_fields=rule.get("compare_fields", ""),
                expected_file_keywords=rule.get("file_keywords") or None,
            )

            if verdict.get("ok") and not is_manual:
                processing_status["detail"] = "adversarial-перепроверка..."
                verdict = adversarial_recheck(
                    api_key, item, {}, evidence, verdict, model=model,
                    comparison=bool(rule.get("comparison", False)),
                    compare_fields=rule.get("compare_fields", ""),
                )

            # Детерминированные пост-чеки (числа/даты) поверх вердикта модели
            if not is_manual:
                verdict = _svod_deterministic_postcheck(idx, verdict, svod_doc_text, all_texts)

            # Ручные пункты: GigaChat извлёк данные, но авто-вердикт не ставим —
            # помечаем «ПРОВЕРИТЬ ВРУЧНУЮ» и оставляем обе клетки ОК/NOK пустыми.
            if is_manual:
                extracted = (verdict.get("reason") or "").strip()
                verdict["ok"] = False
                verdict["nok"] = False
                verdict["reason"] = ("ПРОВЕРИТЬ ВРУЧНУЮ. " + extracted).strip()

            verdict.setdefault("ii_data_found", "")
            checklist_results.append(verdict)
            mark = "РУЧНАЯ" if is_manual else ("OK" if verdict.get("ok") else "NOK")
            print(f"[svod {item_no}/{total}] {mark}")
        except Exception as e:
            print(f"[svod {item_no}/{total}] ошибка: {e}")
            checklist_results.append({
                "ok": False,
                "nok": True,
                "reason": f"Ошибка проверки пункта: {e}",
                "ii_data_found": "",
                "evidence_quote": "",
                "source_file": "",
            })

    processing_status.update({
        "stage": "done",
        "current": total,
        "total": total,
        "message": "[Сводный акт] Обработка завершена",
        "detail": "",
    })

    return {"header": {}, "checklist": checklist_results}


def validate_filled_document(output_path: str, checklist_structure: list[dict], extracted_data: dict) -> dict:
    """
    Проверка корректности заполненного документа.
    Возвращает отчёт о валидации.
    """
    doc = Document(output_path)
    issues = []
    warnings = []
    ok_count = 0
    nok_count = 0
    filled_count = 0
    total_checklist_rows = 0
    
    checklist_data = extracted_data.get("checklist", [])
    
    for table in doc.tables:
        first_row = table.rows[0]
        first_cell_text = first_row.cells[0].text.strip()
        cols_count = len(first_row.cells)
        
        # === Проверка шапки ===
        if cols_count == 2:
            header_data = extracted_data.get("header", {})
            for row in table.rows:
                cells_text = [cell.text.strip() for cell in row.cells]
                full_row_text = " ".join(cells_text)
                
                # Проверяем заполнено ли каждое поле шапки
                if any(key in full_row_text for key in ["Наименования Заявителя", "Вид аудита", "Даты проведения", "РЭГ"]):
                    # Есть ли данные в правой ячейке
                    right_cell = row.cells[1].text.strip() if len(row.cells) > 1 else ""
                    left_cell = row.cells[0].text.strip()
                    
                    # Определяем какое поле ожидаем
                    expected_field = None
                    if "Наименования Заявителя" in full_row_text:
                        expected_field = "Наименование Заявителя"
                    elif "Вид аудита" in full_row_text:
                        expected_field = "Вид аудита"
                    elif "Даты проведения" in full_row_text:
                        expected_field = "Даты проведения"
                    elif "РЭГ" in full_row_text and len(full_row_text) < 50:
                        expected_field = "РЭГ"
                    
                    if expected_field and expected_field in header_data:
                        # Проверяем что значение записано
                        combined = left_cell + " " + right_cell
                        if header_data[expected_field] not in combined and not right_cell:
                            warnings.append(f"Поле шапки '{expected_field}' возможно не заполнено")
        
        # === Проверка чек-листа ===
        if cols_count >= 5 and "Область проверки" in first_cell_text:
            row_idx = 0
            for row in table.rows[1:]:  # Пропускаем заголовок
                row_idx += 1
                total_checklist_rows += 1
                
                area = row.cells[0].text.strip()
                ok_cell = row.cells[2].text.strip()
                nok_cell = row.cells[3].text.strip()
                prob_cell = row.cells[4].text.strip()
                
                if not area:
                    continue
                
                has_ok = "☑" in ok_cell or "OK" in ok_cell.upper()
                has_nok = "☒" in nok_cell or "NOK" in nok_cell.upper()
                
                # Проверяем что стоит ровно один маркер
                if has_ok and has_nok:
                    issues.append(f"Строка {row_idx} ('{area[:40]}...'): стоят ОБА маркера OK и NOK!")
                elif not has_ok and not has_nok:
                    issues.append(f"Строка {row_idx} ('{area[:40]}...'): НЕТ маркера OK или NOK!")
                elif has_ok:
                    ok_count += 1
                elif has_nok:
                    nok_count += 1
                
                filled_count += 1
    
    # Проверяем что все строки из структуры прошли валидацию
    if total_checklist_rows < len(checklist_structure):
        issues.append(f"В документе заполено {total_checklist_rows} строк, ожидалось {len(checklist_structure)}")
    
    # Проверяем количество пунктов от GigaChat
    if len(checklist_data) != len(checklist_structure):
        warnings.append(f"GigaChat вернул {len(checklist_data)} пунктов, в документе {len(checklist_structure)}")
    
    is_valid = len(issues) == 0
    
    return {
        "valid": is_valid,
        "issues": issues,
        "warnings": warnings,
        "stats": {
            "total_rows": total_checklist_rows,
            "expected_rows": len(checklist_structure),
            "ok_count": ok_count,
            "nok_count": nok_count,
            "filled_count": filled_count
        }
    }


@app.get("/")
async def root():
    return {"message": "Audit Plan Filler API", "version": "1.0.0"}


@app.get("/api/validate/{filename:path}")
async def validate_result(filename: str):
    """Проверка корректности заполненного документа.
    filename может быть '<session_id>/<file>' или просто '<file>'."""
    file_path = (OUTPUT_DIR / filename).resolve()
    try:
        file_path.relative_to(OUTPUT_DIR.resolve())
    except ValueError:
        raise HTTPException(status_code=400, detail="Некорректный путь")
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден")
    
    # Загружаем последнюю структуру чек-листа из сохранённых данных
    # Для этого прочитаем сохранённый результат если есть
    # Или просто вердим базовую валидацию
    try:
        doc = Document(str(file_path))
        checklist_rows = []
        header_filled = []
        
        for table in doc.tables:
            cols = len(table.rows[0].cells)
            
            if cols == 2:  # Шапка
                for row in table.rows:
                    text = " ".join(cell.text.strip() for cell in row.cells)
                    right = row.cells[1].text.strip() if len(row.cells) > 1 else ""
                    if any(k in text for k in ["Заявител", "Вид аудита", "Даты", "РЭГ"]):
                        header_filled.append(bool(right and len(right) > 2))
            
            if cols >= 5:
                first_cell = table.rows[0].cells[0].text.strip()
                if "Область проверки" in first_cell:
                    for row in table.rows[1:]:
                        area = row.cells[0].text.strip()
                        ok = row.cells[2].text.strip()
                        nok = row.cells[3].text.strip()
                        if area:
                            has_ok = "☑" in ok
                            has_nok = "☒" in nok
                            checklist_rows.append({
                                "area": area[:60],
                                "ok": has_ok,
                                "nok": has_nok,
                                "both": has_ok and has_nok,
                                "none": not has_ok and not has_nok
                            })
        
        issues = []
        for i, row in enumerate(checklist_rows):
            if row["both"]:
                issues.append(f"Строка {i+1}: ОБА маркера OK и NOK")
            elif row["none"]:
                issues.append(f"Строка {i+1}: НЕТ маркера")

        # Структурные проблемы заполнения
        if len(checklist_rows) == 0:
            issues.append("Чек-лист не распознан: ни одной строки не найдено (проверь шаблон)")
        if header_filled and not all(header_filled):
            missing = len(header_filled) - sum(header_filled)
            issues.append(f"Шапка заполнена не полностью: {missing} из {len(header_filled)} полей пусты")

        ok_count = sum(1 for r in checklist_rows if r["ok"])
        nok_count = sum(1 for r in checklist_rows if r["nok"])

        # NOK в чек-листе — это содержательная проблема, тоже отражаем как issue.
        if nok_count > 0:
            issues.append(f"Обнаружены замечания (NOK): {nok_count} из {len(checklist_rows)} пунктов")

        return {
            "valid": len(issues) == 0,
            "header_filled": f"{sum(header_filled)}/{len(header_filled)} полей",
            "checklist_total": len(checklist_rows),
            "ok_count": ok_count,
            "nok_count": nok_count,
            "issues": issues
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка валидации: {str(e)}")


@app.post("/api/settings/gigachat")
async def save_gigachat_settings(settings: GigaChatSettings):
    """Сохранение настроек GigaChat (ключ + выбранная модель)"""
    settings_file = Path(__file__).parent / "gigachat_settings.json"
    existing = _load_settings()
    existing["api_key"] = settings.api_key
    if settings.model is not None:
        existing["model"] = settings.model.strip()
    with open(settings_file, "w", encoding="utf-8") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)
    return {"status": "ok", "message": "Настройки сохранены"}


@app.get("/api/settings/gigachat")
async def get_gigachat_settings():
    """Получение настроек GigaChat"""
    settings = _load_settings()
    return {
        "api_key": settings.get("api_key", ""),
        "model": settings.get("model", "") or DEFAULT_GIGACHAT_MODEL,
        "default_model": DEFAULT_GIGACHAT_MODEL,
    }


@app.post("/api/session/new")
async def create_session():
    """Создать новую пустую сессию загрузок."""
    sid, _ = _resolve_session_dir(None)
    return {"status": "ok", "session_id": sid}


@app.post("/api/upload")
async def upload_files(
    files: list[UploadFile] = File(...),
    session_id: Optional[str] = Form(None),
    block: Optional[str] = Form(None),
):
    """Загрузка документов (docx, docm, pdf, xlsx) в папку сессии.
    block — необязательный тэг блока в UI: 'checklist' | 'plan' | 'sources'."""
    sid, sdir = _resolve_session_dir(session_id)

    meta = _load_session_meta(sid)
    if not meta.get("created_at"):
        _save_session_meta(sid, created_at=datetime.now().isoformat(timespec="seconds"))

    uploaded = []
    block_files = list(meta.get("blocks", {}).get(block, [])) if block else []
    for file in files:
        if not file.filename:
            continue

        file_path = sdir / file.filename
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        uploaded.append({
            "filename": file.filename,
            "path": str(file_path),
            "size": file_path.stat().st_size
        })
        if block and file.filename not in block_files:
            block_files.append(file.filename)

    if block:
        blocks = dict(meta.get("blocks") or {})
        blocks[block] = block_files
        _save_session_meta(sid, blocks=blocks)

    return {"status": "ok", "session_id": sid, "uploaded_files": uploaded, "count": len(uploaded)}


@app.post("/api/process")
async def process_documents(
    api_key: str = Form(...),
    template_file: str = Form(...),
    session_id: str = Form(...),
    plan_doc_file: Optional[str] = Form(None),
    plan_doc_file_2: Optional[str] = Form(None),
    template_file_2: Optional[str] = Form(None),
):
    """
    Обработка документов через GigaChat и заполнение чек-листа(ов).
    В поле 1 может быть один или два чек-листа: «План АУДИТА» и/или «Сводный акт».
    Тип каждого определяется автоматически; обрабатываются по очереди (План → Сводный),
    на выходе — отдельный заполненный файл на каждый чек-лист.
    """
    if not _SESSION_ID_RE.match(session_id or ""):
        raise HTTPException(status_code=400, detail="Некорректный session_id")
    sdir = UPLOAD_DIR / session_id
    if not sdir.is_dir():
        raise HTTPException(status_code=400, detail="Сессия не найдена")

    # Проверяем наличие загруженных файлов в папке сессии
    docx_files = list(sdir.glob("*.docx"))
    docm_files = list(sdir.glob("*.docm"))
    pdf_files = list(sdir.glob("*.pdf"))
    xlsx_files = list(sdir.glob("*.xlsx"))
    all_docs = docx_files + docm_files + pdf_files + xlsx_files

    if not all_docs:
        raise HTTPException(status_code=400, detail="Нет загруженных документов")

    # Поле 2 (проверяемый документ: План аудита или Сводный акт) обязательно.
    if not plan_doc_file:
        raise HTTPException(
            status_code=400,
            detail="Не загружен проверяемый документ (поле 2). Без него сверка невозможна — пайплайн остановлен."
        )

    # === Чек-листы из поля 1: один или два. Определяем тип каждого. ===
    checklist_names = [template_file] + ([template_file_2] if template_file_2 else [])
    checklists = []
    for name in checklist_names:
        cpath = sdir / name
        if not cpath.exists():
            possible_names = ["ИИ шаблон плана.docm", "ИИ -ЧК -План АУДИТА.docx", "ЧК -Сводный акт.docx"]
            for alt in possible_names:
                if (sdir / alt).exists():
                    cpath = sdir / alt
                    break
            else:
                raise HTTPException(status_code=400, detail=f"Чек-лист '{name}' не найден в сессии")
        ctype = detect_checklist_type(str(cpath))
        checklists.append({"name": cpath.name, "path": cpath, "type": ctype})
        print(f"Чек-лист '{cpath.name}' → тип: {ctype}")

    recognized = [c for c in checklists if c["type"] in ("plan", "svod")]
    if not recognized:
        raise HTTPException(
            status_code=400,
            detail="Не удалось распознать тип чек-листа в поле 1 (ожидался «План АУДИТА» или «Сводный акт»)."
        )
    # Порядок обработки: сначала План, затем Сводный акт
    recognized.sort(key=lambda c: 0 if c["type"] == "plan" else 1)

    # Источники (поле 3) — всё, кроме самих чек-листов. Должен быть хотя бы 1 файл
    # помимо проверяемого документа из поля 2.
    checklist_set = {c["name"] for c in checklists}
    field2_set = {f for f in (plan_doc_file, plan_doc_file_2) if f}
    excluded_names = checklist_set | field2_set
    sources_count = sum(1 for d in all_docs if d.name not in excluded_names)
    if sources_count == 0:
        raise HTTPException(
            status_code=400,
            detail="Нет файлов-источников в поле 3. Загрузите Договор, Приказ ЭГ, Заявку и пр. — без них сверка невозможна."
        )

    # Извлекаем текст из всех документов КРОМЕ чек-листов.
    # ВАЖНО: эта работа CPU-bound (особенно OCR) — выносим в thread, иначе блокируется
    # event loop и /api/status перестаёт отвечать (фронт не видит прогресс).
    import asyncio

    source_docs = [d for d in all_docs if d.name not in checklist_set]
    total_files = len(source_docs)
    processing_status.update({
        "stage": "extract",
        "current": 0,
        "total": total_files,
        "message": f"Извлечение текста из {total_files} файлов",
        "detail": "",
    })

    EXTRACT_WORKERS = 3

    def _extract_one(doc_file: Path):
        """Извлечь текст одного файла. Возвращает (filename, text) или (filename, None) при ошибке."""
        suffix = doc_file.suffix.lower()
        try:
            if suffix in ('.docx', '.docm'):
                text = extract_text_from_docx(str(doc_file))
            elif suffix == '.pdf':
                text = extract_text_from_pdf(str(doc_file))
            elif suffix == '.xlsx':
                text = extract_text_from_xlsx(str(doc_file))
            else:
                print(f"Пропущен файл {doc_file.name} — неподдерживаемый формат")
                return (doc_file.name, None)
            return (doc_file.name, text)
        except Exception as e:
            print(f"ERR Ошибка чтения {doc_file.name}: {e}")
            return (doc_file.name, None)

    def _extract_all_texts_parallel():
        """Параллельное извлечение в пуле потоков (3 файла одновременно)."""
        from concurrent.futures import ThreadPoolExecutor, as_completed
        import threading

        all_texts = {}
        analyzed_files = []
        in_progress = set()
        lock = threading.Lock()
        completed = 0

        def update_status():
            with lock:
                names = ", ".join(sorted(in_progress)[:3])
                processing_status.update({
                    "stage": "extract",
                    "current": completed,
                    "total": total_files,
                    "message": f"Файлы {completed}/{total_files} (параллельно x{EXTRACT_WORKERS})",
                    "detail": f"в обработке: {names}" if names else "",
                })

        with ThreadPoolExecutor(max_workers=EXTRACT_WORKERS) as ex:
            future_to_file = {}
            for doc_file in source_docs:
                fut = ex.submit(_extract_one, doc_file)
                future_to_file[fut] = doc_file
                with lock:
                    in_progress.add(doc_file.name[:60])

            update_status()

            for fut in as_completed(future_to_file):
                doc_file = future_to_file[fut]
                fname, text = fut.result()
                with lock:
                    in_progress.discard(doc_file.name[:60])
                    completed += 1
                    if text is not None:
                        all_texts[fname] = text
                        analyzed_files.append(fname)
                        print(f"OK Проанализирован ({completed}/{total_files}): {fname}")
                update_status()

        return all_texts, analyzed_files

    all_texts, analyzed_files = await asyncio.to_thread(_extract_all_texts_parallel)

    print(f"\nВсего проанализировано файлов: {len(analyzed_files)}")
    print(f"Файлы: {analyzed_files}\n")

    processing_status.update({
        "stage": "extract",
        "current": total_files,
        "total": total_files,
        "message": "Чтение проверяемого документа (поле 2)...",
        "detail": "",
    })

    def _read_doc(p: Path) -> str:
        suffix = p.suffix.lower()
        if suffix in ('.docx', '.docm'):
            return extract_text_from_docx(str(p))
        elif suffix == '.pdf':
            return extract_text_from_pdf(str(p))
        elif suffix == '.xlsx':
            return extract_text_from_xlsx(str(p))
        return ""

    # Проверяемые документы из поля 2 (1 или 2 файла: План аудита и/или Сводный акт).
    field2_files = [plan_doc_file] + ([plan_doc_file_2] if plan_doc_file_2 else [])
    field2_texts = []  # [(name, text), ...]
    for name in field2_files:
        p = sdir / name
        if not p.exists():
            raise HTTPException(status_code=400, detail=f"Файл (поле 2) '{name}' не найден в сессии")
        try:
            txt = await asyncio.to_thread(_read_doc, p)
        except Exception as e:
            print(f"WARN: Не удалось прочитать документ поля 2 {name}: {e}")
            txt = ""
        field2_texts.append((name, txt))

    # Распределяем по типу: Сводный акт vs остальное (План аудита).
    field2_svod_text = ""
    field2_plan_text = ""
    for name, txt in field2_texts:
        if _looks_like_svod_doc(txt):
            if not field2_svod_text:
                field2_svod_text = txt
        elif not field2_plan_text:
            field2_plan_text = txt
    # Фолбэк: если План не распознан, берём первый файл поля 2.
    if not field2_plan_text and field2_texts:
        field2_plan_text = field2_texts[0][1]

    # Активная модель из настроек (с фолбэком на дефолт)
    active_model = _get_active_model()

    # Pre-flight: быстрая проверка GigaChat. Если auth/связь упали — ошибаемся сразу.
    processing_status.update({
        "stage": "preflight",
        "current": 0,
        "total": 1,
        "message": f"Проверка соединения с GigaChat ({active_model})...",
        "detail": "",
    })
    print("Preflight GigaChat...")
    preflight = await asyncio.to_thread(gigachat_preflight, api_key, 8.0, active_model)
    if not preflight["ok"]:
        raise HTTPException(
            status_code=503,
            detail=f"GigaChat недоступен ({preflight['stage']}): {preflight['detail']}"
        )
    print(f"Preflight OK. Доступные модели: {preflight['models']}")

    def _normalize_counts(extracted: dict, structure: list):
        """Приводим число результатов к числу пунктов чек-листа (обрезка/дополнение)."""
        if "checklist" not in extracted:
            return
        actual = extracted["checklist"]
        n = len(structure)
        if len(actual) > n:
            print(f"GigaChat вернул {len(actual)} пунктов, в документе {n}. Обрезано.")
            extracted["checklist"] = actual[:n]
        elif len(actual) < n:
            for i in range(len(actual), n):
                area_name = structure[i]['area'][:60] if i < len(structure) else f"пункт {i+1}"
                extracted["checklist"].append({
                    "ok": False, "nok": True,
                    "reason": f"Нет данных для проверки: {area_name}",
                    "problems": f"Нет данных для проверки: {area_name}",
                    "ii_data_found": "",
                })
        for item in extracted["checklist"]:
            item.setdefault("reason", item.get("problems", ""))
            item.setdefault("ii_data_found", "")

    session_out_dir = OUTPUT_DIR / session_id
    session_out_dir.mkdir(parents=True, exist_ok=True)

    outputs = []
    first_extracted = None
    first_validation = None

    # === Обрабатываем каждый распознанный чек-лист (План → Сводный акт) ===
    for c in recognized:
        cpath = c["path"]
        ctype = c["type"]
        print(f"Запуск пайплайна для чек-листа '{cpath.name}' (тип {ctype}, модель {active_model})...")

        if ctype == "plan":
            # Проверяемый документ для Плана = файл поля 2, распознанный как План аудита
            plan_doc_text = field2_plan_text
            if not plan_doc_text or len(plan_doc_text.strip()) < 100:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "Не удалось извлечь текст «Плана аудита» из поля 2. "
                        f"Прочитано символов: {len(plan_doc_text.strip()) if plan_doc_text else 0}. "
                        "Проверьте, что в поле 2 загружен правильный файл с Планом аудита."
                    )
                )
            try:
                template_text = await asyncio.to_thread(extract_text_from_docx, str(cpath))
            except Exception as e:
                print(f"WARN: текст шаблона Плана не извлечён: {e}")
                template_text = ""
            checklist_structure = extract_checklist_from_template(str(cpath))
            ii_references = extract_ii_references(str(cpath))
            print(f"[План] маркеры ИИ: {list(ii_references.keys())}, пунктов: {len(checklist_structure)}")

            try:
                extracted_data = await asyncio.to_thread(
                    process_checklist_advanced,
                    api_key, all_texts, checklist_structure, ii_references, active_model,
                    plan_doc_text
                )
            except Exception as e:
                processing_status.update({"stage": "error", "message": str(e)[:200]})
                return ProcessingResult(status="error", message=f"Ошибка пайплайна (План): {str(e)}")
            if "error" in extracted_data:
                return ProcessingResult(status="error", message=extracted_data["error"])

            _normalize_counts(extracted_data, checklist_structure)

            output_filename = "Заполненный_План_АУДИТА.docx"
            output_path = session_out_dir / output_filename
            output_ref = f"{session_id}/{output_filename}"
            processing_status.update({"stage": "fill", "message": "Заполнение Плана АУДИТА...", "detail": ""})
            try:
                await asyncio.to_thread(
                    fill_plan_with_checklist,
                    str(cpath), extracted_data, str(output_path), checklist_structure
                )
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Ошибка заполнения Плана: {str(e)}")

            processing_status.update({"stage": "validate", "message": "Валидация Плана...", "detail": ""})
            validation = await validate_result(output_ref)

            # Доп. сверка наименования юр.лица (План vs источники)
            try:
                applicant_check = await asyncio.to_thread(
                    cross_check_applicant_name, api_key, plan_doc_text, all_texts, active_model
                )
            except Exception as e:
                print(f"WARN: cross-check заявителя не выполнен: {e}")
                applicant_check = None
            if applicant_check and applicant_check.get("match") is False:
                note = (
                    f"Наименование заявителя в Плане и в источниках различается. "
                    f"План (поле 2): «{applicant_check.get('plan_name') or '—'}». "
                    f"Источники (поле 3): «{applicant_check.get('sources_name') or '—'}»."
                )
                if applicant_check.get("note"):
                    note += f" {applicant_check['note']}"
                notes = list(validation.get("notes") or [])
                notes.append(note)
                validation["notes"] = notes
                validation["applicant_check"] = applicant_check

        else:  # ctype == "svod"
            # Проверяемый документ для Сводного акта = файл поля 2, распознанный как
            # Сводный акт; иначе ищем среди источников (поле 3).
            svod_doc_text = field2_svod_text
            if not svod_doc_text:
                for fn, txt in all_texts.items():
                    if _looks_like_svod_doc(txt):
                        svod_doc_text = txt
                        print(f"[Сводный акт] проверяемый документ найден среди источников: {fn}")
                        break
            if not svod_doc_text or len(svod_doc_text.strip()) < 100:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "Не найден проверяемый документ «Сводный акт исследования (итог)». "
                        "Загрузите его в поле 2 либо в поле 3 (источники)."
                    )
                )
            checklist_structure = extract_checklist_svod(str(cpath))
            print(f"[Сводный акт] пунктов: {len(checklist_structure)}")

            try:
                extracted_data = await asyncio.to_thread(
                    process_checklist_svod,
                    api_key, all_texts, checklist_structure, active_model, svod_doc_text
                )
            except Exception as e:
                processing_status.update({"stage": "error", "message": str(e)[:200]})
                return ProcessingResult(status="error", message=f"Ошибка пайплайна (Сводный акт): {str(e)}")

            _normalize_counts(extracted_data, checklist_structure)

            output_filename = "Заполненный_Сводный_акт.docx"
            output_path = session_out_dir / output_filename
            output_ref = f"{session_id}/{output_filename}"
            processing_status.update({"stage": "fill", "message": "Заполнение Сводного акта...", "detail": ""})
            try:
                await asyncio.to_thread(
                    fill_svod_with_checklist,
                    str(cpath), extracted_data, str(output_path), checklist_structure
                )
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Ошибка заполнения Сводного акта: {str(e)}")

            # Лёгкая валидация: счётчики (validate_result заточен под План)
            results = extracted_data.get("checklist", [])
            validation = {
                "valid": True,
                "issues": [],
                "notes": [],
                "ok_count": sum(1 for it in results if it.get("ok")),
                "nok_count": sum(1 for it in results if it.get("nok")),
                "manual_count": sum(1 for it in results if not it.get("ok") and not it.get("nok")),
                "total": len(results),
            }

        # Сводка по этому чек-листу
        results = extracted_data.get("checklist", []) if isinstance(extracted_data, dict) else []
        outputs.append({
            "type": ctype,
            "checklist_file": cpath.name,
            "output_file": output_ref,
            "ok_count": sum(1 for it in results if it.get("ok")),
            "nok_count": sum(1 for it in results if it.get("nok")),
            "manual_count": sum(1 for it in results if not it.get("ok") and not it.get("nok")),
            "total_items": len(results),
            "header": extracted_data.get("header") if isinstance(extracted_data, dict) else None,
            "checklist": results,
            "validation": validation,
        })
        if first_extracted is None:
            first_extracted = extracted_data
            first_validation = validation

    # === Сохраняем мета-данные сессии для админки ===
    try:
        first = outputs[0] if outputs else {}
        _save_session_meta(
            session_id,
            finished_at=datetime.now().isoformat(timespec="seconds"),
            template_file=template_file,
            template_file_2=template_file_2,
            plan_doc_file=plan_doc_file,
            plan_doc_file_2=plan_doc_file_2,
            source_files=[d.name for d in source_docs if d.name not in field2_set],
            output_file=first.get("output_file"),
            outputs=outputs,
            model=active_model,
            ok_count=first.get("ok_count"),
            nok_count=first.get("nok_count"),
            total_items=first.get("total_items"),
            header=first_extracted.get("header") if isinstance(first_extracted, dict) else None,
            checklist=first_extracted.get("checklist") if isinstance(first_extracted, dict) else None,
            validation=first_validation,
        )
    except Exception as e:
        print(f"WARN: meta сохранить не удалось: {e}")

    types_done = ", ".join("Сводный акт" if o["type"] == "svod" else "План АУДИТА" for o in outputs)
    return ProcessingResult(
        status="success",
        message=f"Обработано чек-листов: {len(outputs)} ({types_done}). Проанализировано файлов: {len(analyzed_files)}",
        extracted_data=first_extracted,
        output_file=outputs[0]["output_file"] if outputs else None,
        outputs=outputs,
        validation=first_validation,
        analyzed_files=analyzed_files,
    )


@app.get("/api/admin/sessions")
async def admin_list_sessions(limit: int = 50):
    """Список сессий с краткой сводкой, отсортирован по времени убывания."""
    items = []
    for sdir in UPLOAD_DIR.iterdir():
        if not sdir.is_dir() or not _SESSION_ID_RE.match(sdir.name):
            continue
        meta = _load_session_meta(sdir.name)
        files = _list_session_files(sdir)
        ts = meta.get("finished_at") or meta.get("created_at") or ""
        if not ts:
            try:
                ts = datetime.fromtimestamp(sdir.stat().st_mtime).isoformat(timespec="seconds")
            except OSError:
                ts = ""
        items.append({
            "session_id": sdir.name,
            "created_at": meta.get("created_at"),
            "finished_at": meta.get("finished_at"),
            "ts_sort": ts,
            "files_count": len(files),
            "ok_count": meta.get("ok_count"),
            "nok_count": meta.get("nok_count"),
            "total_items": meta.get("total_items"),
            "model": meta.get("model"),
            "output_file": meta.get("output_file"),
            "applicant": (meta.get("header") or {}).get("Наименование Заявителя") if meta.get("header") else None,
        })
    items.sort(key=lambda x: x.get("ts_sort") or "", reverse=True)
    return {"sessions": items[:limit], "total": len(items)}


@app.get("/api/admin/sessions/{session_id}")
async def admin_get_session(session_id: str):
    """Детали сессии: meta + список всех файлов."""
    if not _SESSION_ID_RE.match(session_id or ""):
        raise HTTPException(status_code=400, detail="Некорректный session_id")
    sdir = UPLOAD_DIR / session_id
    if not sdir.is_dir():
        raise HTTPException(status_code=404, detail="Сессия не найдена")
    meta = _load_session_meta(session_id)
    files = _list_session_files(sdir)
    out_dir = OUTPUT_DIR / session_id
    output_files = []
    if out_dir.is_dir():
        for f in sorted(out_dir.iterdir()):
            if f.is_file():
                try:
                    output_files.append({"name": f.name, "size": f.stat().st_size,
                                         "ref": f"{session_id}/{f.name}"})
                except OSError:
                    pass
    return {"session_id": session_id, "meta": meta, "files": files, "output_files": output_files}


@app.get("/api/admin/sessions/{session_id}/zip")
async def admin_download_zip(session_id: str):
    """Скачать ZIP-архив со всеми файлами сессии + результатом + meta.json."""
    if not _SESSION_ID_RE.match(session_id or ""):
        raise HTTPException(status_code=400, detail="Некорректный session_id")
    sdir = UPLOAD_DIR / session_id
    if not sdir.is_dir():
        raise HTTPException(status_code=404, detail="Сессия не найдена")

    import zipfile, tempfile
    tmp = tempfile.NamedTemporaryFile(prefix=f"session_{session_id}_", suffix=".zip", delete=False)
    tmp.close()
    try:
        with zipfile.ZipFile(tmp.name, "w", zipfile.ZIP_DEFLATED) as zf:
            for f in sdir.iterdir():
                if f.is_file():
                    zf.write(f, arcname=f"inputs/{f.name}")
            out_dir = OUTPUT_DIR / session_id
            if out_dir.is_dir():
                for f in out_dir.iterdir():
                    if f.is_file():
                        zf.write(f, arcname=f"outputs/{f.name}")
        return FileResponse(
            path=tmp.name,
            filename=f"session_{session_id}.zip",
            media_type="application/zip",
        )
    except Exception as e:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass
        raise HTTPException(status_code=500, detail=f"Не удалось собрать zip: {e}")


@app.delete("/api/admin/sessions/{session_id}")
async def admin_delete_session(session_id: str):
    """Удалить папку сессии (входы) + папку результатов."""
    if not _SESSION_ID_RE.match(session_id or ""):
        raise HTTPException(status_code=400, detail="Некорректный session_id")
    sdir = UPLOAD_DIR / session_id
    out_dir = OUTPUT_DIR / session_id
    removed = []
    for d in (sdir, out_dir):
        if d.is_dir():
            try:
                shutil.rmtree(d)
                removed.append(str(d.name))
            except OSError as e:
                raise HTTPException(status_code=500, detail=f"Не удалось удалить {d}: {e}")
    return {"status": "ok", "removed": removed}


@app.get("/api/download/{filename:path}")
async def download_file(filename: str):
    """Скачивание готового файла. filename может быть '<session_id>/<file>'."""
    file_path = (OUTPUT_DIR / filename).resolve()
    # Защита от path traversal — file_path должен быть внутри OUTPUT_DIR.
    try:
        file_path.relative_to(OUTPUT_DIR.resolve())
    except ValueError:
        raise HTTPException(status_code=400, detail="Некорректный путь")
    if not file_path.exists() or not file_path.is_file():
        raise HTTPException(status_code=404, detail="Файл не найден")

    return FileResponse(
        path=str(file_path),
        filename=file_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/api/download-outputs/{session_id}")
async def download_outputs_zip(session_id: str):
    """ZIP только с заполненными чек-листами сессии (План и/или Сводный акт)."""
    if not _SESSION_ID_RE.match(session_id or ""):
        raise HTTPException(status_code=400, detail="Некорректный session_id")
    out_dir = OUTPUT_DIR / session_id
    if not out_dir.is_dir():
        raise HTTPException(status_code=404, detail="Результаты не найдены")

    import zipfile, tempfile
    tmp = tempfile.NamedTemporaryFile(prefix=f"outputs_{session_id}_", suffix=".zip", delete=False)
    tmp.close()
    try:
        count = 0
        with zipfile.ZipFile(tmp.name, "w", zipfile.ZIP_DEFLATED) as zf:
            for f in sorted(out_dir.iterdir()):
                if f.is_file() and f.suffix.lower() == ".docx":
                    zf.write(f, arcname=f.name)
                    count += 1
        if count == 0:
            os.unlink(tmp.name)
            raise HTTPException(status_code=404, detail="Нет заполненных файлов для скачивания")
        return FileResponse(
            path=tmp.name,
            filename="Заполненные_чек-листы.zip",
            media_type="application/zip",
        )
    except HTTPException:
        raise
    except Exception as e:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass
        raise HTTPException(status_code=500, detail=f"Не удалось собрать zip: {e}")


@app.get("/api/debug/extract")
async def debug_extract_text(filename: str, session_id: str):
    """Диагностика: возвращает текст, который бэк реально извлекает из файла
    (тот же текст, что уйдёт в модель). Помогает понять, упирается ли проблема
    в извлечение/OCR или в саму модель."""
    if not _SESSION_ID_RE.match(session_id):
        raise HTTPException(status_code=400, detail="Некорректный session_id")
    sdir = UPLOAD_DIR / session_id
    fpath = sdir / filename
    if not fpath.exists() or not fpath.is_file():
        raise HTTPException(status_code=404, detail="Файл не найден в сессии")
    suffix = fpath.suffix.lower()
    try:
        if suffix in ('.docx', '.docm'):
            text = extract_text_from_docx(str(fpath))
        elif suffix == '.pdf':
            text = extract_text_from_pdf(str(fpath))
        elif suffix == '.xlsx':
            text = extract_text_from_xlsx(str(fpath))
        else:
            raise HTTPException(status_code=400, detail=f"Неподдерживаемый формат: {suffix}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка извлечения: {e}")
    return {"filename": filename, "length": len(text), "text": text}


@app.get("/api/files")
async def list_files(session_id: Optional[str] = None):
    """Список загруженных файлов в сессии."""
    if not session_id or not _SESSION_ID_RE.match(session_id):
        raise HTTPException(status_code=400, detail="Некорректный session_id")
    sdir = UPLOAD_DIR / session_id
    if not sdir.is_dir():
        return {"files": []}
    files = []
    for f in sdir.iterdir():
        if f.is_file():
            files.append({
                "name": f.name,
                "size": f.stat().st_size,
                "type": "template" if "шаблон" in f.name.lower() or "план" in f.name.lower() else "source"
            })
    return {"files": files}


@app.delete("/api/files/{filename}")
async def delete_file(filename: str, session_id: Optional[str] = None):
    """Удаление загруженного файла из папки сессии."""
    if not session_id or not _SESSION_ID_RE.match(session_id):
        raise HTTPException(status_code=400, detail="Некорректный session_id")
    file_path = UPLOAD_DIR / session_id / filename
    if file_path.exists():
        file_path.unlink()
        return {"status": "ok", "message": f"Файл {filename} удален"}
    raise HTTPException(status_code=404, detail="Файл не найден")


@app.post("/api/upload-from-path")
async def upload_from_path(
    source_path: str = Form(...),
    session_id: Optional[str] = Form(None),
    block: Optional[str] = Form(None),
):
    """Загрузка документов из указанной папки (все форматы) в папку сессии."""
    if os.environ.get("ENABLE_LOCAL_UPLOAD", "0") != "1":
        raise HTTPException(status_code=404, detail="Endpoint disabled")
    path = Path(source_path)
    if not path.exists():
        raise HTTPException(status_code=400, detail="Путь не существует")

    sid, sdir = _resolve_session_dir(session_id)

    meta = _load_session_meta(sid)
    if not meta.get("created_at"):
        _save_session_meta(sid, created_at=datetime.now().isoformat(timespec="seconds"))
    block_files = list(meta.get("blocks", {}).get(block, [])) if block else []

    uploaded = []
    extensions = ['*.docx', '*.docm', '*.pdf', '*.xlsx']
    for ext in extensions:
        for doc_file in path.glob(ext):
            dest = sdir / doc_file.name
            shutil.copy2(str(doc_file), str(dest))
            uploaded.append({
                "filename": doc_file.name,
                "path": str(dest),
                "size": dest.stat().st_size
            })
            if block and doc_file.name not in block_files:
                block_files.append(doc_file.name)

    if block:
        blocks = dict(meta.get("blocks") or {})
        blocks[block] = block_files
        _save_session_meta(sid, blocks=blocks)

    return {"status": "ok", "session_id": sid, "uploaded_files": uploaded, "count": len(uploaded)}


@app.get("/api/status")
def get_processing_status():
    """Текущий статус обработки (для polling из UI во время /api/process)."""
    return processing_status


@app.get("/api/gigachat/diagnose")
def diagnose_gigachat():
    """
    Диагностика связи с GigaChat:
      - OAuth на 9443
      - Список доступных моделей
      - Проверка что DEFAULT_GIGACHAT_MODEL есть в списке
    Читает ключ из сохранённых настроек.
    """
    settings = _load_settings()
    key = settings.get("api_key", "")
    active_model = _get_active_model()
    if not key:
        return {"ok": False, "stage": "no_key",
                "detail": "API ключ не сохранён. Сначала сохраните ключ в настройках.",
                "models": [], "current_model": active_model}
    return gigachat_preflight(key, timeout=8.0, model=active_model)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
