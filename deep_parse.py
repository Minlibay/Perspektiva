#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Deep analysis of .docx XML structure - looking for form fields, macros, checkboxes etc.
"""

import zipfile
import xml.etree.ElementTree as ET
import json
import re

DOC_PATH = r"D:\Perpektiva\Пакет 2\ИИ -ЧК -План АУДИТА.docx"

def deep_xml_analysis():
    """Deep analysis of all XML parts in the .docx"""
    
    print("=" * 80)
    print("ГЛУБОКИЙ АНАЛИЗ XML СТРУКТУРЫ")
    print("=" * 80)
    
    with zipfile.ZipFile(DOC_PATH, 'r') as z:
        namelist = z.namelist()
        print(f"\nФайлы в архиве .docx:")
        for name in namelist:
            size = len(z.read(name))
            print(f"  {name} ({size} bytes)")
        
        # Parse document.xml
        print("\n" + "=" * 80)
        print("АНАЛИЗ word/document.xml")
        print("=" * 80)
        
        xml_content = z.read('word/document.xml')
        root = ET.fromstring(xml_content)
        
        ns = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
            'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
            'o': 'urn:schemas-microsoft-com:office:office',
            'v': 'urn:schemas-microsoft-com:vml',
        }
        
        # Count elements
        all_elements = list(root.iter())
        print(f"\nВсего XML элементов: {len(all_elements)}")
        
        # Element types
        tag_counts = {}
        for elem in all_elements:
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            tag_counts[tag] = tag_counts.get(tag, 0) + 1
        
        print("\nРаспределение по типам:")
        for tag, count in sorted(tag_counts.items(), key=lambda x: -x[1])[:30]:
            print(f"  {tag}: {count}")
        
        # Search for specific form-related elements
        print("\n" + "=" * 80)
        print("ПОИСК ФОРМ-ЭЛЕМЕНТОВ")
        print("=" * 80)
        
        form_tags = ['fldChar', 'ffData', 'formField', 'checkBox', 'textInput', 
                     'ddList', 'calculate', 'textForm', 'sdt', 'contentPart',
                     'control', 'object', 'OLE']
        
        for tag in form_tags:
            found = list(root.iter(f'{{{ns["w"]}}}{tag}'))
            if found:
                print(f"\n  Найдено '{tag}': {len(found)} шт.")
                for i, elem in enumerate(found[:5]):
                    attrs = {k: v for k, v in elem.attrib.items()}
                    print(f"    [{i}] {attrs}")
        
        # Look for complex fields (fldSimple)
        print("\n" + "=" * 80)
        print("ПОИСК FLDCHAR (поля форм)")
        print("=" * 80)
        
        fld_chars = list(root.iter(f'{{{ns["w"]}}}fldChar'))
        print(f"  Найдено fldChar: {len(fld_chars)}")
        
        # Find all runs with fldChar and their context
        for i, fld in enumerate(fld_chars):
            fldCharType = fld.get(f'{{{ns["w"]}}}fldCharType', '')
            parent = fld.getparent()
            grandparent = parent.getparent() if parent is not None else None
            
            # Get surrounding text
            all_text_elems = list(root.iter(f'{{{ns["w"]}}}t'))
            # Find nearby text
            run_text = []
            if parent is not None:
                for t in parent.iter(f'{{{ns["w"]}}}t'):
                    if t.text:
                        run_text.append(t.text)
            
            print(f"  [{i}] type={fldCharType}, text={''.join(run_text)[:80]}")
        
        # Look for checkbox symbols or form characters
        print("\n" + "=" * 80)
        print("ПОИСК СПЕЦИАЛЬНЫХ СИМВОЛОВ (checkboxes, symbols)")
        print("=" * 80)
        
        sym_elements = list(root.iter(f'{{{ns["w"]}}}sym'))
        print(f"  Найдено sym элементов: {len(sym_elements)}")
        for i, sym in enumerate(sym_elements[:20]):
            font = sym.get(f'{{{ns["w"]}}}font', '')
            char = sym.get(f'{{{ns["w"]}}}char', '')
            print(f"  [{i}] font={font}, char={char}")
        
        # Look for proofState (macros, forms protection)
        print("\n" + "=" * 80)
        print("ЗАЩИТА И НАСТРОЙКИ ДОКУМЕНТА")
        print("=" * 80)
        
        # Check settings.xml
        if 'word/settings.xml' in namelist:
            settings_xml = z.read('word/settings.xml')
            settings_root = ET.fromstring(settings_xml)
            
            # Document protection
            doc_protect = settings_root.find(f'{{{ns["w"]}}}documentProtection')
            if doc_protect is not None:
                print(f"  Защита документа: {dict(doc_protect.attrib)}")
            
            # Proof state
            proof = settings_root.find(f'{{{ns["w"]}}}proofState')
            if proof is not None:
                print(f"  Proof state: {dict(proof.attrib)}")
        
        # Look for vbaProject.bin (macros)
        if 'word/vbaProject.bin' in namelist:
            print("\n  *** ОБНАРУЖЕНЫ МАКРОСЫ (vbaProject.bin) ***")
        
        # Check for custom XML / metadata
        print("\n" + "=" * 80)
        print("ПОЛЬЗОВАТЕЛЬСКИЕ СВОЙСТВА ДОКУМЕНТА")
        print("=" * 80)
        
        if 'docProps/core.xml' in namelist:
            core_xml = z.read('docProps/core.xml')
            print(f"  Core XML:")
            core_text = core_xml.decode('utf-8')[:500]
            for line in core_text.split('\n')[:10]:
                print(f"    {line.strip()}")
        
        if 'docProps/custom.xml' in namelist:
            custom_xml = z.read('docProps/custom.xml')
            print(f"\n  Custom XML:")
            custom_text = custom_xml.decode('utf-8')
            for line in custom_text.split('\n')[:15]:
                print(f"    {line.strip()}")
        
        # Look for header/footer
        print("\n" + "=" * 80)
        print("ЗАГОЛОВКИ И ПОДВАЛЫ")
        print("=" * 80)
        
        header_files = [n for n in namelist if 'header' in n.lower()]
        footer_files = [n for n in namelist if 'footer' in n.lower()]
        
        if header_files:
            print(f"  Header файлы: {header_files}")
            for hf in header_files:
                hf_content = z.read(hf).decode('utf-8')
                # Extract text content
                texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', hf_content)
                print(f"    {hf}: {' '.join(texts)[:100]}")
        
        if footer_files:
            print(f"  Footer файлы: {footer_files}")
        
        # Look for numberings (for checklists)
        print("\n" + "=" * 80)
        print("НУМЕРАЦИЯ И МАРКЕРЫ")
        print("=" * 80)
        
        if 'word/numbering.xml' in namelist:
            num_xml = z.read('word/numbering.xml')
            print(f"  Numbering XML size: {len(num_xml)} bytes")
            num_root = ET.fromstring(num_xml)
            abstract_nums = num_root.findall(f'{{{ns["w"]}}}abstractNum')
            nums = num_root.findall(f'{{{ns["w"]}}}num')
            print(f"  AbstractNum: {len(abstract_nums)}, Num: {len(nums)}")


def extract_full_text_with_structure():
    """Extract full text with paragraph/table structure mapping"""
    
    import docx
    doc = docx.Document(DOC_PATH)
    
    print("\n" + "=" * 80)
    print("ПОЛНАЯ ТЕКСТОВАЯ СТРУКТУРА ДОКУМЕНТА")
    print("=" * 80)
    
    # All text content
    full_text = []
    
    # From paragraphs
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            full_text.append({
                "type": "paragraph",
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else None
            })
    
    # From tables  
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            row_text = []
            for c_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    row_text.append(f"[{c_idx}]: {cell.text.strip()}")
            if row_text:
                full_text.append({
                    "type": "table_row",
                    "table": t_idx,
                    "row": r_idx,
                    "cells": " | ".join(row_text)
                })
    
    # Print structured view
    for item in full_text:
        if item["type"] == "paragraph":
            style_marker = f" [{item['style']}]" if item.get('style') and item['style'] != 'Normal' else ""
            print(f"\n  ПАРАГРАФ{style_marker}: {item['text']}")
        else:
            print(f"\n  ТАБЛИЦА {item['table']}, строка {item['row']}:")
            print(f"    {item['cells']}")


if __name__ == "__main__":
    deep_xml_analysis()
    extract_full_text_with_structure()
