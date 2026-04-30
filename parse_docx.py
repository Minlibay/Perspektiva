#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script to parse and analyze the structure of "ИИ -ЧК -План АУДИТА.docx"
"""

import docx
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import xml.etree.ElementTree as ET
import json
import os

DOC_PATH = r"D:\Perpektiva\Пакет 2\ИИ -ЧК -План АУДИТА.docx"

def parse_paragraphs(doc):
    """Extract all paragraphs with detailed formatting info"""
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        p_info = {
            "index": i,
            "text": para.text,
            "style": para.style.name if para.style else None,
            "alignment": str(para.alignment) if para.alignment else None,
            "runs": []
        }
        for run in para.runs:
            run_info = {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_size": str(run.font.size) if run.font.size else None,
                "font_name": run.font.name,
            }
            p_info["runs"].append(run_info)
        paragraphs.append(p_info)
    return paragraphs


def parse_tables(doc):
    """Extract all tables with cell-level detail"""
    tables = []
    for t_idx, table in enumerate(doc.tables):
        t_info = {
            "table_index": t_idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "cells": []
        }
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_info = {
                    "row": r_idx,
                    "col": c_idx,
                    "text": cell.text.strip(),
                    "paragraphs": [p.text for p in cell.paragraphs],
                    "merged": cell._element.get(qn('w:gridSpan')) is not None,
                }
                # Check for content controls
                for elem in cell._element.iter():
                    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                    if tag == 'sdt':
                        cell_info["has_content_control"] = True
                        alias_elem = elem.find(f'.//{{{nsdecls("w")}}}alias')
                        if alias_elem is not None:
                            cell_info["content_control_alias"] = alias_elem.get(qn('w:val'), '')
                    if tag == 'checkBox':
                        cell_info["has_checkbox"] = True
                    if tag == 'text' or tag == 'textForm':
                        cell_info["has_text_form"] = True
                    if tag == 'ddList':
                        cell_info["has_dropdown"] = True
                t_info["cells"].append(cell_info)
        tables.append(t_info)
    return tables


def parse_content_controls_raw():
    """Parse raw XML to find structured document tags (SDT / content controls)"""
    content_controls = []
    try:
        with zipfile.ZipFile(DOC_PATH, 'r') as z:
            # Parse document.xml
            if 'word/document.xml' in z.namelist():
                xml_content = z.read('word/document.xml')
                root = ET.fromstring(xml_content)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                for sdt in root.iter(f'{{{ns["w"]}}}sdt'):
                    ctrl = {}
                    # Get alias (tag name)
                    alias_elem = sdt.find(f'w:sdtPr/w:alias', ns)
                    if alias_elem is not None:
                        ctrl['alias'] = alias_elem.get(qn('w:val'), '')
                    
                    # Get tag
                    tag_elem = sdt.find(f'w:sdtPr/w:tag', ns)
                    if tag_elem is not None:
                        ctrl['tag'] = tag_elem.get(qn('w:val'), '')
                    
                    # Get id
                    id_elem = sdt.find(f'w:sdtPr/w:id', ns)
                    if id_elem is not None:
                        ctrl['id'] = id_elem.get(qn('w:val'), '')
                    
                    # Determine type
                    if sdt.find(f'w:sdtPr/w:checkBox', ns) is not None:
                        ctrl['type'] = 'checkbox'
                    elif sdt.find(f'w:sdtPr/w:ddList', ns) is not None:
                        ctrl['type'] = 'dropdown'
                        list_items = []
                        for item in sdt.findall(f'w:sdtPr/w:ddList/w:ddListItem', ns):
                            list_items.append(item.text)
                        ctrl['list_items'] = list_items
                    elif sdt.find(f'w:sdtPr/w:date', ns) is not None:
                        ctrl['type'] = 'date'
                    elif sdt.find(f'w:sdtPr/w:text', ns) is not None:
                        ctrl['type'] = 'text'
                    else:
                        ctrl['type'] = 'rich_text'
                    
                    # Get content text
                    content_elem = sdt.find(f'w:sdtContent', ns)
                    if content_elem is not None:
                        texts = content_elem.findall(f'.//w:t', ns)
                        ctrl['content'] = ' '.join([t.text for t in texts if t.text])
                    
                    content_controls.append(ctrl)
    except Exception as e:
        print(f"Error parsing content controls: {e}")
    
    return content_controls


def parse_form_fields_raw():
    """Parse raw XML to find legacy form fields"""
    form_fields = []
    try:
        with zipfile.ZipFile(DOC_PATH, 'r') as z:
            if 'word/document.xml' in z.namelist():
                xml_content = z.read('word/document.xml')
                root = ET.fromstring(xml_content)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                for ffdata in root.iter(f'{{{ns["w"]}}}ffData'):
                    field = {}
                    name_elem = ffdata.find(f'w:name', ns)
                    if name_elem is not None:
                        field['name'] = name_elem.get(qn('w:val'), '')
                    
                    if ffdata.find(f'w:checkBox', ns) is not None:
                        field['type'] = 'checkbox'
                        checked = ffdata.find(f'w:checkBox/w:default', ns)
                        if checked is not None:
                            field['default'] = checked.get(qn('w:val'), '0')
                    elif ffdata.find(f'w:textInput', ns) is not None:
                        field['type'] = 'text'
                        default_elem = ffdata.find(f'w:textInput/w:default', ns)
                        if default_elem is not None:
                            field['default'] = default_elem.get(qn('w:val'), '')
                    elif ffdata.find(f'w:ddList', ns) is not None:
                        field['type'] = 'dropdown'
                        items = []
                        for item in ffdata.findall(f'w:ddList/w:ddListItem', ns):
                            items.append(item.text)
                        field['list_items'] = items
                    
                    form_fields.append(field)
    except Exception as e:
        print(f"Error parsing form fields: {e}")
    
    return form_fields


def analyze_document_structure():
    """Main analysis function"""
    print("=" * 80)
    print("АНАЛИЗ ДОКУМЕНТА: ИИ -ЧК -План АУДИТА.docx")
    print("=" * 80)
    
    doc = docx.Document(DOC_PATH)
    
    # 1. Paragraphs analysis
    print("\n" + "=" * 80)
    print("РАЗДЕЛ 1: ПАРАГРАФЫ ДОКУМЕНТА")
    print("=" * 80)
    paragraphs = parse_paragraphs(doc)
    for p in paragraphs:
        if p['text'].strip():
            style_info = f" [{p['style']}]" if p['style'] else ""
            print(f"  [{p['index']}] {p['text'][:120]}{style_info}")
    
    # 2. Tables analysis
    print("\n" + "=" * 80)
    print("РАЗДЕЛ 2: ТАБЛИЦЫ")
    print("=" * 80)
    tables = parse_tables(doc)
    for t in tables:
        print(f"\n  Таблица {t['table_index'] + 1}: {t['rows']} строк x {t['cols']} колонок")
        for cell in t['cells']:
            extra = ""
            if cell.get('has_content_control'):
                extra += f" [ContentControl: {cell.get('content_control_alias', '')}]"
            if cell.get('has_checkbox'):
                extra += " [CHECKBOX]"
            if cell.get('has_text_form'):
                extra += " [TEXT_FORM]"
            if cell.get('has_dropdown'):
                extra += " [DROPDOWN]"
            if cell['text']:
                print(f"    [{cell['row']},{cell['col']}] {cell['text'][:80]}{extra}")
            elif extra:
                print(f"    [{cell['row']},{cell['col']}] (пусто){extra}")
    
    # 3. Content Controls (SDT)
    print("\n" + "=" * 80)
    print("РАЗДЕЛ 3: ЭЛЕМЕНТЫ УПРАВЛЕНИЯ СОДЕРЖИМЫМ (Content Controls / SDT)")
    print("=" * 80)
    content_controls = parse_content_controls_raw()
    for i, ctrl in enumerate(content_controls):
        print(f"  [{i+1}] Тип: {ctrl.get('type', 'unknown')}")
        if ctrl.get('alias'):
            print(f"       Название: {ctrl['alias']}")
        if ctrl.get('tag'):
            print(f"       Тег: {ctrl['tag']}")
        if ctrl.get('content'):
            print(f"       Содержимое: {ctrl['content'][:80]}")
        if ctrl.get('list_items'):
            print(f"       Варианты: {ctrl['list_items']}")
        print()
    
    # 4. Legacy Form Fields
    print("\n" + "=" * 80)
    print("РАЗДЕЛ 4: УСТАРЕВШИЕ ПОЛЯ ФОРМ (Legacy Form Fields)")
    print("=" * 80)
    form_fields = parse_form_fields_raw()
    for i, field in enumerate(form_fields):
        print(f"  [{i+1}] Имя: {field.get('name', 'N/A')}, Тип: {field.get('type', 'unknown')}")
        if field.get('default'):
            print(f"       По умолчанию: {field['default']}")
        if field.get('list_items'):
            print(f"       Варианты: {field['list_items']}")
        print()
    
    # 5. Full JSON dump for reference
    print("\n" + "=" * 80)
    print("РАЗДЕЛ 5: ПОЛНАЯ СТРУКТУРА (JSON)")
    print("=" * 80)
    
    full_structure = {
        "paragraphs": paragraphs,
        "tables": tables,
        "content_controls": content_controls,
        "form_fields": form_fields
    }
    
    output_path = r"D:\Perpektiva\Web\docx_structure.json"
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(full_structure, f, ensure_ascii=False, indent=2)
    print(f"\n  Полная структура сохранена в: {output_path}")
    
    return full_structure


if __name__ == "__main__":
    analyze_document_structure()
