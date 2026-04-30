#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Final comprehensive structure extraction
"""

import docx
import json

DOC_PATH = r"D:\Perpektiva\Пакет 2\ИИ -ЧК -План АУДИТА.docx"

doc = docx.Document(DOC_PATH)

print("=" * 100)
print("ПОЛНАЯ СТРУКТУРА ДОКУМЕНТА: План АУДИТА (ИИ-ЧК)")
print("=" * 100)

# ==================== ТАБЛИЦА 0: Шапка ====================
print("\n" + "=" * 100)
print("ТАБЛИЦА 0 — ШАПКА ДОКУМЕНТА (4 строки × 2 колонки)")
print("=" * 100)

table0 = doc.tables[0]
for r_idx, row in enumerate(table0.rows):
    for c_idx, cell in enumerate(row.cells):
        text = cell.text.strip()
        print(f"  [{r_idx},{c_idx}] LABEL: '{text}'" if text else f"  [{r_idx},{c_idx}] VALUE: (пусто — поле для ввода)")

# ==================== ТАБЛИЦА 1: Основная таблица аудита ====================
print("\n" + "=" * 100)
print("ТАБЛИЦА 1 — ОСНОВНАЯ ТАБЛИЦА ПРОВЕРКИ (17 строк × 5 колонок)")
print("=" * 100)

table1 = doc.tables[1]

# Full cell content
for r_idx, row in enumerate(table1.rows):
    cells = [cell.text.strip() for cell in row.cells]
    col0 = cells[0] if len(cells) > 0 else ""
    col1 = cells[1] if len(cells) > 1 else ""
    col2 = cells[2] if len(cells) > 2 else ""
    col3 = cells[3] if len(cells) > 3 else ""
    col4 = cells[4] if len(cells) > 4 else ""
    
    print(f"\n  --- СТРОКА {r_idx} ---")
    print(f"  [0] Область проверки:     {col0[:80]}")
    print(f"  [1] Комментарии:          {col1[:80]}")
    print(f"  [2] ОК:                   {col2}  <- чекбокс")
    print(f"  [3] NOK:                  {col3}  <- чекбокс")
    print(f"  [4] Проблемные зоны:      {col4[:100]}")

# ==================== Определить структуру веб-формы ====================
print("\n" + "=" * 100)
print("СТРУКТУРА ВЕБ-ФОРМЫ (вывод для разработки)")
print("=" * 100)

web_form_structure = {
    "document_title": "План АУДИТА (ИИ-ЧК)",
    "sections": {
        "header": {
            "title": "Шапка документа",
            "fields": [
                {"id": "applicant_name", "label": "Наименования Заявителя", "type": "text", "required": True},
                {"id": "audit_type", "label": "Вид аудита", "type": "select", "options": ["Аудит", "Расширение", "Дополнительный аудит"], "required": True},
                {"id": "audit_dates", "label": "Даты проведения", "type": "date_range", "required": True},
                {"id": "reg_number", "label": "РЭГ", "type": "text", "required": True}
            ]
        },
        "audit_checklist": {
            "title": "Основная таблица проверки",
            "columns": [
                {"id": "check_area", "label": "Область проверки", "type": "label"},
                {"id": "comments", "label": "Комментарии", "type": "textarea"},
                {"id": "ok", "label": "ОК", "type": "checkbox"},
                {"id": "nok", "label": "NOK", "type": "checkbox"},
                {"id": "problem_zones", "label": "Проблемные зоны", "type": "textarea"}
            ],
            "rows": []
        }
    }
}

# Extract rows for web form
row_data = []
for r_idx, row in enumerate(table1.rows):
    cells = [cell.text.strip() for cell in row.cells]
    row_data.append({
        "row_index": r_idx,
        "check_area": cells[0] if len(cells) > 0 else "",
        "comments": cells[1] if len(cells) > 1 else "",
        "ok": "",
        "nok": "",
        "problem_zones": cells[4] if len(cells) > 4 else ""
    })

web_form_structure["sections"]["audit_checklist"]["rows"] = row_data

# Save full JSON
output_path = r"D:\Perpektiva\Web\audit_plan_full_structure.json"
with open(output_path, 'w', encoding='utf-8') as f:
    json.dump(web_form_structure, f, ensure_ascii=False, indent=2)

print(f"\nПолная JSON-структура сохранена: {output_path}")

# Print row summaries
print("\n" + "=" * 100)
print("СВОДКА СТРОК ПРОВЕРОЧНОЙ ТАБЛИЦЫ")  
print("=" * 100)

for r in row_data:
    print(f"\n  Строка {r['row_index']}: {r['check_area'][:70]}")
    if r['comments']:
        print(f"    → Что проверять: {r['comments'][:80]}")
    if r['problem_zones']:
        print(f"    → Подсказка: {r['problem_zones'][:80]}")
